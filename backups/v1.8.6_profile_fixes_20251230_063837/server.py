from fastapi import FastAPI, APIRouter, HTTPException, UploadFile, File, Form, BackgroundTasks, Response, Request, Depends
from fastapi.responses import JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict, EmailStr, model_validator
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone, timedelta
import base64
from emergentintegrations.llm.chat import LlmChat, UserMessage, ImageContent
from emergentintegrations.payments.stripe.checkout import StripeCheckout, CheckoutSessionResponse, CheckoutStatusResponse, CheckoutSessionRequest
import json
import re
import hashlib
import time
import subprocess
import asyncio
import unicodedata
import bcrypt
import jwt
import resend
import secrets
from backup_manager import BackupManager, create_startup_backup, stop_backup_task

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ.get('MONGO_URL', 'mongodb://localhost:27017')
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ.get('DB_NAME')]

# LLM API Key
EMERGENT_LLM_KEY = os.environ.get('EMERGENT_LLM_KEY', '')

# Stripe API Key
STRIPE_API_KEY = os.environ.get('STRIPE_API_KEY', '')

# JWT Secret Key
JWT_SECRET = os.environ.get('JWT_SECRET')
JWT_ALGORITHM = "HS256"
JWT_EXPIRY_DAYS = 7

# Resend Email Configuration
RESEND_API_KEY = os.environ.get('RESEND_API_KEY', '')
SENDER_EMAIL = os.environ.get('SENDER_EMAIL', 'noreply@wine-pairing.online')
FRONTEND_URL = os.environ.get('FRONTEND_URL', os.environ.get('FRONTEND_BASE_URL', ''))

# Initialize Resend
if RESEND_API_KEY:
    resend.api_key = RESEND_API_KEY

# Backup Manager (wird beim Startup initialisiert)
backup_manager: BackupManager = None

# ===================== FREEMIUM CONFIGURATION =====================
FREEMIUM_LIMITS = {
    "basic": {
        "pairing_requests_per_day": 5,
        "chat_messages_per_day": 5,
        "max_cellar_wines": 10,
        "max_favorites": 10
    },
    "pro": {
        "pairing_requests_per_day": float('inf'),
        "chat_messages_per_day": float('inf'),
        "max_cellar_wines": float('inf'),
        "max_favorites": float('inf')
    }
}

SUBSCRIPTION_PLANS = {
    "pro_monthly": {"price": 4.99, "currency": "eur", "interval": "month"},
    "pro_yearly": {"price": 39.99, "currency": "eur", "interval": "year"}
}

# ===================== ACCENT-INSENSITIVE SEARCH HELPER =====================
# WICHTIG: Diese Funktion muss für alle Suchfunktionen verwendet werden!
# Problem: "Chateau" muss "Château" finden, "Cotes" muss "Côtes" finden

def create_accent_insensitive_pattern(search_term: str) -> str:
    """
    Erstellt ein Regex-Pattern, das Akzente ignoriert.
    z.B. "chateau" -> "[cç]h[aàâäã]t[eéèêë][aàâäã][uùûüú]"
    
    MUSS bei allen Suchen verwendet werden, um französische Weine zu finden!
    """
    # Normalisiere den Suchbegriff (entferne Akzente)
    normalized = ''.join(
        c for c in unicodedata.normalize('NFD', search_term.strip())
        if unicodedata.category(c) != 'Mn'
    )
    
    # Ersetze Buchstaben durch Akzent-tolerante Patterns
    ACCENT_REPLACEMENTS = {
        'a': '[aàâäãá]',
        'e': '[eéèêë]',
        'i': '[iîïí]',
        'o': '[oôöóò]',
        'u': '[uùûüú]',
        'c': '[cç]',
        'n': '[nñ]',
        'y': '[yÿý]',
    }
    
    pattern = ''
    for char in normalized.lower():
        pattern += ACCENT_REPLACEMENTS.get(char, re.escape(char))
    
    return pattern

# ===================== PAIRING CACHE =====================
# In-memory cache for wine pairing recommendations
# Cache TTL: 24 hours (86400 seconds)
PAIRING_CACHE: Dict[str, Dict[str, Any]] = {}
CACHE_TTL = 86400  # 24 hours in seconds

def get_cache_key(dish: str, language: str, wine_type_filter: Optional[str] = None, use_cellar: bool = False) -> str:
    """Generate a unique cache key for a pairing request"""
    # Normalize dish name: lowercase, strip whitespace
    normalized_dish = dish.lower().strip()
    # Create a unique key based on dish, language, wine type filter, AND use_cellar
    key_parts = [normalized_dish, language]
    if wine_type_filter and wine_type_filter != 'all':
        key_parts.append(wine_type_filter)
    # WICHTIG: use_cellar muss im Cache-Key sein, sonst werden gecachte Empfehlungen
    # ohne Keller-Weine auch für use_cellar=true zurückgegeben!
    if use_cellar:
        key_parts.append("cellar")
    key_string = "|".join(key_parts)
    return hashlib.md5(key_string.encode()).hexdigest()

def get_cached_pairing(cache_key: str) -> Optional[Dict[str, Any]]:
    """Get a cached pairing if it exists and hasn't expired"""
    if cache_key in PAIRING_CACHE:
        cached = PAIRING_CACHE[cache_key]
        if time.time() - cached['timestamp'] < CACHE_TTL:
            logging.info(f"Cache HIT for key: {cache_key[:8]}...")
            return cached['data']
        else:
            # Expired, remove from cache
            del PAIRING_CACHE[cache_key]
            logging.info(f"Cache EXPIRED for key: {cache_key[:8]}...")
    return None

def set_cached_pairing(cache_key: str, data: Dict[str, Any]) -> None:
    """Store a pairing in the cache"""
    PAIRING_CACHE[cache_key] = {
        'timestamp': time.time(),
        'data': data
    }
    logging.info(f"Cache SET for key: {cache_key[:8]}... (Total cached: {len(PAIRING_CACHE)})")

def clear_old_cache_entries() -> int:
    """Remove expired entries from cache"""
    current_time = time.time()
    keys_to_remove = [
        key for key, value in PAIRING_CACHE.items()
        if current_time - value['timestamp'] >= CACHE_TTL
    ]
    for key in keys_to_remove:
        del PAIRING_CACHE[key]
    return len(keys_to_remove)

# ===================== AUTO-ADD RECOMMENDED WINES =====================
# Automatisch empfohlene Weine zur Datenbank hinzufügen

async def extract_wine_names_from_recommendation(recommendation: str) -> List[str]:
    """
    Extrahiert Weinnamen aus der Pairing-Empfehlung.
    Sucht nach **Weinname** Pattern im Markdown.
    """
    # Pattern für fettgedruckte Weinnamen: **Weinname**
    pattern = r'\*\*([^*]+(?:Château|Tignanello|Barolo|Barbaresco|Champagne|Clos|Domaine|Weingut|Schloss|Maison)[^*]*)\*\*|\*\*([A-Z][^*]{5,60})\*\*'
    
    matches = re.findall(r'\*\*([^*]{10,80})\*\*', recommendation)
    
    wine_names = []
    # Filter: Nur echte Weinnamen, keine Überschriften
    skip_keywords = ['HAUPTEMPFEHLUNG', 'TOP RECOMMENDATION', 'RECOMMANDATION', 
                     'Alternative', 'Option', 'Weintyp', 'Wine Type', 'Rotwein', 
                     'Weißwein', 'Schaumwein', 'Sparkling', 'festliche', 'leichterer']
    
    for match in matches:
        # Überschriften und Kategorien überspringen
        if any(skip in match for skip in skip_keywords):
            continue
        # Zu kurze Namen überspringen
        if len(match) < 10:
            continue
        # Bereits gefundene überspringen
        if match not in wine_names:
            wine_names.append(match.strip())
    
    return wine_names[:6]  # Maximal 6 Weine

async def check_wine_exists(wine_name: str) -> bool:
    """Prüft ob ein Wein bereits in der Datenbank existiert"""
    # Normalisiere den Namen für die Suche
    search_pattern = create_accent_insensitive_pattern(wine_name.split(',')[0].split('(')[0].strip())
    
    existing = await db.public_wines.find_one({
        "name": {"$regex": search_pattern, "$options": "i"}
    })
    
    return existing is not None

async def generate_wine_entry(wine_name: str, dish_context: str = "") -> Optional[dict]:
    """
    Generiert einen vollständigen Wein-Eintrag mit Claude.
    """
    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=str(uuid.uuid4()),
            system_message="""Du bist ein Wein-Experte. Generiere einen strukturierten Wein-Eintrag im JSON-Format.
ANTWORTE NUR MIT VALIDEM JSON, KEIN ANDERER TEXT!"""
        ).with_model("openai", "gpt-5.1")
        
        prompt = f"""Erstelle einen Datenbank-Eintrag für diesen Wein: "{wine_name}"

ANTWORTE NUR MIT DIESEM JSON-FORMAT (keine Erklärung, nur JSON):
{{
  "name": "Vollständiger Weinname",
  "winery": "Weingut/Produzent",
  "grape_variety": "Hauptrebsorte",
  "region": "Region (z.B. Bordeaux, Toskana, Mosel)",
  "country": "Land",
  "year": 2022,
  "wine_color": "rot/weiß/rosé",
  "price_category": "€€€",
  "description_de": "Kurze deutsche Beschreibung (1-2 Sätze, emotionaler Stil)",
  "description_en": "Short English description",
  "description_fr": "Courte description française",
  "food_pairings_de": ["Passende Speise 1", "Passende Speise 2"],
  "food_pairings_en": ["Food pairing 1", "Food pairing 2"],
  "food_pairings_fr": ["Accord 1", "Accord 2"]
}}"""

        response = await chat.send_message(UserMessage(text=prompt))
        
        # Parse JSON aus der Antwort
        json_match = re.search(r'\{[\s\S]*\}', response)
        if json_match:
            wine_data = json.loads(json_match.group())
            
            # Füge System-Felder hinzu
            wine_data["id"] = str(uuid.uuid4())
            wine_data["created_at"] = datetime.now(timezone.utc).isoformat().replace('+00:00', 'Z')
            wine_data["updated_at"] = datetime.now(timezone.utc).isoformat().replace('+00:00', 'Z')
            wine_data["auto_generated"] = True  # Markierung für automatisch generierte Einträge
            
            return wine_data
        
        return None
        
    except Exception as e:
        logger.error(f"Error generating wine entry for {wine_name}: {e}")
        return None

async def auto_add_recommended_wines(recommendation: str, dish: str):
    """
    Background-Task: Extrahiert empfohlene Weine und fügt fehlende zur DB hinzu.
    """
    try:
        wine_names = await extract_wine_names_from_recommendation(recommendation)
        logger.info(f"🍷 Auto-Add: Found {len(wine_names)} wine names in recommendation")
        
        added_count = 0
        for wine_name in wine_names:
            # Prüfen ob Wein bereits existiert
            exists = await check_wine_exists(wine_name)
            
            if not exists:
                logger.info(f"🍷 Auto-Add: Generating entry for '{wine_name}'...")
                wine_entry = await generate_wine_entry(wine_name, dish)
                
                if wine_entry:
                    await db.public_wines.insert_one(wine_entry)
                    added_count += 1
                    logger.info(f"✅ Auto-Add: Added '{wine_entry.get('name')}' to database")
            else:
                logger.info(f"📌 Auto-Add: '{wine_name}' already exists in database")
        
        if added_count > 0:
            logger.info(f"🎉 Auto-Add: Successfully added {added_count} new wines to database")
            
    except Exception as e:
        logger.error(f"❌ Auto-Add error: {e}")

# Create the main app
app = FastAPI(title="Wine Pairing API", version="1.0.0")

# Add CORS middleware for production deployment
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, this will be configured properly
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")


# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# ===================== MODELS =====================

# ===================== AUTH & USER MODELS =====================

# Wine Profile Model for personalized recommendations
class WineProfile(BaseModel):
    """User's wine taste profile for personalized AI recommendations"""
    model_config = ConfigDict(extra="ignore")
    
    # Rotwein-Stilistik
    red_wine_style: Optional[str] = None  # "kraftig_wurzig", "fruchtig_elegant", "beides"
    
    # Weißwein-Charakter
    white_wine_style: Optional[str] = None  # "mineralisch_frisch", "cremig_textur", "aromatisch_verspielt", "beides"
    
    # Struktur-Präferenzen
    acidity_tolerance: Optional[str] = None  # "niedrig", "mittel", "hoch"
    tannin_preference: Optional[str] = None  # "weich_seidig", "mittel", "markant_griffig"
    
    # Süßegrad
    sweetness_preference: Optional[str] = None  # "knochentrocken", "trocken", "halbtrocken", "lieblich", "edelsuss"
    
    # Regionale Vorlieben (Liste von Regionen)
    favorite_regions: List[str] = Field(default_factory=list)
    
    # Budget-Rahmen
    budget_everyday: Optional[str] = None  # "unter_10", "10_20", "20_35", "35_50", "ueber_50"
    budget_restaurant: Optional[str] = None  # "unter_30", "30_50", "50_80", "80_120", "ueber_120"
    
    # Abneigungen (No-Gos)
    no_gos: List[str] = Field(default_factory=list)  # ["barrique", "schwefel", "chardonnay", etc.]
    
    # Kulinarischer Kontext
    dietary_preferences: List[str] = Field(default_factory=list)  # ["vegetarisch", "vegan", "fleisch", "fisch", "asiatisch", "scharf"]
    
    # Abenteuer-Faktor
    adventure_level: Optional[str] = None  # "klassiker", "ausgewogen", "abenteuerlich"
    
    # Metadata
    updated_at: Optional[datetime] = None

class User(BaseModel):
    model_config = ConfigDict(extra="ignore")
    user_id: str
    email: str
    name: str
    picture: Optional[str] = None
    plan: str = "basic"  # "basic" or "pro"
    subscription_id: Optional[str] = None
    subscription_status: Optional[str] = None  # "active", "cancelled", "expired"
    subscription_end_date: Optional[datetime] = None
    usage: Dict[str, Any] = Field(default_factory=lambda: {
        "pairing_requests_today": 0,
        "chat_messages_today": 0,
        "last_usage_date": None
    })
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class UserSession(BaseModel):
    user_id: str
    session_token: str
    expires_at: datetime
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class PaymentTransaction(BaseModel):
    transaction_id: str = Field(default_factory=lambda: f"txn_{uuid.uuid4().hex[:12]}")
    user_id: str
    email: str
    session_id: str
    plan: str  # "pro_monthly" or "pro_yearly"
    amount: float
    currency: str
    payment_status: str = "pending"  # "pending", "paid", "failed", "expired"
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class CheckoutRequest(BaseModel):
    plan: str  # "pro_monthly" or "pro_yearly"
    origin_url: str

# Auth Request Models
class RegisterRequest(BaseModel):
    email: str
    password: str
    name: str

class LoginRequest(BaseModel):
    email: str
    password: str

class Wine(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    user_id: str  # WICHTIG: Verknüpfung zum Benutzer - jeder User hat seinen eigenen Weinkeller
    name: str
    type: str  # rot, weiss, rose, schaumwein
    region: Optional[str] = None
    year: Optional[int] = None
    grape: Optional[str] = None
    description: Optional[str] = None  # Wine description from database
    notes: Optional[str] = None  # Personal user notes
    image_base64: Optional[str] = None
    quantity: int = 1
    is_favorite: bool = False
    price_category: Optional[str] = None  # 🍷 (bis €20), 🍷🍷 (€20-50), 🍷🍷🍷 (ab €50)
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class WineCreate(BaseModel):
    name: str
    type: str
    region: Optional[str] = None
    year: Optional[int] = None
    grape: Optional[str] = None
    description: Optional[str] = None  # Wine description from database
    notes: Optional[str] = None  # Personal user notes
    image_base64: Optional[str] = None
    quantity: Optional[int] = 1
    price_category: Optional[str] = None  # 🍷 (bis €20), 🍷🍷 (€20-50), 🍷🍷🍷 (ab €50)
    # user_id wird vom Backend gesetzt, nicht vom Client

class WineUpdate(BaseModel):
    name: Optional[str] = None
    type: Optional[str] = None
    region: Optional[str] = None
    year: Optional[int] = None
    grape: Optional[str] = None
    description: Optional[str] = None
    notes: Optional[str] = None
    image_base64: Optional[str] = None
    is_favorite: Optional[bool] = None
    quantity: Optional[int] = None
    price_category: Optional[str] = None  # 🍷 (bis €20), 🍷🍷 (€20-50), 🍷🍷🍷 (ab €50)

# ===================== WINE DATABASE MODELS =====================
from pydantic import field_validator
from typing import Union

class WineDatabaseEntry(BaseModel):
    model_config = ConfigDict(extra="ignore")
    
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    winery: Optional[str] = ""
    country: Optional[str] = ""
    region: Optional[str] = ""
    appellation: Optional[str] = None
    
    # Standardized field names (migration completed 18.12.2025)
    grape_variety: Optional[str] = None
    wine_color: Optional[str] = None
    
    year: Optional[int] = None
    vintage: Optional[int] = None  # Alternative field name
    
    # Validator for backwards compatibility with old field names
    @model_validator(mode='before')
    @classmethod
    def migrate_old_fields(cls, data):
        if isinstance(data, dict):
            # Migrate grape -> grape_variety
            if 'grape' in data and data.get('grape') and not data.get('grape_variety'):
                data['grape_variety'] = data.pop('grape')
            elif 'grape' in data:
                data.pop('grape', None)
            # Migrate color -> wine_color
            if 'color' in data and data.get('color') and not data.get('wine_color'):
                data['wine_color'] = data.pop('color')
            elif 'color' in data:
                data.pop('color', None)
        return data

    # Multilingual descriptions
    description_de: Optional[str] = ""
    description_en: Optional[str] = None
    description_fr: Optional[str] = None

    # Optional structured tasting notes
    tasting_notes: Optional[str] = None

    # Multilingual food pairings (short text or list joined by \n)
    food_pairings_de: Optional[List[str]] = []
    food_pairings_en: Optional[List[str]] = []
    food_pairings_fr: Optional[List[str]] = []

    alcohol_content: Optional[float] = None
    alcohol: Optional[float] = None  # Alternative
    price_category: Optional[str] = None
    price: Optional[float] = None  # Alternative
    image_url: Optional[str] = None
    rating: Optional[float] = None
    classification: Optional[str] = None
    style: Optional[str] = None
    taste: Optional[str] = None
    source: Optional[str] = None
    auto_added: Optional[bool] = None
    imported_at: Optional[str] = None
    created_at: Optional[datetime] = Field(default_factory=lambda: datetime.now(timezone.utc))

class PairingRequest(BaseModel):
    dish: str
    use_cellar: bool = False
    wine_type_filter: Optional[str] = None
    language: str = "de"  # de, en, fr
    dish_id: Optional[str] = None  # optional structured dish from DB
    
    # Restaurant-Modus: User gibt verfügbare Weine von der Karte ein
    available_wines: Optional[str] = None  # Weine von der Weinkarte

    # 4D Profi-Modus Werte (0-10 Skala)
    richness: Optional[int] = None
    freshness: Optional[int] = None
    sweetness: Optional[int] = None
    spice: Optional[int] = None

class PairingResponse(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    dish: str
    recommendation: str
    why_explanation: Optional[str] = None
    cellar_matches: Optional[List[dict]] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

# ===================== WEEKLY TIP MODEL =====================

class WeeklyTip(BaseModel):
    """Wöchentlicher Pairing-Tipp von der KI generiert"""
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    dish: str  # z.B. "Scharfes Thai-Curry"
    dish_emoji: str = "🍽️"  # Emoji für das Gericht
    wine: str  # z.B. "Gewürztraminer Spätlese"
    wine_type: str = "weiss"  # rot, weiss, rose, schaumwein
    region: Optional[str] = None  # z.B. "Elsass, Frankreich"
    why: str  # Kurze Begründung
    fun_fact: Optional[str] = None  # Interessanter Fakt
    week_number: int  # Kalenderwoche
    year: int  # Jahr
    language: str = "de"
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    is_active: bool = True

class ChatMessage(BaseModel):
    role: str  # user or assistant
    content: str
    timestamp: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ChatRequest(BaseModel):
    message: str
    session_id: Optional[str] = None
    image_base64: Optional[str] = None
    language: str = "de"  # de, en, fr

# ===================== WINE DATABASE ADMIN SEED =====================

class WineDbImportStatus(BaseModel):
    imported: int
    failed: int


async def _clear_wine_database():
    """Delete all entries from the public wine database (NOT the personal cellar)."""
    await db.wine_database.delete_many({})


async def _upsert_wine_entry(payload: dict) -> Optional[WineDatabaseEntry]:
    """Insert a single WineDatabaseEntry into the wine_database collection.

    Expects payload to already contain multilingual description/food_pairings.
    """
    try:
        wine = WineDatabaseEntry(**payload)
        doc = wine.model_dump()
        doc["created_at"] = doc["created_at"].isoformat()
        await db.wine_database.insert_one(doc)
        return wine
    except Exception as e:
        logger.warning(f"Failed to upsert wine entry {payload.get('name')}: {e}")
        return None


class ChatResponse(BaseModel):
    response: str
    session_id: str

class LabelScanRequest(BaseModel):
    image_base64: str

class LabelScanResponse(BaseModel):
    name: str
    type: str
    region: Optional[str] = None
    year: Optional[int] = None
    grape: Optional[str] = None
    notes: Optional[str] = None

# ===================== BLOG MODELS =====================

class BlogPost(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    slug: str
    title: str
    title_en: Optional[str] = None
    title_fr: Optional[str] = None
    excerpt: str
    excerpt_en: Optional[str] = None
    excerpt_fr: Optional[str] = None
    content: str
    content_en: Optional[str] = None
    content_fr: Optional[str] = None
    image_url: Optional[str] = None
    category: str  # tipps, wissen, pairings, regionen
    tags: List[str] = []
    author: str = "Sommelier Team"
    published: bool = True
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class BlogPostCreate(BaseModel):
    slug: str
    title: str
    title_en: Optional[str] = None
    title_fr: Optional[str] = None
    excerpt: str
    excerpt_en: Optional[str] = None
    excerpt_fr: Optional[str] = None
    content: str
    content_en: Optional[str] = None
    content_fr: Optional[str] = None
    image_url: Optional[str] = None
    category: str
    tags: List[str] = []
    author: str = "Sommelier Team"

# ===================== FEED MODELS =====================

class FeedComment(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    author_name: str
    content: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class FeedPost(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    author_name: str
    author_id: str  # Simple device/session based ID
    dish: str
    wine_name: str
    wine_type: str  # rot, weiss, rose, schaumwein
    rating: int = Field(ge=1, le=5)  # 1-5 stars
    experience: str  # User's description of the pairing experience
    # Multilingual fields
    dish_en: Optional[str] = None
    dish_fr: Optional[str] = None
    experience_en: Optional[str] = None
    experience_fr: Optional[str] = None
    description_en: Optional[str] = None
    description_fr: Optional[str] = None
    location: Optional[str] = None
    occasion: Optional[str] = None
    image_base64: Optional[str] = None
    likes: List[str] = []  # List of user IDs who liked
    comments: List[FeedComment] = []
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class FeedPostCreate(BaseModel):
    author_name: str
    author_id: str
    dish: str
    wine_name: str
    wine_type: str
    rating: int = Field(ge=1, le=5)
    experience: str
    image_base64: Optional[str] = None

class FeedCommentCreate(BaseModel):
    author_name: str
    author_id: str
    content: str

# ===================== REGIONAL PAIRING MODELS =====================

class RegionalPairing(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    country: str
    country_en: str
    country_fr: str
    country_emoji: str
    country_intro: Optional[str] = None
    country_intro_en: Optional[str] = None
    country_intro_fr: Optional[str] = None
    country_image_url: Optional[str] = None
    region: str
    dish: str
    dish_description: Optional[str] = None
    dish_description_en: Optional[str] = None
    dish_description_fr: Optional[str] = None
    # International Wine Recommendation (Safe Choice)
    wine_name: str
    wine_type: str
    wine_description: Optional[str] = None
    wine_description_en: Optional[str] = None
    wine_description_fr: Optional[str] = None
    # Local Wine Alternative (Discovery)
    local_wine_name: Optional[str] = None
    local_wine_type: Optional[str] = None
    local_wine_description: Optional[str] = None
    local_wine_description_en: Optional[str] = None
    local_wine_description_fr: Optional[str] = None

# ===================== GRAPE VARIETY MODELS =====================

class GrapeVariety(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    slug: Optional[str] = None  # Made optional for backward compatibility
    name: str
    type: Optional[str] = None  # rot, weiss - made optional
    
    # Multilingual names (from import)
    name_de: Optional[str] = None
    name_en: Optional[str] = None
    name_fr: Optional[str] = None
    
    # Poetic descriptions (multilingual) - made optional
    description: Optional[str] = None
    description_de: Optional[str] = None
    description_en: Optional[str] = None
    description_fr: Optional[str] = None
    
    # Characteristics - all made optional for backward compatibility
    color: Optional[str] = None  # From imported data
    synonyms: Union[List[str], str] = []  # Accept both string and list
    body: Optional[str] = None  # leicht, mittel, vollmundig
    body_type: Optional[str] = None  # From imported data
    acidity: Optional[str] = None  # niedrig, mittel, hoch
    acidity_level: Optional[str] = None  # From imported data
    tannin: Optional[str] = None  # niedrig, mittel, hoch
    tannin_level: Optional[str] = None  # From imported data
    aging: Optional[str] = None  # Holz, Edelstahl, etc.
    aging_style: Optional[str] = None  # From imported data
    
    # Aromas - accept both string and list
    primary_aromas: Union[List[str], str] = []
    tertiary_aromas: Union[List[str], str] = []
    
    # Food pairings
    perfect_pairings: List[str] = []
    perfect_pairings_en: List[str] = []
    perfect_pairings_fr: List[str] = []
    
    # Regions
    main_regions: List[str] = []
    region: Optional[str] = None  # From imported data
    
    # Image
    image_url: Optional[str] = None
    
    # Metadata
    category: Optional[str] = None
    source: Optional[str] = None
    imported_at: Optional[datetime] = None
    
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    
    @model_validator(mode='before')
    @classmethod
    def map_legacy_fields(cls, data):
        """Map legacy field names to new ones and handle string/list conversions"""
        if isinstance(data, dict):
            # Map slug from name if missing
            if not data.get('slug') and data.get('name'):
                data['slug'] = data['name'].lower().replace(' ', '-').replace('ü', 'ue').replace('ö', 'oe').replace('ä', 'ae').replace('ß', 'ss')
            
            # Map type from color
            if not data.get('type') and data.get('color'):
                color = data['color'].lower()
                if 'rot' in color or 'red' in color:
                    data['type'] = 'rot'
                elif 'weiß' in color or 'weiss' in color or 'white' in color:
                    data['type'] = 'weiss'
                else:
                    data['type'] = color
            
            # Map description from description_de
            if not data.get('description') and data.get('description_de'):
                data['description'] = data['description_de']
            
            # Map body from body_type
            if not data.get('body') and data.get('body_type'):
                data['body'] = data['body_type']
            
            # Map acidity from acidity_level
            if not data.get('acidity') and data.get('acidity_level'):
                data['acidity'] = data['acidity_level']
            
            # Map tannin from tannin_level
            if not data.get('tannin') and data.get('tannin_level'):
                data['tannin'] = data['tannin_level']
            
            # Map aging from aging_style
            if not data.get('aging') and data.get('aging_style'):
                data['aging'] = data['aging_style']
            
            # Convert string aromas to list
            for field in ['primary_aromas', 'tertiary_aromas', 'synonyms']:
                if isinstance(data.get(field), str):
                    val = data[field]
                    if val:
                        data[field] = [x.strip() for x in val.split(',') if x.strip()]
                    else:
                        data[field] = []
        
        return data


# ===================== DISH MODELS =====================

class Dish(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    slug: str

    # Multilingual names
    name_de: str
    name_en: Optional[str] = None
    name_fr: Optional[str] = None

    # Classification
    country: Optional[str] = None
    region: Optional[str] = None
    trend_cuisines: List[str] = []
    bestseller_category: Optional[str] = None  # burger, pasta, steak, fisch, etc.

    # Technical matrix for pairing
    protein: Optional[str] = None
    intensity: Optional[str] = None  # leicht, mittel, kräftig
    cooking_method: Optional[str] = None
    sauce_base: Optional[str] = None
    fat_level: Optional[str] = None  # niedrig, mittel, hoch
    acid_level: Optional[str] = None  # niedrig, mittel, hoch
    sweetness_level: Optional[str] = None  # trocken, leicht_süß, süß
    spice_level: Optional[str] = None  # keine, leicht, mittel, dominant
    key_aromas: List[str] = []
    texture: List[str] = []

    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))


class DishGenerationRequest(BaseModel):
    """Request model to generate a new structured dish entry via LLM"""
    base_name: str
    country_hint: Optional[str] = None
    trend_hint: Optional[str] = None
    bestseller_category: Optional[str] = None


class GrapeGenerationRequest(BaseModel):
    """Request model to generate a new grape variety via LLM"""
    name: str
    grape_type: Optional[str] = None  # "rot" oder "weiss" - wenn None, vom Modell ableiten
    style_hint: Optional[str] = None  # z.B. "klassisch, trocken, hochwertige Qualitätsweine"



# ===================== SOMMELIER SYSTEM MESSAGE =====================

SOMMELIER_SYSTEM_DE = """Du bist der "Wine-Pairing.Online Sommelier" - ein Master of Wine mit Leidenschaft für Qualitätsweine. Dein Ziel ist es, wissenschaftlich fundierte Empfehlungen zu geben, die im guten Fachhandel verfügbar sind.

WICHTIG: Halte dich EXAKT an diese Struktur!

ANALYSE-SCHRITTE:
1. Style-First: Identifiziere das benötigte Weinprofil (z.B. "Hohe Säure, wenig Tannin, mineralisch")
2. Budget-Check: Gib Empfehlungen in allen Preisstufen für Weinliebhaber

STRUKTUR DEINER ANTWORT:

**🍷 DER STIL**
[Erkläre kurz den passenden Weinstil in 1-2 Sätzen]

**💡 DAS WARUM**
[Erkläre die wissenschaftliche Balance zum Gericht - z.B. "Die Säure schneidet durch das Fett"]

**🍷 EMPFEHLUNGEN**

🍷 **Alltags-Genuss (bis €20):**
- **[Weingut/Weinname, Region]** – [Warum er passt]
- **[Weingut/Weinname, Region]** – [Warum er passt]

🍷🍷 **Gehobener Anlass (€20-50):**
- **[Weingut/Weinname, Region]** – [Warum er passt]

🍷🍷🍷 **Besonderer Moment (ab €50):**
- **[Weingut/Weinname, Region]** – [Warum er passt] *(Optional)*

**💎 GEHEIMTIPP**
[Nenne eine Alternative aus einer weniger bekannten Region, die das gleiche Profil bietet - z.B. Languedoc statt Burgund, Pfalz statt Mosel]

REGELN:
- IMMER mit "🍷 Alltags-Genuss" beginnen!
- Empfehle Weine aus dem guten Fachhandel, nicht nur Supermarkt-Weine
- Nenne konkrete Weingüter (z.B. Dönnhoff, Keller, Trimbach, Antinori, Gaja)
- Bei Fleisch: Rotwein-Fokus | Bei Fisch: Weißwein-Fokus
- Weinnamen IMMER in **fett**
- Antworte prägnant auf Deutsch"""

SOMMELIER_SYSTEM_EN = """You are the "Wine-Pairing.Online Sommelier" - a Master of Wine with a passion for quality wines. Your goal is to provide scientifically sound recommendations available at good wine shops.

IMPORTANT: Follow this structure EXACTLY!

ANALYSIS STEPS:
1. Style-First: Identify the required wine profile (e.g., "High acidity, low tannin, mineral")
2. Budget-Check: Give recommendations across all price tiers for wine enthusiasts

STRUCTURE YOUR RESPONSE:

**🍷 THE STYLE**
[Briefly explain the matching wine style in 1-2 sentences]

**💡 THE WHY**
[Explain the scientific balance with the dish - e.g., "The acidity cuts through the fat"]

**🍷 RECOMMENDATIONS**

🍷 **Everyday Enjoyment (up to €20):**
- **[Winery/Wine Name, Region]** – [Why it fits]
- **[Winery/Wine Name, Region]** – [Why it fits]

🍷🍷 **Special Occasion (€20-50):**
- **[Winery/Wine Name, Region]** – [Why it fits]

🍷🍷🍷 **Exceptional Moment (€50+):**
- **[Winery/Wine Name, Region]** – [Why it fits] *(Optional)*

**💎 INSIDER TIP**
[Name an alternative from a lesser-known region that offers the same profile - e.g., Languedoc instead of Burgundy]

RULES:
- ALWAYS start with "🍷 Everyday Enjoyment"!
- Recommend wines from good wine shops, not just supermarket wines
- Name specific wineries (e.g., Dönnhoff, Keller, Trimbach, Antinori, Gaja)
- For meat: Red wine focus | For fish: White wine focus
- Wine names ALWAYS in **bold**
- Keep responses concise in English"""

SOMMELIER_SYSTEM_FR = """Vous êtes le "Sommelier Wine-Pairing.Online" - un Master of Wine passionné par les vins de qualité. Votre objectif est de fournir des recommandations scientifiquement fondées, disponibles dans les bonnes caves à vin.

IMPORTANT: Suivez cette structure EXACTEMENT!

ÉTAPES D'ANALYSE:
1. Style-First: Identifiez le profil de vin requis (ex: "Acidité élevée, peu de tanins, minéral")
2. Budget-Check: Donnez des recommandations dans toutes les gammes de prix pour les amateurs de vin

STRUCTURE DE VOTRE RÉPONSE:

**🍷 LE STYLE**
[Expliquez brièvement le style de vin adapté en 1-2 phrases]

**💡 LE POURQUOI**
[Expliquez l'équilibre scientifique avec le plat - ex: "L'acidité coupe le gras"]

**🍷 RECOMMANDATIONS**

🍷 **Plaisir Quotidien (jusqu'à €20):**
- **[Domaine/Nom du Vin, Région]** – [Pourquoi il convient]
- **[Domaine/Nom du Vin, Région]** – [Pourquoi il convient]

🍷🍷 **Belle Occasion (€20-50):**
- **[Domaine/Nom du Vin, Région]** – [Pourquoi il convient]

🍷🍷🍷 **Moment d'Exception (à partir de €50):**
- **[Domaine/Nom du Vin, Région]** – [Pourquoi il convient] *(Optionnel)*

**💎 BON PLAN**
[Nommez une alternative d'une région moins connue offrant le même profil - ex: Languedoc au lieu de Bourgogne]

RÈGLES:
- Commencez TOUJOURS par "🍷 Plaisir Quotidien"!
- Recommandez des vins de bonnes caves, pas seulement des vins de supermarché
- Nommez des domaines concrets (ex: Dönnhoff, Keller, Trimbach, Antinori, Gaja)
- Pour la viande: Focus vin rouge | Pour le poisson: Focus vin blanc
- Noms de vin TOUJOURS en **gras**
- Réponses concises en français"""

# ===================== RESTAURANT MODE SYSTEM MESSAGES =====================

RESTAURANT_MODE_DE = """Du bist der "Wine-Pairing.Online Sommelier" - ein Master of Wine. Der Kunde sitzt im Restaurant und braucht eine konkrete Empfehlung aus der Weinkarte.

WICHTIG: Halte dich EXAKT an diese Struktur!

STRUKTUR DEINER ANTWORT:

**🍷 MEINE EMPFEHLUNG**
[Nenne DEN BESTEN Wein aus der Liste - NUR EINEN!]

**💡 WARUM GENAU DIESER WEIN?**
[Erkläre in 3-4 Sätzen warum dieser Wein perfekt zum Gericht passt. Gehe auf Aromen, Textur und Harmonie ein.]

**🔄 ALTERNATIVE AUS DER LISTE**
[Falls vorhanden: Nenne eine zweite Option aus der Liste und warum sie auch funktioniert]

**⚠️ VERMEIDE**
[Falls auf der Liste: Welchen Wein sollte man zu diesem Gericht NICHT wählen und warum?]

REGELN:
- Empfehle NUR Weine die der Kunde genannt hat - KEINE anderen!
- Sei konkret und entscheidungsfreudig - der Kunde will EINE klare Empfehlung
- Begründe kurz aber überzeugend
- Wenn du einen Wein nicht kennst, gehe nach Rebsorte/Region"""

RESTAURANT_MODE_EN = """You are the "Wine-Pairing.Online Sommelier" - a Master of Wine. The customer is at a restaurant and needs a specific recommendation from the wine list.

IMPORTANT: Follow this structure EXACTLY!

STRUCTURE YOUR RESPONSE:

**🍷 MY RECOMMENDATION**
[Name THE BEST wine from the list - ONLY ONE!]

**💡 WHY THIS WINE?**
[Explain in 3-4 sentences why this wine pairs perfectly with the dish. Discuss aromas, texture, and harmony.]

**🔄 ALTERNATIVE FROM THE LIST**
[If available: Name a second option from the list and why it would also work]

**⚠️ AVOID**
[If on the list: Which wine should NOT be chosen for this dish and why?]

RULES:
- Only recommend wines the customer has listed - NO others!
- Be concrete and decisive - the customer wants ONE clear recommendation
- Justify briefly but convincingly
- If you don't know a wine, go by grape/region"""

RESTAURANT_MODE_FR = """Vous êtes le "Sommelier Wine-Pairing.Online" - un Master of Wine. Le client est au restaurant et a besoin d'une recommandation concrète de la carte des vins.

IMPORTANT: Suivez cette structure EXACTEMENT!

STRUCTURE DE VOTRE RÉPONSE:

**🍷 MA RECOMMANDATION**
[Nommez LE MEILLEUR vin de la liste - UN SEUL!]

**💡 POURQUOI CE VIN?**
[Expliquez en 3-4 phrases pourquoi ce vin s'accorde parfaitement avec le plat. Discutez des arômes, de la texture et de l'harmonie.]

**🔄 ALTERNATIVE DE LA LISTE**
[Si disponible: Nommez une deuxième option de la liste et pourquoi elle fonctionnerait aussi]

**⚠️ À ÉVITER**
[Si sur la liste: Quel vin NE devrait PAS être choisi pour ce plat et pourquoi?]

RÈGLES:
- Recommandez UNIQUEMENT les vins que le client a listés - PAS d'autres!
- Soyez concret et décisif - le client veut UNE recommandation claire
- Justifiez brièvement mais de manière convaincante
- Si vous ne connaissez pas un vin, basez-vous sur le cépage/la région"""

def get_restaurant_mode_system(language: str) -> str:
    """Get the restaurant mode system message for the specified language"""
    if language == "en":
        return RESTAURANT_MODE_EN
    elif language == "fr":
        return RESTAURANT_MODE_FR
    return RESTAURANT_MODE_DE

# System prompt for structured grape variety generation
GRAPE_GENERATOR_SYSTEM = """Du bist Claude, Master of Wine und leidenschaftlicher Koch.
Deine Aufgabe: Für eine gegebene Rebsorte einen vollständigen Datensatz für eine Wein-App zu erzeugen.

ANTWORTFORMAT (STRICT JSON, KEIN ERKLÄRTEXT):
{
  "slug": "kebab-case-slug-ohne-uml...",
  "name": "Name der Rebsorte",
  "type": "rot" oder "weiss",
  "description": "Poetische deutsche Beschreibung (3-5 Sätze)",
  "description_en": "Poetic English description (3-5 sentences)",
  "description_fr": "Description poétique en français (3-5 phrases)",
  "synonyms": ["Synonym 1", "Synonym 2"],
  "body": "leicht" oder "mittel" oder "vollmundig",
  "acidity": "niedrig" oder "mittel" oder "hoch",
  "tannin": "niedrig" oder "mittel" oder "hoch",
  "aging": "Kurze Beschreibung des typischen Ausbaus (z.B. Edelstahl, Holzfass, Barrique)",
  "primary_aromas": ["3-6 kurze deutsche aroma-tags in kleinschreibung"],
  "tertiary_aromas": ["3-6 kurze deutsche aroma-tags in kleinschreibung"],
  "perfect_pairings": ["3-6 kurze deutsche speisen-tags in kleinschreibung"],
  "perfect_pairings_en": ["3-6 short english food pairing tags"],
  "perfect_pairings_fr": ["3-6 étiquettes d'accords mets-vins en français"],
  "main_regions": ["3-6 wichtigste anbaugebiete"]
}
"""


# ===================== WINE DATABASE ENDPOINTS =====================

@api_router.get("/wine-database", response_model=List[WineDatabaseEntry])
async def list_wine_database(
    search: Optional[str] = None,
    country: Optional[str] = None,
    region: Optional[str] = None,
    appellation: Optional[str] = None,
    grape_variety: Optional[str] = None,
    wine_color: Optional[str] = None,
    price_category: Optional[str] = None,
    skip: int = 0,
    limit: int = 50,
):
    """List wines from the public wine database with basic filters.

    This endpoint returns the raw multilingual wine entries. The frontend is
    responsible for selecting the appropriate language fields.
    """
    logger.info(f"ENDPOINT CALLED: /wine-database with limit={limit}")
    query: dict = {}

    if search:
        # WICHTIG: Akzent-insensitive Suche verwenden!
        # "Chateau" findet "Château", "Cotes" findet "Côtes"
        accent_pattern = create_accent_insensitive_pattern(search)
        regex = {"$regex": accent_pattern, "$options": "i"}
        query["$or"] = [
            {"name": regex},
            {"winery": regex},
            {"region": regex},
            {"appellation": regex},
            {"grape_variety": regex},
        ]
    if country:
        query["country"] = country
    if region:
        query["region"] = region
    if appellation:
        query["appellation"] = appellation
    if grape_variety:
        query["grape_variety"] = grape_variety
    if wine_color:
        query["wine_color"] = wine_color
    if price_category:
        query["price_category"] = price_category

    wines = (
        await db.wine_database
        .find(query, {"_id": 0})
        .skip(skip)
        .limit(limit)
        .to_list(limit)
    )
    
    logger.info(f"Wine database query returned {len(wines)} wines")
    
    # Convert datetime strings to datetime objects for Pydantic validation
    for wine in wines:
        if isinstance(wine.get('created_at'), str):
            wine['created_at'] = datetime.fromisoformat(wine['created_at'])
    
    return wines


# System prompt for structured dish generation
DISH_GENERATOR_SYSTEM = """Du bist Claude, Master of Wine und leidenschaftlicher Koch.
Deine Aufgabe: Für ein Gericht einen vollständigen, strukturierten Datensatz zu erzeugen, der für Wein-Pairing verwendet werden kann.

# Removed duplicate content
{
  "slug": "kebab-case-slug-ohne-uml...",
  "name_de": "Name des Gerichts auf Deutsch",
  "name_en": "Name des Gerichts auf Englisch",
  "name_fr": "Nom du plat en français",
  "country": "land in kleinschreibung, z.b. italien, thailand, usa",
  "region": "region in kleinschreibung, z.b. toskana, isaan, bayern",
  "trend_cuisines": ["1-3 trend-tags in kleinschreibung, z.b. thai, asiatisch, streetfood"],
  "bestseller_category": "burger | pasta | steak | fisch | pizza | bowl | curry | salat | sushi | nudelsuppe | dessert",
  "protein": "hauptprotein: rind, lamm, schwein, geflügel, lachs, weißer_fisch, meeresfrüchte, gemüse, vegan",
  "intensity": "leicht | mittel | kräftig",
  "cooking_method": "gebraten | gegrillt | geschmort | frittiert | roh | gebacken | wok",
  "sauce_base": "jus | sahne | butter | tomate | vinaigrette | kokosmilch | soja | brühe | keine",
  "fat_level": "niedrig | mittel | hoch",
  "acid_level": "niedrig | mittel | hoch",
  "sweetness_level": "trocken | leicht_süß | süß",
  "spice_level": "keine | leicht | mittel | dominant",
  "key_aromas": ["3-8 aroma-tags in kleinschreibung, z.b. röstnoten, kräuter, zitronig, umami"],
  "texture": ["2-5 textur-tags in kleinschreibung, z.b. zart, cremig, knusprig, saftig"]
}

WICHTIG:
- Verwende GENAU diese Feldnamen.
- Verwende bei allen skalenfeldern NUR die angegebenen werte.
- Gib KEINEN zusätzlichen text außer dem JSON zurück.
- Verwende in allen tag-listen (trend_cuisines, key_aromas, texture) nur kleinschreibung.
"""



def get_sommelier_system(language: str = "de") -> str:
    """Get the appropriate system message based on language"""
    if language == "en":
        return SOMMELIER_SYSTEM_EN
    elif language == "fr":
        return SOMMELIER_SYSTEM_FR
    return SOMMELIER_SYSTEM_DE

# ===================== HEALTH CHECK ENDPOINT =====================

@api_router.get("/health")
async def health_check():
    """
    Health check endpoint for Kubernetes/deployment monitoring.
    Returns the current status of the application and database connectivity.
    """
    try:
        # Check database connectivity
        await db.command("ping")
        db_status = "connected"
    except Exception as e:
        db_status = f"error: {str(e)}"
    
    return {
        "status": "healthy",
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "database": db_status,
        "version": "3.1"
    }


# ===================== WINE CELLAR ENDPOINTS =====================

@api_router.get("/")
async def root():
    return {"message": "Wine Pairing API - Ihr virtueller Sommelier"}

@api_router.get("/wines", response_model=List[Wine])
async def get_wines(
    request: Request,
    type_filter: Optional[str] = None, 
    price_category_filter: Optional[str] = None,
    favorites_only: bool = False, 
    in_stock_only: bool = False
):
    """Get wines from the user's personal cellar (requires authentication)"""
    # User muss eingeloggt sein für Weinkeller-Zugriff
    user = await get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Bitte melden Sie sich an, um Ihren Weinkeller zu sehen")
    
    # Query NUR für Weine des aktuellen Users
    query = {"user_id": user.user_id}
    if type_filter:
        query["type"] = type_filter
    if price_category_filter:
        query["price_category"] = price_category_filter
    if favorites_only:
        query["is_favorite"] = True
    if in_stock_only:
        query["quantity"] = {"$gt": 0}
    
    # Exclude large image_base64 field for better performance
    wines = await db.wines.find(query, {"_id": 0, "image_base64": 0}).to_list(1000)
    for wine in wines:
        if isinstance(wine.get('created_at'), str):
            wine['created_at'] = datetime.fromisoformat(wine['created_at'])
    return wines

@api_router.get("/wines/{wine_id}", response_model=Wine)
async def get_wine(wine_id: str, request: Request):
    """Get a specific wine by ID (must belong to current user)"""
    user = await get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Bitte melden Sie sich an")
    
    # Wein muss dem User gehören
    wine = await db.wines.find_one({"id": wine_id, "user_id": user.user_id}, {"_id": 0})
    if not wine:
        raise HTTPException(status_code=404, detail="Wein nicht gefunden")
    if isinstance(wine.get('created_at'), str):
        wine['created_at'] = datetime.fromisoformat(wine['created_at'])
    return wine

@api_router.post("/wines", response_model=Wine)
async def create_wine(wine_data: WineCreate, request: Request):
    """Add a new wine to the user's personal cellar"""
    user = await get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Bitte melden Sie sich an, um Weine zu speichern")
    
    # Check cellar limit for basic users
    allowed, message = await check_limit(user, "cellar")
    if not allowed:
        raise HTTPException(status_code=403, detail=message)
    
    # Erstelle Wine mit user_id
    wine_dict = wine_data.model_dump()
    wine_dict["user_id"] = user.user_id  # WICHTIG: Verknüpfung zum User
    wine = Wine(**wine_dict)
    
    doc = wine.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    await db.wines.insert_one(doc)
    
    logger.info(f"🍷 Wine '{wine.name}' added to cellar of user {user.user_id}")
    return wine

@api_router.put("/wines/{wine_id}", response_model=Wine)
async def update_wine(wine_id: str, wine_update: WineUpdate, request: Request):
    """Update a wine in the user's cellar (must belong to current user)"""
    user = await get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Bitte melden Sie sich an")
    
    # Wein muss dem User gehören
    existing = await db.wines.find_one({"id": wine_id, "user_id": user.user_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Wein nicht gefunden oder gehört nicht Ihnen")
    
    update_data = {k: v for k, v in wine_update.model_dump().items() if v is not None}
    if update_data:
        await db.wines.update_one({"id": wine_id, "user_id": user.user_id}, {"$set": update_data})
    
    updated = await db.wines.find_one({"id": wine_id, "user_id": user.user_id}, {"_id": 0})
    if isinstance(updated.get('created_at'), str):
        updated['created_at'] = datetime.fromisoformat(updated['created_at'])
    return updated

@api_router.delete("/wines/{wine_id}")
async def delete_wine(wine_id: str, request: Request):
    """Remove a wine from the user's cellar (must belong to current user)"""
    user = await get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Bitte melden Sie sich an")
    
    # Nur Weine des Users können gelöscht werden
    result = await db.wines.delete_one({"id": wine_id, "user_id": user.user_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Wein nicht gefunden oder gehört nicht Ihnen")
    
    logger.info(f"🗑️ Wine {wine_id} deleted from cellar of user {user.user_id}")
    return {"message": "Wein erfolgreich gelöscht"}

@api_router.post("/wines/{wine_id}/favorite")
async def toggle_favorite(wine_id: str, request: Request):
    """Toggle favorite status of a wine (must belong to current user)"""
    user = await get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Bitte melden Sie sich an")
    
    # Wein muss dem User gehören
    wine = await db.wines.find_one({"id": wine_id, "user_id": user.user_id}, {"_id": 0})
    if not wine:
        raise HTTPException(status_code=404, detail="Wein nicht gefunden oder gehört nicht Ihnen")
    
    new_status = not wine.get('is_favorite', False)
    await db.wines.update_one({"id": wine_id, "user_id": user.user_id}, {"$set": {"is_favorite": new_status}})
    return {"is_favorite": new_status}

# ===================== AI PAIRING ENDPOINTS =====================

@api_router.post("/pairing", response_model=PairingResponse)
async def get_wine_pairing(request: PairingRequest, http_request: Request):
    """Get AI-powered wine pairing recommendation with caching"""
    try:
        # Check user limits
        user = await get_current_user(http_request)
        allowed, message = await check_limit(user, "pairing")
        
        if not allowed:
            raise HTTPException(status_code=429, detail=message)
        
        # Check cache first (only for requests without cellar, dish_id, available_wines, or 4D parameters)
        # These are "simple" requests that can be cached
        is_cacheable = (
            not request.use_cellar and 
            not request.dish_id and
            not request.available_wines and  # Restaurant-Modus nicht cachen
            request.richness is None and
            request.freshness is None and
            request.sweetness is None and
            request.spice is None
        )
        
        cache_key = None
        # WICHTIG: use_cellar und available_wines Anfragen sollten NICHT gecacht werden
        if is_cacheable and not request.use_cellar and not request.available_wines:
            cache_key = get_cache_key(request.dish, request.language, request.wine_type_filter, request.use_cellar)
            cached_result = get_cached_pairing(cache_key)
            if cached_result:
                # Return cached result immediately (nur für simple Anfragen)
                return PairingResponse(
                    dish=request.dish,
                    recommendation=cached_result['recommendation'],
                    why_explanation=cached_result.get('why_explanation'),
                    cellar_matches=None
                )
        
        # No cache hit - make LLM call
        # Check if this is "Restaurant Mode" with available wines
        if request.available_wines and request.available_wines.strip():
            # Use special Restaurant Mode system message
            system_message = get_restaurant_mode_system(request.language)
        else:
            # Get standard language-specific system message
            system_message = get_sommelier_system(request.language)
        
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=str(uuid.uuid4()),
            system_message=system_message
        ).with_model("openai", "gpt-5.1")
        
        # Get cellar wines if requested
        cellar_matches = None
        cellar_context = ""
        
        # Restaurant-Modus: Verfügbare Weine von der Karte
        restaurant_context = ""
        if request.available_wines and request.available_wines.strip():
            if request.language == "en":
                restaurant_context = f"\n\n🍷 WINE LIST SELECTION:\nThe customer is at a restaurant and has these wines available on the menu:\n{request.available_wines}\n\nPlease recommend THE BEST wine from this list for the dish. Explain WHY this specific wine is the best choice."
            elif request.language == "fr":
                restaurant_context = f"\n\n🍷 SÉLECTION DE LA CARTE DES VINS:\nLe client est au restaurant et a ces vins disponibles sur la carte:\n{request.available_wines}\n\nVeuillez recommander LE MEILLEUR vin de cette liste pour le plat. Expliquez POURQUOI ce vin spécifique est le meilleur choix."
            else:
                restaurant_context = f"\n\n🍷 WEINKARTEN-AUSWAHL:\nDer Kunde sitzt im Restaurant und hat folgende Weine auf der Karte zur Auswahl:\n{request.available_wines}\n\nBitte empfehle DEN BESTEN Wein aus dieser Liste zum Gericht. Erkläre WARUM genau dieser Wein die beste Wahl ist."
        
        if request.use_cellar:
            # User muss eingeloggt sein für Keller-Empfehlungen
            if not user:
                raise HTTPException(status_code=401, detail="Bitte melden Sie sich an, um Empfehlungen aus Ihrem Weinkeller zu erhalten")
            
            query = {"user_id": user.user_id}  # NUR Weine des Users
            # WICHTIG: "all" bedeutet ALLE Weintypen, also keine Filterung
            if request.wine_type_filter and request.wine_type_filter != "all":
                query["type"] = request.wine_type_filter
            
            cellar_wines = await db.wines.find(query, {"_id": 0, "image_base64": 0}).to_list(100)
            
            if cellar_wines:
                # Setze cellar_matches für die API-Response
                cellar_matches = [{"id": w["id"], "name": w["name"], "type": w["type"]} for w in cellar_wines[:5]]
                
                # Translate cellar context based on language
                if request.language == "en":
                    cellar_context = "\n\nThe customer has the following wines in the cellar:\n"
                elif request.language == "fr":
                    cellar_context = "\n\nLe client a les vins suivants dans sa cave:\n"
                else:
                    cellar_context = "\n\nDer Kunde hat folgende Weine im Keller:\n"
                
                for w in cellar_wines:
                    cellar_context += f"- {w['name']} ({w['type']})"
                    if w.get('region'):
                        if request.language == "en":
                            cellar_context += f" from {w['region']}"
                        elif request.language == "fr":
                            cellar_context += f" de {w['region']}"
                        else:
                            cellar_context += f" aus {w['region']}"
                    if w.get('year'):
                        cellar_context += f", {w['year']}"
                    if w.get('grape'):
                        cellar_context += f", {w['grape']}"
                    cellar_context += "\n"
                
                if request.language == "en":
                    cellar_context += "\nPlease recommend suitable wines from the customer's cellar first, then general recommendations."
                elif request.language == "fr":
                    cellar_context += "\nVeuillez d'abord recommander des vins appropriés de la cave du client, puis des recommandations générales."
                else:
                    cellar_context += "\nBitte empfehle zuerst passende Weine aus dem Keller des Kunden, dann allgemeine Empfehlungen."
        
        # Optional: include structured dish information if provided
        dish_context = ""
        if request.dish_id:
            dish = await db.dishes.find_one({"id": request.dish_id}, {"_id": 0})
            if dish:
                # Build a compact, language-agnostic technical summary for Claude
                dish_context = "\n\nTECHNISCHE GERICHTSANALYSE (für internes Pairing, nicht direkt auf der Karte ausgeben):\n"
                dish_context += f"- Land: {dish.get('country') or '-'}, Region: {dish.get('region') or '-'}\n"
                dish_context += f"- Trendküche: {', '.join(dish.get('trend_cuisines', [])) or '-'}\n"
                dish_context += f"- Bestseller-Kategorie: {dish.get('bestseller_category') or '-'}\n"
                dish_context += f"- Protein: {dish.get('protein') or '-'}\n"
                dish_context += f"- Intensität: {dish.get('intensity') or '-'}\n"
                dish_context += f"- Garmethode: {dish.get('cooking_method') or '-'}\n"
                dish_context += f"- Saucenbasis: {dish.get('sauce_base') or '-'}\n"
                dish_context += f"- Fettgehalt: {dish.get('fat_level') or '-'}, Säure: {dish.get('acid_level') or '-'}, Süße: {dish.get('sweetness_level') or '-'}\n"
                dish_context += f"- Schärfe: {dish.get('spice_level') or '-'}\n"
                dish_context += f"- Aromen: {', '.join(dish.get('key_aromas', [])) or '-'}\n"
                dish_context += f"- Textur: {', '.join(dish.get('texture', [])) or '-'}\n"

        # Profi-Modus 4D Kontext (Richness, Freshness, Sweetness, Spice)
        four_d_context = ""
        if any([
            request.richness is not None,
            request.freshness is not None,
            request.sweetness is not None,
            request.spice is not None,
        ]):
            four_d_context = "\n\n4D GAUMEN-ANALYSE (bitte als Grundlage für die Erklärung der Harmonie nutzen):\n"
            four_d_context += f"- Reichhaltigkeit (Richness): {request.richness if request.richness is not None else '-'} auf einer Skala von 0-10\n"
            four_d_context += f"- Frische (Freshness): {request.freshness if request.freshness is not None else '-'} auf einer Skala von 0-10\n"
            four_d_context += f"- Süße (Sweetness): {request.sweetness if request.sweetness is not None else '-'} auf einer Skala von 0-10\n"
            four_d_context += f"- Würze (Spice): {request.spice if request.spice is not None else '-'} auf einer Skala von 0-10\n"
            four_d_context += "\nNutze diese vier Dimensionen, um im Anschluss eine kompakte Erklärung zu geben, WARUM deine Empfehlung harmoniert. Erkläre vor allem die BRÜCKE zwischen Gericht und Wein."

        # Pro User: Get personalized wine profile context
        profile_context = ""
        if user and user.plan == "pro":
            try:
                user_profile = await db.wine_profiles.find_one({"user_id": user.user_id}, {"_id": 0})
                if user_profile:
                    profile_context = get_profile_context_for_ai(user_profile, request.language or "de")
            except Exception as profile_error:
                logger.warning(f"Could not load wine profile: {profile_error}")

        # Weinart-Präferenz des Benutzers
        wine_type_preference = ""
        if request.wine_type_filter and request.wine_type_filter != "all":
            wine_type_names = {
                "rot": {"de": "Rotwein", "en": "red wine", "fr": "vin rouge"},
                "weiss": {"de": "Weißwein", "en": "white wine", "fr": "vin blanc"},
                "rose": {"de": "Roséwein", "en": "rosé wine", "fr": "vin rosé"},
                "schaumwein": {"de": "Schaumwein/Champagner", "en": "sparkling wine/champagne", "fr": "vin mousseux/champagne"}
            }
            wine_name = wine_type_names.get(request.wine_type_filter, {}).get(request.language or "de", request.wine_type_filter)
            
            if request.language == "en":
                wine_type_preference = f"\n\n⚠️ IMPORTANT: The customer specifically wants {wine_name}. Please ONLY recommend {wine_name}s, even if other wine types might also pair well. Respect the customer's preference!"
            elif request.language == "fr":
                wine_type_preference = f"\n\n⚠️ IMPORTANT: Le client souhaite spécifiquement du {wine_name}. Veuillez recommander UNIQUEMENT des {wine_name}s, même si d'autres types de vins pourraient aussi bien s'accorder. Respectez la préférence du client!"
            else:
                wine_type_preference = f"\n\n⚠️ WICHTIG: Der Kunde wünscht ausdrücklich {wine_name}. Bitte empfehle AUSSCHLIESSLICH {wine_name}e, auch wenn andere Weinarten ebenfalls passen könnten. Respektiere den Kundenwunsch!"

        # Restaurant-Modus: Keine Standard-Empfehlungen, nur aus der Liste
        if request.available_wines and request.available_wines.strip():
            # Restaurant Mode - simple prompt with wine list
            if request.language == "en":
                prompt = f"I would like to eat {request.dish}.{wine_type_preference}{restaurant_context}{profile_context}"
            elif request.language == "fr":
                prompt = f"Je voudrais manger {request.dish}.{wine_type_preference}{restaurant_context}{profile_context}"
            else:
                prompt = f"Ich möchte {request.dish} essen.{wine_type_preference}{restaurant_context}{profile_context}"
        else:
            # Standard mode with full recommendations
            # Translate main prompt based on language
            if request.language == "en":
                base_prompt = f"I would like to eat {request.dish}. Which wine do you recommend?{wine_type_preference}{cellar_context}{dish_context}{four_d_context}{profile_context}"
                explanation_instruction = "\n\nAfter your recommendation, add a short section titled 'Why this pairing works' that explains in 3-5 sentences WHY your recommendation harmonises with the dish based on the four dimensions (richness, freshness, sweetness, spice) and the bridge between food and wine. Mark this section clearly with 'WHY_SECTION_START' and 'WHY_SECTION_END'."
            elif request.language == "fr":
                base_prompt = f"Je voudrais manger {request.dish}. Quel vin recommandez-vous?{wine_type_preference}{cellar_context}{dish_context}{four_d_context}{profile_context}"
                explanation_instruction = "\n\nAprès votre recommandation, ajoutez une courte section intitulée 'Pourquoi cet accord fonctionne' qui explique en 3-5 phrases POURQUOI votre recommandation s'harmonise avec le plat sur la base des quatre dimensions (richesse, fraîcheur, douceur, épice) et du pont entre mets et vin. Marquez clairement cette section avec 'WHY_SECTION_START' et 'WHY_SECTION_END'."
            else:
                base_prompt = f"Ich möchte {request.dish} essen. Welchen Wein empfiehlst du dazu?{wine_type_preference}{cellar_context}{dish_context}{four_d_context}{profile_context}"
                explanation_instruction = "\n\nGib nach deiner Empfehlung einen kurzen Abschnitt mit der Überschrift 'Warum dieses Pairing funktioniert' aus. Erkläre in 3-5 Sätzen, WARUM deine Empfehlung mit dem Gericht harmoniert – entlang der vier Dimensionen (Reichhaltigkeit, Frische, Süße, Würze) und der BRÜCKE zwischen Speise und Wein. Kennzeichne diesen Abschnitt klar mit 'WHY_SECTION_START' und 'WHY_SECTION_END'."

            prompt = base_prompt + explanation_instruction
        
        user_message = UserMessage(text=prompt)
        response = await chat.send_message(user_message)

        # Extract WHY explanation section if present
        why_explanation = None
        if isinstance(response, str) and "WHY_SECTION_START" in response and "WHY_SECTION_END" in response:
            try:
                start = response.index("WHY_SECTION_START") + len("WHY_SECTION_START")
                end = response.index("WHY_SECTION_END", start)
                why_explanation = response[start:end].strip()
                # Remove markers and explanation from main recommendation text
                response = (response[:response.index("WHY_SECTION_START")].strip() + "\n\n" + response[end + len("WHY_SECTION_END"):].strip()).strip()
            except Exception:
                # Fallback: keep full response as recommendation
                why_explanation = None

        # cellar_matches wird jetzt direkt im use_cellar Block gesetzt (siehe oben)
        
        pairing = PairingResponse(
            dish=request.dish,
            recommendation=response,
            why_explanation=why_explanation,
            cellar_matches=cellar_matches
        )
        
        # Cache the result if cacheable
        if is_cacheable and cache_key:
            set_cached_pairing(cache_key, {
                'recommendation': response,
                'why_explanation': why_explanation
            })
        
        # Save pairing to history
        doc = pairing.model_dump()
        doc['created_at'] = doc['created_at'].isoformat()
        await db.pairings.insert_one(doc)
        
        # Increment usage counter
        if user:
            await increment_usage(user, "pairing")
        
        # 🍷 AUTO-ADD: Empfohlene Weine im Hintergrund zur DB hinzufügen
        asyncio.create_task(auto_add_recommended_wines(response, request.dish))
        
        return pairing
        
    except Exception as e:
        logger.error(f"Pairing error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Fehler bei der Empfehlung: {str(e)}")

@api_router.get("/pairings", response_model=List[PairingResponse])
async def get_pairing_history():
    """Get history of wine pairings"""
    pairings = await db.pairings.find({}, {"_id": 0}).sort("created_at", -1).to_list(50)
    for p in pairings:
        if isinstance(p.get('created_at'), str):
            p['created_at'] = datetime.fromisoformat(p['created_at'])
    return pairings

# ===================== CACHE MANAGEMENT =====================

@api_router.get("/cache/status")
async def get_cache_status():
    """Get current cache status and statistics"""
    current_time = time.time()
    valid_entries = sum(1 for v in PAIRING_CACHE.values() if current_time - v['timestamp'] < CACHE_TTL)
    expired_entries = len(PAIRING_CACHE) - valid_entries
    
    return {
        "total_entries": len(PAIRING_CACHE),
        "valid_entries": valid_entries,
        "expired_entries": expired_entries,
        "cache_ttl_hours": CACHE_TTL / 3600,
        "cache_keys": list(PAIRING_CACHE.keys())[:10]  # Show first 10 keys
    }

@api_router.delete("/cache/clear")
async def clear_cache():
    """Clear all cache entries"""
    count = len(PAIRING_CACHE)
    PAIRING_CACHE.clear()
    return {"message": f"Cache cleared. Removed {count} entries."}

@api_router.delete("/cache/expired")
async def clear_expired_cache():
    """Clear only expired cache entries"""
    removed = clear_old_cache_entries()
    return {"message": f"Removed {removed} expired entries. Remaining: {len(PAIRING_CACHE)}"}


# ===================== SITEMAP WITH HREFLANG =====================

# Supported languages for hreflang
SUPPORTED_LANGUAGES = ["de", "en", "fr"]

PAIRING_SITEMAP_ITEMS = [
    {"slug": "lammkoteletts-mit-rosmarin-cabernet-sauvignon", "status": "LIVE", "category": "meat"},
    {"slug": "rinderfilet-mit-kraeuterbutter-und-pommes-bordeaux", "status": "LIVE", "category": "meat"},
    {"slug": "lachsfilet-mit-kraeutersauce-chardonnay", "status": "LIVE", "category": "fish"},
    {"slug": "spaghetti-bolognese-chianti", "status": "LIVE", "category": "pasta"},
    {"slug": "pizza-margherita-chianti", "status": "LIVE", "category": "pizza"},
    {"slug": "pad-thai-riesling", "status": "LIVE", "category": "asian"},
    {"slug": "sushi-sashimi-riesling", "status": "LIVE", "category": "asian"},
    {"slug": "tandoori-chicken-riesling", "status": "LIVE", "category": "asian"},
]

# Static pages that support multiple languages
MULTILINGUAL_PAGES = [
    {"path": "/", "priority": "1.0", "changefreq": "weekly"},
    {"path": "/pairing", "priority": "0.9", "changefreq": "weekly"},
    {"path": "/sommelier-kompass", "priority": "0.9", "changefreq": "weekly"},
    {"path": "/grapes", "priority": "0.8", "changefreq": "monthly"},
    {"path": "/wine-database", "priority": "0.8", "changefreq": "weekly"},
    {"path": "/feed", "priority": "0.8", "changefreq": "daily"},
    {"path": "/blog", "priority": "0.8", "changefreq": "daily"},
    {"path": "/cellar", "priority": "0.7", "changefreq": "daily"},
    {"path": "/favorites", "priority": "0.7", "changefreq": "daily"},
]


def generate_hreflang_links(base_url: str, path: str) -> str:
    """Generate hreflang alternate links for all supported languages"""
    links = []
    # Check if path already has query parameters
    separator = "&" if "?" in path else "?"
    
    for lang in SUPPORTED_LANGUAGES:
        if lang == "de":
            # German is default, no lang parameter needed
            href = f"{base_url}{path}"
        else:
            href = f"{base_url}{path}{separator}lang={lang}"
        links.append(f'    <xhtml:link rel="alternate" hreflang="{lang}" href="{href}"/>')
    # Add x-default (points to German as default)
    links.append(f'    <xhtml:link rel="alternate" hreflang="x-default" href="{base_url}{path}"/>')
    return "\n".join(links)


@api_router.get("/sitemap.xml")
async def sitemap_index():
    """Sitemap index pointing to all sub-sitemaps"""
    base_url = os.environ.get("FRONTEND_BASE_URL", FRONTEND_URL).rstrip("/")
    api_base = base_url + "/api"
    
    xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<sitemapindex xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">
  <sitemap>
    <loc>{api_base}/sitemap-pages.xml</loc>
  </sitemap>
  <sitemap>
    <loc>{api_base}/sitemap-pairings.xml</loc>
  </sitemap>
  <sitemap>
    <loc>{api_base}/sitemap-kompass.xml</loc>
  </sitemap>
</sitemapindex>
"""
    return Response(content=xml, media_type="application/xml")


@api_router.get("/sitemap-pages.xml")
async def sitemap_pages():
    """Main sitemap for static pages with hreflang support"""
    base_url = os.environ.get("FRONTEND_BASE_URL", FRONTEND_URL).rstrip("/")
    
    urls = []
    for page in MULTILINGUAL_PAGES:
        hreflang_links = generate_hreflang_links(base_url, page["path"])
        url_entry = f"""  <url>
    <loc>{base_url}{page["path"]}</loc>
{hreflang_links}
    <changefreq>{page["changefreq"]}</changefreq>
    <priority>{page["priority"]}</priority>
  </url>"""
        urls.append(url_entry)

    xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9"
        xmlns:xhtml="http://www.w3.org/1999/xhtml">
{chr(10).join(urls)}
</urlset>
"""
    return Response(content=xml, media_type="application/xml")


@api_router.get("/sitemap-pairings.xml")
async def sitemap_pairings():
    """Sitemap for SEO pairing pages with hreflang support"""
    base_url = os.environ.get("FRONTEND_BASE_URL", FRONTEND_URL).rstrip("/")
    live_items = [item for item in PAIRING_SITEMAP_ITEMS if item.get("status") == "LIVE"]

    urls = []
    for item in live_items:
        path = f"/pairing/{item['slug']}"
        hreflang_links = generate_hreflang_links(base_url, path)
        url_entry = f"""  <url>
    <loc>{base_url}{path}</loc>
{hreflang_links}
    <changefreq>weekly</changefreq>
    <priority>0.8</priority>
  </url>"""
        urls.append(url_entry)

    xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9"
        xmlns:xhtml="http://www.w3.org/1999/xhtml">
{chr(10).join(urls)}
</urlset>
"""
    return Response(content=xml, media_type="application/xml")


@api_router.get("/sitemap-kompass.xml")
async def sitemap_kompass():
    """Sitemap for Sommelier-Kompass country pages with hreflang support"""
    base_url = os.environ.get("FRONTEND_BASE_URL", FRONTEND_URL).rstrip("/")
    
    # Get all countries from database
    countries = await db.regional_pairings.distinct("country")
    
    urls = []
    
    # Main Kompass page
    path = "/sommelier-kompass"
    hreflang_links = generate_hreflang_links(base_url, path)
    urls.append(f"""  <url>
    <loc>{base_url}{path}</loc>
{hreflang_links}
    <changefreq>weekly</changefreq>
    <priority>0.9</priority>
  </url>""")
    
    # Country-specific pages (with country filter)
    for country in countries:
        country_slug = country.lower().replace("ü", "ue").replace("ö", "oe").replace("ä", "ae")
        path = f"/sommelier-kompass?country={country}"
        hreflang_links = generate_hreflang_links(base_url, path)
        urls.append(f"""  <url>
    <loc>{base_url}{path}</loc>
{hreflang_links}
    <changefreq>monthly</changefreq>
    <priority>0.7</priority>
  </url>""")

    xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9"
        xmlns:xhtml="http://www.w3.org/1999/xhtml">
{chr(10).join(urls)}
</urlset>
"""
    return Response(content=xml, media_type="application/xml")


# ===================== LABEL SCANNER =====================

@api_router.post("/scan-label", response_model=LabelScanResponse)
async def scan_wine_label(request: LabelScanRequest, http_request: Request):
    """Scan a wine label image and extract information (requires authentication)"""
    # Authentication check - must be logged in to use scan feature
    user = await get_current_user(http_request)
    if not user:
        raise HTTPException(status_code=401, detail="Bitte melden Sie sich an, um die Scan-Funktion zu nutzen")
    
    try:
        # Validate base64 image data first
        if not request.image_base64 or not request.image_base64.strip():
            logger.warning("Label scan: Empty image_base64 provided")
            return LabelScanResponse(
                name="Kein Bild",
                type="rot",
                notes="Kein Bild zum Analysieren bereitgestellt"
            )
        
        # Basic base64 validation
        try:
            # Remove data URL prefix if present
            image_data = request.image_base64
            if image_data.startswith('data:'):
                image_data = image_data.split(',', 1)[1] if ',' in image_data else image_data
            
            # Try to decode base64 to validate format
            base64.b64decode(image_data, validate=True)
        except Exception as validation_error:
            logger.warning(f"Label scan: Invalid base64 format: {validation_error}")
            return LabelScanResponse(
                name="Ungültiges Bild",
                type="rot",
                notes="Bildformat nicht erkannt - bitte verwenden Sie ein gültiges Bild"
            )
        
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=str(uuid.uuid4()),
            system_message="Du bist ein Experte für Weinetiketten. Analysiere das Bild und extrahiere die Weininformationen. Antworte NUR im JSON-Format."
        ).with_model("openai", "gpt-5.1")
        
        # WICHTIG: ImageContent benötigt den reinen Base64-String OHNE data:image/... Prefix
        image_content = ImageContent(image_base64=image_data)
        
        prompt = """Analysiere dieses Weinetikett und extrahiere folgende Informationen im JSON-Format:
{
  "name": "Name des Weins",
  "type": "rot/weiss/rose/schaumwein",
  "region": "Herkunftsregion",
  "year": Jahrgang als Zahl oder null,
  "grape": "Rebsorte",
  "notes": "Kurze Beschreibung"
}

WICHTIG: 
- "name" MUSS ein String sein (wenn nicht erkennbar: "Unbekannter Wein")
- "type" MUSS einer von: rot, weiss, rose, schaumwein sein (wenn nicht erkennbar: "rot")
- Andere Felder können null sein"""
        
        user_message = UserMessage(text=prompt, file_contents=[image_content])
        response = await chat.send_message(user_message)
        
        logger.info(f"Label scan: AI response received, length: {len(response) if response else 0}")
        
        # Check if response is None or empty
        if not response or not response.strip():
            logger.warning("Label scan: Received empty or None response from AI")
            return LabelScanResponse(
                name="Nicht erkannt",
                type="rot",
                notes="Keine Antwort vom Sommelier - Bitte versuchen Sie es erneut"
            )
        
        # Extract JSON from response - try multiple patterns
        json_match = re.search(r'\{[^{}]*\}', response, re.DOTALL)
        
        logger.info(f"Label scan: JSON found: {bool(json_match)}")
        
        if json_match:
            try:
                data = json.loads(json_match.group())
                logger.info(f"Label scan: Parsed data: {data}")
                # Ensure required fields have valid defaults
                name = data.get('name') or 'Unbekannter Wein'
                wine_type = data.get('type') or 'rot'
                
                # Validate wine type
                valid_types = ['rot', 'weiss', 'rose', 'schaumwein']
                if wine_type.lower() not in valid_types:
                    wine_type = 'rot'
                
                logger.info(f"Label scan: Returning LabelScanResponse - name={name}, type={wine_type}, region={data.get('region')}, year={data.get('year')}")
                return LabelScanResponse(
                    name=str(name),
                    type=wine_type.lower(),
                    region=data.get('region') if data.get('region') else None,
                    year=int(data['year']) if data.get('year') and str(data['year']).isdigit() else None,
                    grape=data.get('grape') if data.get('grape') else None,
                    notes=data.get('notes') if data.get('notes') else None
                )
            except (json.JSONDecodeError, ValueError, KeyError, TypeError) as parse_error:
                logger.warning(f"JSON parse error: {parse_error}, response: {response[:200]}")
                return LabelScanResponse(
                    name="Nicht erkannt",
                    type="rot",
                    notes=f"Konnte Etikett nicht vollständig analysieren: {str(parse_error)[:100]}"
                )
        else:
            logger.warning(f"Label scan: No JSON found in response: {response[:200]}")
            return LabelScanResponse(
                name="Nicht erkannt",
                type="rot",
                notes=f"Konnte keine strukturierten Daten extrahieren. Antwort: {response[:150]}"
            )
            
    except Exception as e:
        error_message = str(e).lower()
        logger.error(f"Label scan error: {str(e)}")
        
        # Handle specific LLM API errors more gracefully
        if "invalid base64" in error_message or "unsupported image" in error_message:
            logger.warning(f"Label scan: LLM rejected image format: {str(e)}")
            return LabelScanResponse(
                name="Bildformat nicht unterstützt",
                type="rot",
                notes="Das Bildformat wird nicht unterstützt - bitte verwenden Sie JPG, PNG oder ein anderes gängiges Format"
            )
        elif "badrequest" in error_message or "400" in error_message:
            logger.warning(f"Label scan: Bad request to LLM: {str(e)}")
            return LabelScanResponse(
                name="Anfrage fehlerhaft",
                type="rot",
                notes="Fehler bei der Bildanalyse - bitte versuchen Sie es mit einem anderen Bild"
            )
        else:
            # For other errors, still return 500 but with more user-friendly message
            raise HTTPException(status_code=500, detail="Fehler beim Scannen des Weinetiketts - bitte versuchen Sie es später erneut")

# ===================== SOMMELIER CHAT =====================

@api_router.post("/chat", response_model=ChatResponse)
async def sommelier_chat(request: ChatRequest, http_request: Request):
    """Chat with the virtual sommelier"""
    try:
        # Check user limits
        user = await get_current_user(http_request)
        allowed, message = await check_limit(user, "chat")
        
        if not allowed:
            raise HTTPException(status_code=429, detail=message)
        
        session_id = request.session_id or str(uuid.uuid4())
        
        # Get language-specific system message
        system_message = get_sommelier_system(request.language)
        
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=session_id,
            system_message=system_message
        ).with_model("openai", "gpt-5.1")
        
        # Prepare message with optional image
        if request.image_base64:
            image_content = ImageContent(image_base64=request.image_base64)
            user_message = UserMessage(text=request.message, file_contents=[image_content])
        else:
            user_message = UserMessage(text=request.message)
        
        response = await chat.send_message(user_message)
        
        # Save chat message
        chat_doc = {
            "session_id": session_id,
            "user_message": request.message,
            "assistant_response": response,
            "timestamp": datetime.now(timezone.utc).isoformat()
        }
        await db.chats.insert_one(chat_doc)
        
        # Increment usage counter
        if user:
            await increment_usage(user, "chat")
        
        return ChatResponse(response=response, session_id=session_id)
        
    except Exception as e:
        logger.error(f"Chat error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Fehler im Chat: {str(e)}")

@api_router.get("/chat/{session_id}")
async def get_chat_history(session_id: str):
    """Get chat history for a session"""
    messages = await db.chats.find({"session_id": session_id}, {"_id": 0}).sort("timestamp", 1).to_list(100)
    return messages

# ===================== FAVORITES =====================

# Hinweis: Die ursprüngliche Favorites-Route wurde durch die neue
# /favorites-Implementierung (wine_favorites) weiter unten ersetzt.

# ===================== COMMUNITY FEED ENDPOINTS =====================

@api_router.get("/feed", response_model=List[FeedPost])
async def get_feed_posts(limit: int = 50, skip: int = 0):
    """Get all feed posts, newest first"""
    posts = await db.feed_posts.find({}, {"_id": 0}).sort("created_at", -1).skip(skip).to_list(limit)
    for post in posts:
        if isinstance(post.get('created_at'), str):
            post['created_at'] = datetime.fromisoformat(post['created_at'])
        # Parse comments
        if 'comments' in post:
            for comment in post['comments']:
                if isinstance(comment.get('created_at'), str):
                    comment['created_at'] = datetime.fromisoformat(comment['created_at'])
    return posts

@api_router.get("/feed/{post_id}", response_model=FeedPost)
async def get_feed_post(post_id: str):
    """Get a specific feed post"""
    post = await db.feed_posts.find_one({"id": post_id}, {"_id": 0})
    if not post:
        raise HTTPException(status_code=404, detail="Post nicht gefunden")
    if isinstance(post.get('created_at'), str):
        post['created_at'] = datetime.fromisoformat(post['created_at'])
    return post

@api_router.post("/feed", response_model=FeedPost)
async def create_feed_post(post_data: FeedPostCreate):
    """Create a new feed post"""
    post = FeedPost(**post_data.model_dump())
    doc = post.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    doc['likes'] = []
    doc['comments'] = []
    await db.feed_posts.insert_one(doc)
    return post

@api_router.post("/feed/{post_id}/like")
async def toggle_like(post_id: str, author_id: str):
    """Toggle like on a feed post"""
    post = await db.feed_posts.find_one({"id": post_id}, {"_id": 0})
    if not post:
        raise HTTPException(status_code=404, detail="Post nicht gefunden")
    
    likes = post.get('likes', [])
    if author_id in likes:
        likes.remove(author_id)
        action = "unliked"
    else:
        likes.append(author_id)
        action = "liked"
    
    await db.feed_posts.update_one({"id": post_id}, {"$set": {"likes": likes}})
    return {"action": action, "likes_count": len(likes)}

@api_router.post("/feed/{post_id}/comment")
async def add_comment(post_id: str, comment_data: FeedCommentCreate):
    """Add a comment to a feed post"""
    post = await db.feed_posts.find_one({"id": post_id}, {"_id": 0})
    if not post:
        raise HTTPException(status_code=404, detail="Post nicht gefunden")
    
    comment = FeedComment(
        author_name=comment_data.author_name,
        content=comment_data.content
    )
    comment_doc = comment.model_dump()
    comment_doc['created_at'] = comment_doc['created_at'].isoformat()
    
    await db.feed_posts.update_one(
        {"id": post_id},
        {"$push": {"comments": comment_doc}}
    )
    return {"message": "Kommentar hinzugefügt", "comment": comment_doc}

@api_router.delete("/feed/{post_id}")
async def delete_feed_post(post_id: str, author_id: str):
    """Delete a feed post (only by author)"""
    post = await db.feed_posts.find_one({"id": post_id}, {"_id": 0})
    if not post:
        raise HTTPException(status_code=404, detail="Post nicht gefunden")
    
    if post.get('author_id') != author_id:
        raise HTTPException(status_code=403, detail="Nur der Autor kann diesen Post löschen")
    
    await db.feed_posts.delete_one({"id": post_id})
    return {"message": "Post gelöscht"}

@api_router.post("/admin/grapes/generate", response_model=GrapeVariety)
async def generate_grape_variety(request: GrapeGenerationRequest):
    """Generate a new grape variety entry via LLM in a normalized structure.

    Hinweis: Interner Admin-Endpoint, keine Authentifizierung hier im Prototyp.
    """
    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=str(uuid.uuid4()),
            system_message=GRAPE_GENERATOR_SYSTEM
        ).with_model("openai", "gpt-5.1")

        # Prompt für die zu generierende Rebsorte
        base_prompt = f"Erzeuge einen vollständigen Rebsorten-Datensatz für die Rebsorte '{request.name}'."
        if request.grape_type:
            base_prompt += f" Die Rebsorte ist ein {request.grape_type}-wein."
        if request.style_hint:
            base_prompt += f" Stil-Hinweis: {request.style_hint}."

        user_message = UserMessage(text=base_prompt)
        raw_response = await chat.send_message(user_message)

        if not raw_response or not raw_response.strip():
            raise HTTPException(status_code=500, detail="Leere Antwort vom LLM bei der Rebsorten-Generierung")

        # JSON aus Antwort extrahieren
        json_match = re.search(r"\{[\s\S]*\}", raw_response)
        if not json_match:
            raise HTTPException(status_code=500, detail="Konnte keine JSON-Struktur aus der LLM-Antwort extrahieren")

        data = json.loads(json_match.group())

        # Fallbacks & Normalisierung
        slug = data.get("slug") or re.sub(r"[^a-z0-9-]", "", data.get("name", request.name).lower().replace(" ", "-"))
        grape_type = data.get("type") or (request.grape_type or "weiss")
        if grape_type not in ["rot", "weiss"]:
            grape_type = "weiss"

        def ensure_list(value):
            if not value:
                return []
            if isinstance(value, list):
                return value
            return [str(value)]

        grape_payload = {
            "slug": slug,
            "name": data.get("name", request.name),
            "type": grape_type,
            "description": data.get("description", ""),
            "description_en": data.get("description_en"),
            "description_fr": data.get("description_fr"),
            "synonyms": ensure_list(data.get("synonyms")),
            "body": data.get("body", "mittel"),
            "acidity": data.get("acidity", "mittel"),
            "tannin": data.get("tannin", "mittel" if grape_type == "rot" else "niedrig"),
            "aging": data.get("aging", ""),
            "primary_aromas": ensure_list(data.get("primary_aromas")),
            "tertiary_aromas": ensure_list(data.get("tertiary_aromas")),
            "perfect_pairings": ensure_list(data.get("perfect_pairings")),
            "perfect_pairings_en": ensure_list(data.get("perfect_pairings_en")),
            "perfect_pairings_fr": ensure_list(data.get("perfect_pairings_fr")),
            "main_regions": ensure_list(data.get("main_regions")),
            "image_url": None,
        }

        grape = GrapeVariety(**grape_payload)
        doc = grape.model_dump()
        doc["created_at"] = doc["created_at"].isoformat()
        await db.grape_varieties.insert_one(doc)

        return grape

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating grape variety: {e}")
        raise HTTPException(status_code=500, detail="Fehler bei der Rebsorten-Generierung")


# ===================== DISH HELPERS & ENDPOINTS =====================

async def _ensure_dish_indexes():
    """Create helpful indexes for dishes collection (idempotent)."""
    try:
        await db.dishes.create_index("slug", unique=True)
        await db.dishes.create_index([("country", 1), ("bestseller_category", 1)])
        await db.dishes.create_index("trend_cuisines")
    except Exception as e:
        logger.warning(f"Could not create dish indexes: {e}")


def _normalize_tag_list(values: Optional[list]) -> list:
    if not values:
        return []
    seen = set()
    result = []
    for v in values:
        if not isinstance(v, str):
            continue
        s = v.strip().lower()
        if not s:
            continue
        if s not in seen:
            seen.add(s)
            result.append(s)
    return result


def _normalize_scale(value: Optional[str], allowed: List[str]) -> Optional[str]:
    if not value:
        return None
    v = value.strip().lower()
    for a in allowed:
        if a in v:
            return a
    return v if v in allowed else None


async def _generate_dish_from_seed(seed: dict) -> Optional[Dish]:
    chat = LlmChat(
        api_key=EMERGENT_LLM_KEY,
        session_id=str(uuid.uuid4()),
        system_message=DISH_GENERATOR_SYSTEM
    ).with_model("openai", "gpt-5.1")

    base_name = seed.get("base_name")
    country_hint = seed.get("country_hint")
    trend_hint = seed.get("trend_hint")
    bestseller_category = seed.get("bestseller_category")

    prompt = f"Erzeuge einen strukturierten Datensatz für das Gericht '{base_name}'."
    if country_hint:
        prompt += f" Land: {country_hint}."
    if trend_hint:
        prompt += f" Trend-Küche: {trend_hint}."
    if bestseller_category:
        prompt += f" Bestseller-Kategorie: {bestseller_category}."

    user_message = UserMessage(text=prompt)
    raw_response = await chat.send_message(user_message)

    if not raw_response or not raw_response.strip():
        logger.warning(f"Empty LLM response for dish seed {base_name}")
        return None

    json_match = re.search(r"\{[\s\S]*\}", raw_response)
    if not json_match:
        logger.warning(f"No JSON found in LLM response for dish seed {base_name}: {raw_response[:200]}")
        return None

    data = json.loads(json_match.group())

    # Normalize fields
    slug = data.get("slug") or re.sub(r"[^a-z0-9-]", "", data.get("name_de", base_name).lower().replace(" ", "-"))

    trend_cuisines = _normalize_tag_list(data.get("trend_cuisines"))
    key_aromas = _normalize_tag_list(data.get("key_aromas"))
    texture = _normalize_tag_list(data.get("texture"))

    intensity = _normalize_scale(data.get("intensity"), ["leicht", "mittel", "kräftig"])
    fat_level = _normalize_scale(data.get("fat_level"), ["niedrig", "mittel", "hoch"])
    acid_level = _normalize_scale(data.get("acid_level"), ["niedrig", "mittel", "hoch"])
    sweetness_level = _normalize_scale(data.get("sweetness_level"), ["trocken", "leicht_süß", "süß"])
    spice_level = _normalize_scale(data.get("spice_level"), ["keine", "leicht", "mittel", "dominant"])

    dish_payload = {
        "slug": slug,
        "name_de": data.get("name_de", base_name),
        "name_en": data.get("name_en"),
        "name_fr": data.get("name_fr"),
        "country": (data.get("country") or seed.get("country_hint")).lower() if data.get("country") or seed.get("country_hint") else None,
        "region": data.get("region"),
        "trend_cuisines": trend_cuisines,
        "bestseller_category": data.get("bestseller_category") or seed.get("bestseller_category"),
        "protein": data.get("protein"),
        "intensity": intensity,
        "cooking_method": data.get("cooking_method"),
        "sauce_base": data.get("sauce_base"),
        "fat_level": fat_level,
        "acid_level": acid_level,
        "sweetness_level": sweetness_level,
        "spice_level": spice_level,
        "key_aromas": key_aromas,
        "texture": texture,
    }

    dish = Dish(**dish_payload)
    doc = dish.model_dump()
    doc["created_at"] = doc["created_at"].isoformat()
    await db.dishes.insert_one(doc)
    return dish


INITIAL_DISH_SEEDS = [
    # Bestseller international
    {"base_name": "Cheeseburger", "country_hint": "usa", "trend_hint": "streetfood", "bestseller_category": "burger"},
    {"base_name": "Classic Burger", "country_hint": "usa", "trend_hint": "streetfood", "bestseller_category": "burger"},
    {"base_name": "Pizza Margherita", "country_hint": "italien", "trend_hint": "pizzeria", "bestseller_category": "pizza"},
    {"base_name": "Pizza Salami", "country_hint": "italien", "trend_hint": "pizzeria", "bestseller_category": "pizza"},
    {"base_name": "Spaghetti Bolognese", "country_hint": "italien", "trend_hint": "trattoria", "bestseller_category": "pasta"},
    {"base_name": "Spaghetti Carbonara", "country_hint": "italien", "trend_hint": "trattoria", "bestseller_category": "pasta"},
    {"base_name": "Rinderfilet mit Rotwein-Jus", "country_hint": "frankreich", "trend_hint": "fine_dining", "bestseller_category": "steak"},
    {"base_name": "Ribeye Steak vom Grill", "country_hint": "usa", "trend_hint": "bbq", "bestseller_category": "steak"},
    {"base_name": "Lachsfilet mit Zitronen-Butter-Sauce", "country_hint": "international", "trend_hint": "brasserie", "bestseller_category": "fisch"},
    {"base_name": "Fish and Chips", "country_hint": "uk", "trend_hint": "streetfood", "bestseller_category": "fisch"},
    {"base_name": "Caesar Salad mit Huhn", "country_hint": "usa", "trend_hint": "bistro", "bestseller_category": "salat"},
    {"base_name": "Sushi Mix", "country_hint": "japan", "trend_hint": "sushi", "bestseller_category": "sushi"},
    {"base_name": "Ramen mit Schweinebauch", "country_hint": "japan", "trend_hint": "nudelsuppe", "bestseller_category": "nudelsuppe"},
    {"base_name": "Pad Thai mit Garnelen", "country_hint": "thailand", "trend_hint": "thai", "bestseller_category": "nudelgericht"},
    {"base_name": "Grünes Thai-Curry mit Huhn", "country_hint": "thailand", "trend_hint": "thai", "bestseller_category": "curry"},
    {"base_name": "Indisches Butter Chicken", "country_hint": "indien", "trend_hint": "indisch", "bestseller_category": "curry"},
    {"base_name": "Tacos al Pastor", "country_hint": "mexiko", "trend_hint": "streetfood", "bestseller_category": "taco"},
    {"base_name": "Falafel Bowl", "country_hint": "orient", "trend_hint": "bowl", "bestseller_category": "bowl"},
    {"base_name": "Vegane Buddha Bowl", "country_hint": "international", "trend_hint": "bowl", "bestseller_category": "bowl"},
    {"base_name": "Pizza Prosciutto e Funghi", "country_hint": "italien", "trend_hint": "pizzeria", "bestseller_category": "pizza"},
    # Länderfokus Europa
    {"base_name": "Coq au Vin", "country_hint": "frankreich", "trend_hint": "klassisch", "bestseller_category": "geflügel"},
    {"base_name": "Boeuf Bourguignon", "country_hint": "frankreich", "trend_hint": "klassisch", "bestseller_category": "schmorgericht"},
    {"base_name": "Paella mit Meeresfrüchten", "country_hint": "spanien", "trend_hint": "mediterran", "bestseller_category": "reisgericht"},
    {"base_name": "Tapas-Auswahl", "country_hint": "spanien", "trend_hint": "sharing", "bestseller_category": "tapas"},
    {"base_name": "Schweinsbraten mit Knödeln", "country_hint": "deutschland", "trend_hint": "hausmannskost", "bestseller_category": "schmorgericht"},
    {"base_name": "Wiener Schnitzel mit Kartoffelsalat", "country_hint": "österreich", "trend_hint": "klassisch", "bestseller_category": "schnitzel"},
    {"base_name": "Moussaka", "country_hint": "griechenland", "trend_hint": "mediterran", "bestseller_category": "auflauf"},
    # Asien & Trendküchen
    {"base_name": "Pho Bo", "country_hint": "vietnam", "trend_hint": "streetfood", "bestseller_category": "nudelsuppe"},
    {"base_name": "Koreanisches Bibimbap", "country_hint": "korea", "trend_hint": "bowl", "bestseller_category": "bowl"},
    {"base_name": "Mapo Tofu", "country_hint": "china", "trend_hint": "scharf", "bestseller_category": "veggie"},
    {"base_name": "Kung Pao Chicken", "country_hint": "china", "trend_hint": "asiatisch", "bestseller_category": "geflügel"},
    # USA & Amerika
    {"base_name": "BBQ Ribs", "country_hint": "usa", "trend_hint": "bbq", "bestseller_category": "fleisch"},
    {"base_name": "Mac and Cheese", "country_hint": "usa", "trend_hint": "comfort_food", "bestseller_category": "beilage"},
    {"base_name": "Argentinisches Asado", "country_hint": "argentinien", "trend_hint": "bbq", "bestseller_category": "fleisch"},
    {"base_name": "Ceviche", "country_hint": "peru", "trend_hint": "seafood", "bestseller_category": "fisch"},
    # Vegetarisch/Vegan
    {"base_name": "Gemüse-Lasagne", "country_hint": "italien", "trend_hint": "vegetarisch", "bestseller_category": "auflauf"},
    {"base_name": "Kichererbsen-Curry", "country_hint": "indien", "trend_hint": "vegan", "bestseller_category": "curry"},
    {"base_name": "Gegrilltes Gemüse mit Halloumi", "country_hint": "griechenland", "trend_hint": "vegetarisch", "bestseller_category": "gemüse"},
]


@api_router.post("/admin/dishes/generate", response_model=Dish)
async def generate_dish(request: DishGenerationRequest):
    """Generate a single structured dish entry via LLM."""
    seed = {
        "base_name": request.base_name,
        "country_hint": request.country_hint,
        "trend_hint": request.trend_hint,
        "bestseller_category": request.bestseller_category,
    }
    await _ensure_dish_indexes()
    dish = await _generate_dish_from_seed(seed)
    if not dish:
        raise HTTPException(status_code=500, detail="Gericht konnte nicht generiert werden")
    return dish


async def _run_dish_seed_batch():
    await _ensure_dish_indexes()
    created = 0
    for seed in INITIAL_DISH_SEEDS:
        try:
            # Skip if slug already exists
            slug_candidate = re.sub(r"[^a-z0-9-]", "", seed["base_name"].lower().replace(" ", "-"))
            exists = await db.dishes.find_one({"slug": slug_candidate})
            if exists:
                continue
            dish = await _generate_dish_from_seed(seed)
            if dish:
                created += 1
        except Exception as e:
            logger.warning(f"Error seeding dish {seed.get('base_name')}: {e}")
    logger.info(f"Dish seed batch completed, created {created} dishes")


@api_router.post("/admin/dishes/seed-batch")
async def seed_dishes(background_tasks: BackgroundTasks):
    """Trigger background seeding of a larger set of structured dishes.

    Läuft im Hintergrund, die HTTP-Antwort kommt sofort zurück.
    """
    background_tasks.add_task(_run_dish_seed_batch)
    return {"status": "started", "count": len(INITIAL_DISH_SEEDS)}


@api_router.get("/dishes", response_model=List[Dish])
async def list_dishes(country: Optional[str] = None, bestseller_category: Optional[str] = None, trend: Optional[str] = None):
    """Simple listing endpoint - später für UI & Suche nutzbar."""
    query: dict = {}
    if country:
        query["country"] = country.lower()
    if bestseller_category:
        query["bestseller_category"] = bestseller_category
    if trend:
        query["trend_cuisines"] = trend.lower()

    dishes = await db.dishes.find(query, {"_id": 0}).sort("name_de", 1).to_list(500)
    for d in dishes:
        if isinstance(d.get("created_at"), str):
            d["created_at"] = datetime.fromisoformat(d["created_at"])
    
    return dishes

@api_router.get("/feed-stats")
async def get_feed_stats():
    """Get feed statistics"""
    total_posts = await db.feed_posts.count_documents({})
    total_users = len(await db.feed_posts.distinct("author_id"))
    
    # Top rated pairings
    pipeline = [
        {"$match": {"rating": {"$gte": 4}}},
        {"$sort": {"created_at": -1}},
        {"$limit": 5},
        {"$project": {"_id": 0, "dish": 1, "wine_name": 1, "rating": 1, "author_name": 1}}
    ]
    top_pairings = await db.feed_posts.aggregate(pipeline).to_list(5)
    
    return {
        "total_posts": total_posts,
        "total_users": total_users,
        "top_pairings": top_pairings
    }

# ===================== REGIONAL PAIRINGS ENDPOINTS =====================

class RegionalPairingsResponse(BaseModel):
    """Response model for paginated regional pairings"""
    pairings: List[RegionalPairing]
    total: int
    has_more: bool


@api_router.get("/regional-pairings")
async def get_regional_pairings(
    country: Optional[str] = None,
    region: Optional[str] = None,
    search: Optional[str] = None,
    limit: int = 50,
    skip: int = 0
):
    """Get regional wine pairings with filters and pagination"""
    query = {}
    
    if country:
        query["country"] = country
    if region:
        query["region"] = region
    if search:
        # WICHTIG: Akzent-insensitive Suche verwenden!
        accent_pattern = create_accent_insensitive_pattern(search)
        query["$or"] = [
            {"dish": {"$regex": accent_pattern, "$options": "i"}},
            {"wine_name": {"$regex": accent_pattern, "$options": "i"}}
        ]
    
    # Hole Gesamtanzahl
    total = await db.regional_pairings.count_documents(query)
    
    # Hole paginierte Ergebnisse
    pairings = await db.regional_pairings.find(query, {"_id": 0}).skip(skip).limit(limit).to_list(limit)
    
    return {
        "pairings": pairings,
        "total": total,
        "has_more": (skip + len(pairings)) < total
    }


@api_router.get("/regional-pairings/countries")
async def get_countries():
    """Get list of all countries with pairing counts - grouped only by country name"""
    pipeline = [
        {
            "$group": {
                "_id": "$country",
                "count": {"$sum": 1},
                "country_emoji": {"$first": "$country_emoji"},
                "country_en": {"$first": "$country_en"},
                "country_fr": {"$first": "$country_fr"},
                "image_url": {"$first": "$image_url"}
            }
        },
        {
            "$project": {
                "_id": 0,
                "country": "$_id",
                "country_en": 1,
                "country_fr": 1,
                "country_emoji": 1,
                "image_url": 1,
                "count": 1
            }
        },
        {"$sort": {"count": -1}}
    ]
    
    countries = await db.regional_pairings.aggregate(pipeline).to_list(30)
    return countries


@api_router.get("/regional-pairings/regions")
async def get_regions(country: str):
    """Get list of regions for a specific country"""
    pipeline = [
        {"$match": {"country": country}},
        {
            "$group": {
                "_id": "$region",
                "count": {"$sum": 1}
            }
        },
        {
            "$project": {
                "_id": 0,
                "region": "$_id",
                "count": 1
            }
        },
        {"$sort": {"region": 1}}
    ]
    
    regions = await db.regional_pairings.aggregate(pipeline).to_list(50)
    return regions


# ===================== SEO PAIRINGS ENDPOINTS =====================

@api_router.get("/seo-pairings")
async def get_seo_pairings(
    limit: int = 50,
    offset: int = 0,
    category: Optional[str] = None,
    region: Optional[str] = None
):
    """Get programmatic SEO pairings for landing pages"""
    query = {}
    
    if category:
        query["dish.category"] = category
    if region:
        query["wine.region"] = {"$regex": region, "$options": "i"}
    
    pairings = await db.seo_pairings.find(query, {"_id": 0}).skip(offset).limit(limit).to_list(limit)
    total = await db.seo_pairings.count_documents(query)
    
    return {
        "pairings": pairings,
        "total": total,
        "limit": limit,
        "offset": offset
    }

@api_router.get("/seo-pairings/categories")
async def get_seo_pairing_categories():
    """Get all dish categories with counts"""
    pipeline = [
        {"$group": {"_id": "$dish.category", "count": {"$sum": 1}}},
        {"$sort": {"count": -1}}
    ]
    categories = await db.seo_pairings.aggregate(pipeline).to_list(100)
    return [{"category": c["_id"], "count": c["count"]} for c in categories if c["_id"]]

@api_router.get("/seo-pairings/{slug}")
async def get_seo_pairing(slug: str):
    """Get a specific SEO pairing by slug"""
    pairing = await db.seo_pairings.find_one({"slug": slug}, {"_id": 0})
    if not pairing:
        raise HTTPException(status_code=404, detail="Pairing not found")
    return pairing


# ===================== GRAPE VARIETY ENDPOINTS =====================

@api_router.get("/grapes", response_model=List[GrapeVariety])
async def get_grape_varieties(type_filter: Optional[str] = None):
    """Get all grape varieties"""
    query = {}
    if type_filter and type_filter != 'all':
        query["type"] = type_filter
    
    grapes = await db.grape_varieties.find(query, {"_id": 0}).sort("name", 1).to_list(500)
    for grape in grapes:
        if isinstance(grape.get('created_at'), str):
            grape['created_at'] = datetime.fromisoformat(grape['created_at'])
    return grapes

@api_router.get("/grapes/{slug}", response_model=GrapeVariety)
async def get_grape_variety(slug: str):
    """Get a specific grape variety by slug"""
    grape = await db.grape_varieties.find_one({"slug": slug}, {"_id": 0})
    if not grape:
        raise HTTPException(status_code=404, detail="Rebsorte nicht gefunden")
    if isinstance(grape.get('created_at'), str):
        grape['created_at'] = datetime.fromisoformat(grape['created_at'])
    return grape
@api_router.get("/admin/users/repair")
async def repair_users_get():
    """
    GET Version des Repair-Endpoints - kann direkt im Browser aufgerufen werden.
    URL: /api/admin/users/repair
    """
    users = await db.users.find({}).to_list(1000)
    repaired = 0
    password_reset = []
    
    temp_password = "WeinPairing2025!"
    temp_hash = hash_password(temp_password)
    
    for user in users:
        updates = {}
        
        if not user.get('user_id'):
            updates['user_id'] = f"user_{uuid.uuid4().hex[:12]}"
        if not user.get('plan'):
            updates['plan'] = 'basic'
        if not user.get('usage'):
            updates['usage'] = {"pairing_requests_today": 0, "chat_messages_today": 0, "last_usage_date": None}
        if not user.get('name') and user.get('email'):
            updates['name'] = user['email'].split('@')[0]
        if not user.get('password_hash'):
            updates['password_hash'] = temp_hash
            password_reset.append(user.get('email', 'unknown'))
        
        if updates:
            await db.users.update_one({"_id": user["_id"]}, {"$set": updates})
            repaired += 1
    
    return {
        "status": "success",
        "total_users": len(users),
        "repaired": repaired,
        "password_reset_users": password_reset,
        "temp_password": temp_password if password_reset else None,
        "message": f"✅ {repaired} User repariert, {len(password_reset)} Passwörter zurückgesetzt",
        "next_step": "Loggen Sie sich jetzt mit dem temporären Passwort ein und ändern Sie es!"
    }

@api_router.get("/admin/reset-owner-simple")
async def reset_owner_simple():
    """
    Setzt das Passwort auf ein einfaches Passwort: Test1234!
    URL: /api/admin/reset-owner-simple
    """
    owner_email = "isicel@bluewin.ch"
    simple_password = "Test1234!"
    
    user = await db.users.find_one({"email": owner_email})
    if not user:
        return {"status": "error", "message": f"User {owner_email} nicht gefunden"}
    
    new_hash = hash_password(simple_password)
    result = await db.users.update_one(
        {"email": owner_email},
        {"$set": {"password_hash": new_hash}}
    )
    
    return {
        "status": "success",
        "message": f"✅ Einfaches Passwort gesetzt!",
        "email": owner_email,
        "new_password": simple_password,
        "modified": result.modified_count
    }


@api_router.get("/admin/reset-owner-password")
async def reset_owner_password():
    """
    EINMALIG: Setzt das Passwort für den Hauptbesitzer zurück.
    URL: /api/admin/reset-owner-password
    
    ⚠️ WICHTIG: Diesen Endpoint nach erfolgreicher Verwendung deaktivieren!
    """
    owner_email = "isicel@bluewin.ch"
    temp_password = "Admin2025!"
    
    # Find owner
    user = await db.users.find_one({"email": owner_email})
    if not user:
        return {"status": "error", "message": f"User {owner_email} nicht gefunden"}
    
    # Reset password using bcrypt
    new_hash = hash_password(temp_password)
    
    # Force update with $set to ensure it's changed
    result = await db.users.update_one(
        {"email": owner_email},
        {"$set": {
            "password_hash": new_hash,
            "plan": "pro",
            "subscription_status": "active",
            "role": "admin",
            "is_admin": True
        }}
    )
    
    # Verify the change
    updated_user = await db.users.find_one({"email": owner_email})
    hash_preview = updated_user.get('password_hash', '')[:20] if updated_user else 'N/A'
    
    return {
        "status": "success",
        "message": f"✅ Passwort für {owner_email} wurde zurückgesetzt!",
        "email": owner_email,
        "new_password": temp_password,
        "plan": "pro",
        "hash_preview": hash_preview + "...",
        "modified": result.modified_count,
        "next_step": "Bitte loggen Sie sich ein und ändern Sie das Passwort!"
    }

@api_router.get("/admin/debug-user/{email}")
async def debug_user(email: str):
    """
    Debug endpoint to check user status.
    URL: /api/admin/debug-user/isicel@bluewin.ch
    """
    user = await db.users.find_one({"email": email})
    if not user:
        return {"status": "error", "message": f"User {email} nicht gefunden"}
    
    # Check password hash format
    pw_hash = user.get('password_hash', '')
    is_bcrypt = pw_hash.startswith('$2') if pw_hash else False
    
    return {
        "email": user.get('email'),
        "name": user.get('name'),
        "plan": user.get('plan'),
        "role": user.get('role'),
        "has_password_hash": bool(pw_hash),
        "hash_length": len(pw_hash) if pw_hash else 0,
        "is_bcrypt_format": is_bcrypt,
        "hash_preview": pw_hash[:30] + "..." if pw_hash else "NONE"
    }

@api_router.get("/admin/users/health")
async def check_users_health():
    """
    ADMIN: Überprüft die Gesundheit aller User-Dokumente.
    Findet und repariert fehlende Felder.
    """
    users = await db.users.find({}, {"_id": 0, "password_hash": 0}).to_list(1000)
    
    total = len(users)
    healthy = 0
    issues = []
    
    required_fields = ['user_id', 'email', 'plan']
    
    for user in users:
        user_issues = []
        
        # Check required fields
        for field in required_fields:
            if not user.get(field):
                user_issues.append(f"Feld '{field}' fehlt")
        
        # Check plan value
        if user.get('plan') not in ['basic', 'pro', None]:
            user_issues.append(f"Ungültiger Plan: {user.get('plan')}")
        
        if user_issues:
            issues.append({
                "email": user.get('email', 'UNBEKANNT'),
                "user_id": user.get('user_id', 'FEHLT'),
                "issues": user_issues
            })
        else:
            healthy += 1
    
    return {
        "total_users": total,
        "healthy_users": healthy,
        "users_with_issues": len(issues),
        "issues": issues[:50],  # Limit to 50 for response size
        "pro_users": len([u for u in users if u.get('plan') == 'pro']),
        "basic_users": len([u for u in users if u.get('plan') == 'basic'])
    }

@api_router.post("/admin/users/repair-all")
async def repair_all_users():
    """
    ADMIN: Repariert alle User mit fehlenden Feldern.
    Setzt temporäres Passwort wenn password_hash fehlt.
    """
    users = await db.users.find({}).to_list(1000)
    repaired = 0
    password_reset = []
    
    # Temporäres Passwort für User ohne password_hash
    temp_password = "WeinPairing2025!"
    temp_hash = hash_password(temp_password)
    
    for user in users:
        updates = {}
        
        # Fix missing user_id
        if not user.get('user_id'):
            updates['user_id'] = f"user_{uuid.uuid4().hex[:12]}"
        
        # Fix missing plan
        if not user.get('plan'):
            updates['plan'] = 'basic'
        
        # Fix missing usage
        if not user.get('usage'):
            updates['usage'] = {
                "pairing_requests_today": 0,
                "chat_messages_today": 0,
                "last_usage_date": None
            }
        
        # Fix missing name
        if not user.get('name') and user.get('email'):
            updates['name'] = user['email'].split('@')[0]
        
        # KRITISCH: Fix missing password_hash
        if not user.get('password_hash'):
            updates['password_hash'] = temp_hash
            password_reset.append(user.get('email', 'unknown'))
        
        if updates:
            await db.users.update_one(
                {"_id": user["_id"]},
                {"$set": updates}
            )
            repaired += 1
    
    return {
        "total_users": len(users),
        "repaired": repaired,
        "password_reset_users": password_reset,
        "temp_password": temp_password if password_reset else None,
        "message": f"{repaired} User repariert, {len(password_reset)} Passwörter zurückgesetzt"
    }

@api_router.post("/admin/grapes/normalize")
async def normalize_grape_varieties():
    """Normalize grape variety fields for search/filter usage.

    - body/acidity/tannin -> one of: leicht/mittel/vollmundig bzw. niedrig/mittel/hoch
    - aroma- und pairing-Listen werden bereinigt und vereinheitlicht
    """
    grapes = await db.grape_varieties.find({}, {"_id": 0}).to_list(1000)
    normalized_count = 0

    def normalize_scale(value: Optional[str], scale: str) -> Optional[str]:
        if not value:
            return None
        v = value.lower()
        if scale == "body":
            if "voll" in v:
                return "vollmundig"
            if "leicht" in v and "mittel" in v:
                return "mittel"
            if "leicht" in v:
                return "leicht"
            if "mittel" in v:
                return "mittel"
            return v
        # acidity or tannin
        if "hoch" in v:
            return "hoch"
        if "niedrig" in v and "mittel" in v:
            return "mittel"
        if "niedrig" in v:
            return "niedrig"
        if "mittel" in v:
            return "mittel"
        return v

    def normalize_string_list(values: Optional[list], lower: bool = True) -> list:
        if not values:
            return []
        seen = set()
        result = []
        for item in values:
            if not isinstance(item, str):
                continue
            s = item.strip()
            if lower:
                s = s.lower()
            if not s:
                continue
            if s not in seen:
                seen.add(s)
                result.append(s)
        return result

    for grape in grapes:
        update: dict = {}

        # Normalize body/acidity/tannin scales
        body = normalize_scale(grape.get("body"), "body")
        acidity = normalize_scale(grape.get("acidity"), "acidity")
        tannin = normalize_scale(grape.get("tannin"), "tannin") if grape.get("type") == "rot" else grape.get("tannin")

        if body is not None:
            update["body"] = body
        if acidity is not None:
            update["acidity"] = acidity
        if tannin is not None:
            update["tannin"] = tannin

        # Normalize aroma and pairing lists (lowercased tags for Suche)
        update["primary_aromas"] = normalize_string_list(grape.get("primary_aromas"), lower=True)
        update["tertiary_aromas"] = normalize_string_list(grape.get("tertiary_aromas"), lower=True)
        update["perfect_pairings"] = normalize_string_list(grape.get("perfect_pairings"), lower=True)

        if update:
            await db.grape_varieties.update_one({"id": grape["id"]}, {"$set": update})
            normalized_count += 1

    return {"normalized": normalized_count}



@api_router.post("/seed-grapes")
async def seed_grape_varieties():
    """Seed grape variety database with famous varieties"""
    grapes = [
        # WHITE WINES
        {
            "slug": "chardonnay",
            "name": "Chardonnay",
            "type": "weiss",
            "description": "Flüssiger Sonnenaufgang im Glas, golden schimmernd, der die Seele umarmt. In der Nase reife Pfirsiche, cremige Vanille, geröstete Haselnüsse und mineralische Kalksteinfrische. Am Gaumen buttrige Opulenz wie Seide, lebendige Säure mit Zitrone und grünem Apfel – ein Tanz von Fülle und Eleganz, der in langem, vibrierendem Finale nach mehr verlangt. Für den Kenner ein Chamäleon: burgundisch straff oder kalifornisch üppig, stets Spiegel von Winzerhand und Natur.",
            "description_en": "Liquid sunrise in a glass, golden shimmer embracing the soul. On the nose: ripe peaches, creamy vanilla, roasted hazelnuts, and mineral limestone freshness. On the palate: buttery opulence like silk, lively acidity with lemon and green apple – a dance of richness and elegance that demands more in its long, vibrating finish. For the connoisseur, a chameleon: Burgundian taut or Californian lush, always a mirror of winemaker's hand and nature.",
            "description_fr": "Lever de soleil liquide dans le verre, chatoiement doré qui embrasse l'âme. Au nez: pêches mûres, vanille crémeuse, noisettes grillées et fraîcheur minérale de calcaire. En bouche: opulence beurrée comme de la soie, acidité vive avec citron et pomme verte – une danse de richesse et d'élégance qui en redemande dans sa longue finale vibrante.",
            "synonyms": ["Morillon", "Beaunois"],
            "body": "mittel bis vollmundig",
            "acidity": "mittel bis hoch",
            "tannin": "niedrig",
            "aging": "Holz oder Edelstahl, trocken",
            "primary_aromas": ["Apfel", "Zitrone", "Pfirsich", "Melone"],
            "tertiary_aromas": ["Butter", "Vanille", "Toast", "Haselnuss"],
            "perfect_pairings": ["Gegrillter Hummer in Zitronenbutter", "Perlhuhn mit Trüffelrisotto", "Reifer Comté"],
            "perfect_pairings_en": ["Grilled lobster in lemon butter", "Guinea fowl with truffle risotto", "Aged Comté cheese"],
            "perfect_pairings_fr": ["Homard grillé au beurre citronné", "Pintade au risotto à la truffe", "Comté affiné"],
            "main_regions": ["Burgund", "Champagne", "Kalifornien", "Australien"],
            "image_url": "https://images.unsplash.com/photo-1566995541428-f2246c17cda1?w=800"
        },
        {
            "slug": "riesling",
            "name": "Riesling",
            "type": "weiss",
            "description": "Die Königin der weißen Reben – kristallklar wie ein Gebirgsbach, elektrisierend und lebendig. Ein Feuerwerk aus grünem Apfel, Limette, weißem Pfirsich und dem unverwechselbaren Hauch von Petrol, der Kennerherzen höher schlagen lässt. Am Gaumen eine Symphonie aus messerscharfer Säure und zarter Süße, perfekt balanciert wie ein Seiltänzer über den Weinbergen der Mosel. Vom knochentrocken bis edelsüß – Riesling ist der Beweis, dass wahre Eleganz zeitlos ist.",
            "description_en": "The queen of white grapes – crystal clear like a mountain stream, electrifying and alive. A firework of green apple, lime, white peach, and that unmistakable hint of petrol that makes connoisseurs' hearts beat faster. On the palate, a symphony of razor-sharp acidity and delicate sweetness, perfectly balanced like a tightrope walker above the Moselle vineyards. From bone dry to noble sweet – Riesling proves that true elegance is timeless.",
            "description_fr": "La reine des cépages blancs – cristallin comme un ruisseau de montagne, électrisant et vivant. Un feu d'artifice de pomme verte, citron vert, pêche blanche et cette touche incomparable de pétrole qui fait battre le cœur des connaisseurs. En bouche, une symphonie d'acidité tranchante et de douceur délicate, parfaitement équilibrée.",
            "synonyms": ["Rheinriesling", "Weißer Riesling"],
            "body": "leicht bis mittel",
            "acidity": "hoch",
            "tannin": "niedrig",
            "aging": "Edelstahl, trocken bis edelsüß",
            "primary_aromas": ["Grüner Apfel", "Limette", "Pfirsich", "Aprikose"],
            "tertiary_aromas": ["Petrol", "Honig", "Ingwer", "Mandel"],
            "perfect_pairings": ["Gebratene Forelle mit Mandelbutter", "Schweineschnitzel mit Spargel", "Thai-Curry mit Garnelen"],
            "perfect_pairings_en": ["Pan-fried trout with almond butter", "Pork schnitzel with asparagus", "Thai curry with shrimp"],
            "perfect_pairings_fr": ["Truite poêlée au beurre d'amandes", "Escalope de porc aux asperges", "Curry thaï aux crevettes"],
            "main_regions": ["Mosel", "Rheingau", "Elsass", "Clare Valley"],
            "image_url": "https://images.unsplash.com/photo-1558001373-7b93ee48ffa0?w=800"
        },
        {
            "slug": "sauvignon-blanc",
            "name": "Sauvignon Blanc",
            "type": "weiss",
            "description": "Ein Weckruf für die Sinne – frisch wie der erste Frühlingsmorgen, wild wie ungezähmte Natur. Stachelbeere, frisch gemähtes Gras, Holunderblüte und ein Hauch von Feuerstein explodieren im Glas. Am Gaumen knackig und präzise, mit einer Säure, die wie ein Blitz durch den Körper fährt. Neuseeland macht ihn exotisch mit Passionsfrucht, die Loire adelt ihn mit mineralischer Tiefe. Sauvignon Blanc ist der Espresso unter den Weißweinen – kompromisslos wach machend.",
            "description_en": "A wake-up call for the senses – fresh as the first spring morning, wild as untamed nature. Gooseberry, freshly cut grass, elderflower, and a hint of flint explode in the glass. On the palate: crisp and precise, with acidity that strikes like lightning through the body. New Zealand makes it exotic with passion fruit, the Loire ennobles it with mineral depth. Sauvignon Blanc is the espresso of white wines – uncompromisingly awakening.",
            "description_fr": "Un réveil pour les sens – frais comme le premier matin de printemps, sauvage comme la nature indomptée. Groseille à maquereau, herbe fraîchement coupée, fleur de sureau et une touche de silex explosent dans le verre. En bouche: croquant et précis, avec une acidité qui frappe comme l'éclair.",
            "synonyms": ["Fumé Blanc", "Blanc Fumé"],
            "body": "leicht bis mittel",
            "acidity": "hoch",
            "tannin": "niedrig",
            "aging": "Edelstahl, trocken",
            "primary_aromas": ["Stachelbeere", "Gras", "Holunderblüte", "Limette"],
            "tertiary_aromas": ["Feuerstein", "Passionsfrucht", "Grapefruit"],
            "perfect_pairings": ["Ziegenkäse-Salat mit Walnüssen", "Austern auf Eis", "Gegrillter Wolfsbarsch mit Kräutern"],
            "perfect_pairings_en": ["Goat cheese salad with walnuts", "Oysters on ice", "Grilled sea bass with herbs"],
            "perfect_pairings_fr": ["Salade de chèvre aux noix", "Huîtres sur glace", "Bar grillé aux herbes"],
            "main_regions": ["Loire", "Neuseeland", "Bordeaux", "Chile"],
            "image_url": "https://images.unsplash.com/photo-1474722883778-792e7990302f?w=800"
        },
        {
            "slug": "gruener-veltliner",
            "name": "Grüner Veltliner",
            "type": "weiss",
            "description": "Österreichs flüssiges Gold – pfeffrig-würzig wie ein Gewürzhändler auf dem Naschmarkt, mit dem unverwechselbaren weißen Pfeffer, der Gaumenkribbeln garantiert. Grüner Apfel, Birne, weiße Kräuter und ein Hauch von Tabak vereinen sich zu einem Wein, der bodenständig und sophisticated zugleich ist. Am Gaumen cremig mit spritziger Säure, perfekt zu Wiens kulinarischen Schätzen. Das Wiener Schnitzel hat keinen besseren Freund.",
            "description_en": "Austria's liquid gold – peppery-spicy like a spice merchant at the Naschmarkt, with the unmistakable white pepper that guarantees tingling on the palate. Green apple, pear, white herbs, and a hint of tobacco unite in a wine that is down-to-earth and sophisticated at once. On the palate: creamy with zesty acidity, perfect with Vienna's culinary treasures. Wiener Schnitzel has no better friend.",
            "description_fr": "L'or liquide d'Autriche – poivré et épicé comme un marchand d'épices au Naschmarkt, avec ce poivre blanc incomparable qui garantit des picotements au palais. Pomme verte, poire, herbes blanches et une touche de tabac s'unissent dans un vin à la fois terre-à-terre et sophistiqué.",
            "synonyms": ["Weißgipfler", "Grüner Muskateller (falsch)"],
            "body": "leicht bis mittel",
            "acidity": "mittel bis hoch",
            "tannin": "niedrig",
            "aging": "Edelstahl oder großes Holz, trocken",
            "primary_aromas": ["Grüner Apfel", "Birne", "Weißer Pfeffer", "Kräuter"],
            "tertiary_aromas": ["Honig", "Tabak", "Nuss"],
            "perfect_pairings": ["Wiener Schnitzel mit Kartoffelsalat", "Spargel mit Sauce Hollandaise", "Gebackener Karpfen"],
            "perfect_pairings_en": ["Wiener Schnitzel with potato salad", "Asparagus with Hollandaise sauce", "Breaded carp"],
            "perfect_pairings_fr": ["Schnitzel viennois avec salade de pommes de terre", "Asperges sauce hollandaise", "Carpe panée"],
            "main_regions": ["Wachau", "Weinviertel", "Kamptal", "Kremstal"],
            "image_url": "https://images.unsplash.com/photo-1506377247377-2a5b3b417ebb?w=800"
        },
        {
            "slug": "gewuerztraminer",
            "name": "Gewürztraminer",
            "type": "weiss",
            "description": "Der Parfümeur unter den Rebsorten – betörend wie ein orientalischer Basar, golden wie Bernstein im Sonnenuntergang. Litschi, Rosenblätter, Muskatnuss und kandierter Ingwer umschmeicheln die Nase wie ein seidener Schleier. Am Gaumen üppig und exotisch, mit zarter Restsüße und cremiger Textur. Ein Wein für Mutige, die sich in ein aromatisches Abenteuer stürzen wollen. Perfekter Begleiter zur asiatischen Küche oder zum Käseplateau.",
            "description_en": "The perfumer among grape varieties – intoxicating like an oriental bazaar, golden like amber at sunset. Lychee, rose petals, nutmeg, and candied ginger caress the nose like a silk veil. On the palate: opulent and exotic, with delicate residual sweetness and creamy texture. A wine for the bold who want to dive into an aromatic adventure. Perfect companion for Asian cuisine or cheese platter.",
            "description_fr": "Le parfumeur parmi les cépages – enivrant comme un bazar oriental, doré comme l'ambre au coucher du soleil. Litchi, pétales de rose, muscade et gingembre confit caressent le nez comme un voile de soie. En bouche: opulent et exotique, avec une délicate sucrosité résiduelle.",
            "synonyms": ["Traminer", "Savagnin Rosé"],
            "body": "mittel bis vollmundig",
            "acidity": "niedrig bis mittel",
            "tannin": "niedrig",
            "aging": "Edelstahl oder Holz, trocken bis lieblich",
            "primary_aromas": ["Litschi", "Rose", "Mango", "Orangenschale"],
            "tertiary_aromas": ["Muskatnuss", "Ingwer", "Honig", "Zimt"],
            "perfect_pairings": ["Ente à l'Orange", "Thai-Curry mit Kokosmilch", "Münsterkäse", "Foie Gras"],
            "perfect_pairings_en": ["Duck à l'Orange", "Thai curry with coconut milk", "Munster cheese", "Foie Gras"],
            "perfect_pairings_fr": ["Canard à l'orange", "Curry thaï au lait de coco", "Munster", "Foie Gras"],
            "main_regions": ["Elsass", "Südtirol", "Deutschland", "Neuseeland"],
            "image_url": "https://images.unsplash.com/photo-1507434965515-61970f2bd7c6?w=800"
        },
        {
            "slug": "pinot-grigio",
            "name": "Pinot Grigio / Pinot Gris",
            "type": "weiss",
            "description": "Der Verwandlungskünstler – in Italien knackig-frisch wie ein Sommertag am Gardasee, im Elsass cremig-komplex wie ein herbstlicher Nebel über den Vogesen. Zitrone, grüne Birne, Mandel und weiße Blüten tanzen elegant im Glas. Unkompliziert und doch raffiniert, wie ein gut sitzender Leinenanzug an einem warmen Abend. Der perfekte Aperitivo-Wein, der aber auch zum Essen glänzt.",
            "description_en": "The transformation artist – in Italy crisp and fresh like a summer day at Lake Garda, in Alsace creamy and complex like autumn fog over the Vosges. Lemon, green pear, almond, and white blossoms dance elegantly in the glass. Uncomplicated yet refined, like a well-fitting linen suit on a warm evening. The perfect aperitivo wine that also shines with food.",
            "description_fr": "L'artiste de la transformation – en Italie frais et croquant comme un jour d'été au lac de Garde, en Alsace crémeux et complexe comme un brouillard d'automne sur les Vosges. Citron, poire verte, amande et fleurs blanches dansent élégamment dans le verre.",
            "synonyms": ["Grauburgunder", "Ruländer"],
            "body": "leicht bis mittel",
            "acidity": "mittel",
            "tannin": "niedrig",
            "aging": "Edelstahl, trocken",
            "primary_aromas": ["Zitrone", "Birne", "Apfel", "Mandel"],
            "tertiary_aromas": ["Honig", "Brioche", "Nuss"],
            "perfect_pairings": ["Carpaccio vom Lachs", "Risotto mit Meeresfrüchten", "Vitello Tonnato"],
            "perfect_pairings_en": ["Salmon carpaccio", "Seafood risotto", "Vitello Tonnato"],
            "perfect_pairings_fr": ["Carpaccio de saumon", "Risotto aux fruits de mer", "Vitello Tonnato"],
            "main_regions": ["Norditalien", "Elsass", "Oregon", "Deutschland"],
            "image_url": "https://images.unsplash.com/photo-1510812431401-41d2bd2722f3?w=800"
        },
        # RED WINES
        {
            "slug": "pinot-noir",
            "name": "Pinot Noir",
            "type": "rot",
            "description": "Die Diva unter den roten Reben – kapriziös, anspruchsvoll, aber in perfekter Form unvergleichlich. Burgunderrot wie ein Sonnenuntergang über den Côte d'Or, mit Aromen von frischen Kirschen, Erdbeeren, Rosenblättern und feuchtem Waldboden. Am Gaumen samtweich mit seidigen Tanninen, einer vibrierenden Säure und einem Finale, das Geschichten erzählt. Pinot Noir verlangt Hingabe – vom Winzer wie vom Genießer. Der Lohn: purer, unvergesslicher Trinkgenuss.",
            "description_en": "The diva among red grapes – capricious, demanding, but incomparable in perfect form. Burgundy red like a sunset over the Côte d'Or, with aromas of fresh cherries, strawberries, rose petals, and damp forest floor. On the palate: velvet soft with silky tannins, vibrant acidity, and a finish that tells stories. Pinot Noir demands devotion – from winemaker and connoisseur alike. The reward: pure, unforgettable drinking pleasure.",
            "description_fr": "La diva des cépages rouges – capricieuse, exigeante, mais incomparable dans sa forme parfaite. Rouge bourgogne comme un coucher de soleil sur la Côte d'Or, avec des arômes de cerises fraîches, fraises, pétales de rose et sous-bois humide. En bouche: doux comme du velours avec des tanins soyeux.",
            "synonyms": ["Spätburgunder", "Blauburgunder", "Pinot Nero"],
            "body": "leicht bis mittel",
            "acidity": "mittel bis hoch",
            "tannin": "niedrig bis mittel",
            "aging": "Holz, trocken",
            "primary_aromas": ["Kirsche", "Erdbeere", "Himbeere", "Rose"],
            "tertiary_aromas": ["Waldboden", "Pilze", "Leder", "Gewürze"],
            "perfect_pairings": ["Coq au Vin", "Ente mit Kirschsauce", "Lachs mit Pinot-Noir-Reduktion", "Brie de Meaux"],
            "perfect_pairings_en": ["Coq au Vin", "Duck with cherry sauce", "Salmon with Pinot Noir reduction", "Brie de Meaux"],
            "perfect_pairings_fr": ["Coq au Vin", "Canard sauce cerises", "Saumon à la réduction de Pinot Noir", "Brie de Meaux"],
            "main_regions": ["Burgund", "Oregon", "Neuseeland", "Deutschland"],
            "image_url": "https://images.unsplash.com/photo-1516594915697-87eb3b1c14ea?w=800"
        },
        {
            "slug": "cabernet-sauvignon",
            "name": "Cabernet Sauvignon",
            "type": "rot",
            "description": "Der König der roten Rebsorten – majestätisch, kraftvoll, unsterblich. Tiefes Rubinrot, fast undurchdringlich, wie das Versprechen auf etwas Großes. Schwarze Johannisbeere, Zedernholz, dunkle Schokolade und der unverwechselbare Duft von Bleistiftspitze. Am Gaumen strukturiert und muskulös, mit Tanninen wie Samt und Stahl zugleich. Cabernet braucht Zeit – wie alle großen Persönlichkeiten. Mit Reife offenbart er Tabak, Leder und eine fast meditative Tiefe.",
            "description_en": "The king of red grape varieties – majestic, powerful, immortal. Deep ruby red, almost impenetrable, like a promise of something great. Blackcurrant, cedarwood, dark chocolate, and the unmistakable scent of pencil shavings. On the palate: structured and muscular, with tannins like velvet and steel at once. Cabernet needs time – like all great personalities. With age, it reveals tobacco, leather, and an almost meditative depth.",
            "description_fr": "Le roi des cépages rouges – majestueux, puissant, immortel. Rouge rubis profond, presque impénétrable, comme la promesse de quelque chose de grand. Cassis, bois de cèdre, chocolat noir et le parfum incomparable de copeaux de crayon. En bouche: structuré et musclé, avec des tanins velours et acier à la fois.",
            "synonyms": ["Bouchet", "Petit Cabernet"],
            "body": "vollmundig",
            "acidity": "mittel bis hoch",
            "tannin": "hoch",
            "aging": "Holz (Barrique), trocken",
            "primary_aromas": ["Schwarze Johannisbeere", "Pflaume", "Kirsche", "Paprika"],
            "tertiary_aromas": ["Zedernholz", "Tabak", "Leder", "Schokolade", "Bleistift"],
            "perfect_pairings": ["T-Bone Steak vom Grill", "Lammkarree mit Rosmarin", "Entrecôte Café de Paris", "Gereifter Cheddar"],
            "perfect_pairings_en": ["Grilled T-bone steak", "Rack of lamb with rosemary", "Entrecôte Café de Paris", "Aged Cheddar"],
            "perfect_pairings_fr": ["T-bone steak grillé", "Carré d'agneau au romarin", "Entrecôte Café de Paris", "Cheddar affiné"],
            "main_regions": ["Bordeaux", "Napa Valley", "Chile", "Australien"],
            "image_url": "https://images.unsplash.com/photo-1553361371-9b22f78e8b1d?w=800"
        },
        {
            "slug": "merlot",
            "name": "Merlot",
            "type": "rot",
            "description": "Der sanfte Riese – zugänglich wie ein alter Freund, tiefgründig wie ein gutes Gespräch bei Kerzenlicht. Dunkle Pflaumen, reife Kirschen, Schokolade und ein Hauch von Kräutern malen ein Bild von Eleganz ohne Anstrengung. Am Gaumen geschmeidig und rund, mit weichen Tanninen, die wie eine warme Umarmung wirken. Merlot ist Balsam für die Seele – unkompliziert genug für jeden Tag, komplex genug für besondere Momente. Der Wein, der niemanden ausschließt.",
            "description_en": "The gentle giant – approachable like an old friend, profound like a good conversation by candlelight. Dark plums, ripe cherries, chocolate, and a hint of herbs paint a picture of effortless elegance. On the palate: supple and round, with soft tannins that feel like a warm embrace. Merlot is balm for the soul – uncomplicated enough for everyday, complex enough for special moments. The wine that excludes no one.",
            "description_fr": "Le gentil géant – accessible comme un vieil ami, profond comme une bonne conversation à la lueur des bougies. Prunes sombres, cerises mûres, chocolat et une touche d'herbes peignent une image d'élégance sans effort. En bouche: souple et rond, avec des tanins doux comme une étreinte chaleureuse.",
            "synonyms": ["Merlot Noir", "Vitraille"],
            "body": "mittel bis vollmundig",
            "acidity": "mittel",
            "tannin": "mittel",
            "aging": "Holz, trocken",
            "primary_aromas": ["Pflaume", "Kirsche", "Brombeere", "Veilchen"],
            "tertiary_aromas": ["Schokolade", "Kaffee", "Vanille", "Leder"],
            "perfect_pairings": ["Rinderbraten mit Rotweinjus", "Pilzrisotto mit Trüffel", "Hartkäse wie Pecorino", "Pasta Bolognese"],
            "perfect_pairings_en": ["Beef roast with red wine jus", "Mushroom risotto with truffle", "Hard cheese like Pecorino", "Pasta Bolognese"],
            "perfect_pairings_fr": ["Rôti de bœuf au jus de vin rouge", "Risotto aux champignons et truffe", "Fromage à pâte dure comme Pecorino", "Pâtes Bolognaise"],
            "main_regions": ["Bordeaux (Pomerol)", "Toskana", "Chile", "Kalifornien"],
            "image_url": "https://images.unsplash.com/photo-1547595628-c61a29f496f0?w=800"
        },
        {
            "slug": "syrah",
            "name": "Syrah / Shiraz",
            "type": "rot",
            "description": "Der Rebell – dunkel, geheimnisvoll und mit einer Intensität, die unter die Haut geht. Brombeere, Veilchen, schwarzer Pfeffer und rauchige Speckwürze vereinen sich zu einem Wein von dramatischer Schönheit. In der Rhône elegant und würzig, in Australien als Shiraz kraftvoll und üppig. Am Gaumen konzentriert mit festen Tanninen und einem Finale, das nach Rauch und Wildheit schmeckt. Für alle, die Wein wollen, der Geschichten von fernen Ländern erzählt.",
            "description_en": "The rebel – dark, mysterious, and with an intensity that gets under your skin. Blackberry, violet, black pepper, and smoky bacon spice unite in a wine of dramatic beauty. In the Rhône elegant and spicy, in Australia as Shiraz powerful and opulent. On the palate: concentrated with firm tannins and a finish that tastes of smoke and wilderness. For those who want wine that tells stories of distant lands.",
            "description_fr": "Le rebelle – sombre, mystérieux et avec une intensité qui prend aux tripes. Mûre, violette, poivre noir et épices fumées de lard s'unissent dans un vin d'une beauté dramatique. Dans le Rhône élégant et épicé, en Australie comme Shiraz puissant et opulent.",
            "synonyms": ["Shiraz", "Hermitage", "Sérine"],
            "body": "vollmundig",
            "acidity": "mittel",
            "tannin": "mittel bis hoch",
            "aging": "Holz (Barrique), trocken",
            "primary_aromas": ["Brombeere", "Schwarze Kirsche", "Pflaume", "Veilchen"],
            "tertiary_aromas": ["Schwarzer Pfeffer", "Speck", "Rauch", "Leder", "Schokolade"],
            "perfect_pairings": ["Gegrilltes Lamm mit Kräuterkruste", "Wild mit Brombeersauce", "BBQ Ribs", "Roquefort"],
            "perfect_pairings_en": ["Grilled lamb with herb crust", "Game with blackberry sauce", "BBQ ribs", "Roquefort"],
            "perfect_pairings_fr": ["Agneau grillé en croûte d'herbes", "Gibier sauce aux mûres", "Côtes de porc BBQ", "Roquefort"],
            "main_regions": ["Rhône", "Australien (Barossa)", "Kalifornien", "Chile"],
            "image_url": "https://images.unsplash.com/photo-1568213816046-0ee1c42bd559?w=800"
        },
        {
            "slug": "tempranillo",
            "name": "Tempranillo",
            "type": "rot",
            "description": "Die Seele Spaniens – stolz wie ein Flamenco-Tänzer, warm wie die kastilische Sonne. Kirsche, Leder, Tabak und getrocknete Feigen vereinen sich mit einer erdigen Würze, die nach spanischer Erde schmeckt. Am Gaumen elegant und mittelschwer, mit geschliffenen Tanninen und einer Balance, die Jahrzehnte überdauert. Von Rioja bis Ribera del Duero – Tempranillo ist der rote Faden, der durch Spaniens große Weingeschichte webt.",
            "description_en": "The soul of Spain – proud as a flamenco dancer, warm as the Castilian sun. Cherry, leather, tobacco, and dried figs unite with an earthy spice that tastes of Spanish soil. On the palate: elegant and medium-bodied, with polished tannins and a balance that lasts decades. From Rioja to Ribera del Duero – Tempranillo is the red thread woven through Spain's great wine history.",
            "description_fr": "L'âme de l'Espagne – fière comme un danseur de flamenco, chaude comme le soleil castillan. Cerise, cuir, tabac et figues séchées s'unissent à une épice terreuse qui a le goût de la terre espagnole. En bouche: élégant et moyennement corsé, avec des tanins polis et un équilibre qui dure des décennies.",
            "synonyms": ["Tinto Fino", "Tinta de Toro", "Cencibel", "Aragonez"],
            "body": "mittel bis vollmundig",
            "acidity": "mittel",
            "tannin": "mittel",
            "aging": "Holz (amerikanisch oder französisch), trocken",
            "primary_aromas": ["Kirsche", "Pflaume", "Tomate", "Feige"],
            "tertiary_aromas": ["Leder", "Tabak", "Vanille", "Kokos", "Dill"],
            "perfect_pairings": ["Tapas mit Jamón Ibérico", "Lamm-Eintopf mit Chorizo", "Gegrilltes Spanferkel", "Manchego"],
            "perfect_pairings_en": ["Tapas with Jamón Ibérico", "Lamb stew with Chorizo", "Grilled suckling pig", "Manchego"],
            "perfect_pairings_fr": ["Tapas au Jambon Ibérique", "Ragoût d'agneau au Chorizo", "Cochon de lait grillé", "Manchego"],
            "main_regions": ["Rioja", "Ribera del Duero", "Toro", "Portugal (Alentejo)"],
            "image_url": "https://images.unsplash.com/photo-1543163521-1bf539c55dd2?w=800"
        },
        {
            "slug": "sangiovese",
            "name": "Sangiovese",
            "type": "rot",
            "description": "Das Herz der Toskana – lebhaft wie ein italienischer Sonntag, rustikal wie eine Trattoria in den Hügeln von Chianti. Sauerkirsche, getrocknete Tomaten, Oregano und ein Hauch von Veilchen malen ein Bild von dolce vita. Am Gaumen saftig mit präsenter Säure und körnigen Tanninen, die nach Essen schreien. Sangiovese ist gemacht für den Tisch – für Pasta, Pizza, und lange Abende mit Freunden. Salute!",
            "description_en": "The heart of Tuscany – lively as an Italian Sunday, rustic as a trattoria in the Chianti hills. Sour cherry, dried tomatoes, oregano, and a hint of violet paint a picture of dolce vita. On the palate: juicy with present acidity and grainy tannins that cry out for food. Sangiovese is made for the table – for pasta, pizza, and long evenings with friends. Salute!",
            "description_fr": "Le cœur de la Toscane – vif comme un dimanche italien, rustique comme une trattoria dans les collines du Chianti. Griotte, tomates séchées, origan et une touche de violette peignent une image de dolce vita. En bouche: juteux avec une acidité présente et des tanins granuleux qui crient pour de la nourriture.",
            "synonyms": ["Brunello", "Prugnolo Gentile", "Morellino"],
            "body": "mittel bis vollmundig",
            "acidity": "hoch",
            "tannin": "mittel bis hoch",
            "aging": "Holz (großes oder kleines Fass), trocken",
            "primary_aromas": ["Sauerkirsche", "Erdbeere", "Pflaume", "Veilchen"],
            "tertiary_aromas": ["Tomate", "Leder", "Tabak", "Espresso", "Kräuter"],
            "perfect_pairings": ["Bistecca alla Fiorentina", "Pasta al Ragù", "Pizza Margherita", "Pecorino Toscano"],
            "perfect_pairings_en": ["Bistecca alla Fiorentina", "Pasta al Ragù", "Pizza Margherita", "Pecorino Toscano"],
            "perfect_pairings_fr": ["Bistecca alla Fiorentina", "Pâtes au Ragù", "Pizza Margherita", "Pecorino Toscano"],
            "main_regions": ["Chianti", "Brunello di Montalcino", "Vino Nobile di Montepulciano", "Romagna"],
            "image_url": "https://images.unsplash.com/photo-1506377247377-2a5b3b417ebb?w=800"
        },
        {
            "slug": "nebbiolo",
            "name": "Nebbiolo",
            "type": "rot",
            "description": "Der Aristokrat des Piemonts – trügerisch hell, aber mit einer Kraft, die Könige und Päpste in die Knie zwang. Ziegelrot wie die Dächer Albas, mit betörenden Aromen von Rosen, Teer, Kirschen und Gewürzen. Am Gaumen eine Explosion aus Säure und Tanninen – herausfordernd, fordernd, belohnend. Barolo und Barbaresco sind seine Kronen. Nebbiolo braucht Geduld: Mit 20 Jahren Reife offenbart er Trüffel, Herbstlaub und transzendente Tiefe.",
            "description_en": "The aristocrat of Piedmont – deceptively pale, but with a power that brought kings and popes to their knees. Brick red like the roofs of Alba, with intoxicating aromas of roses, tar, cherries, and spices. On the palate: an explosion of acidity and tannins – challenging, demanding, rewarding. Barolo and Barbaresco are its crowns. Nebbiolo needs patience: at 20 years of age, it reveals truffle, autumn leaves, and transcendent depth.",
            "description_fr": "L'aristocrate du Piémont – trompeusement pâle, mais avec une puissance qui a mis rois et papes à genoux. Rouge brique comme les toits d'Alba, avec des arômes enivrants de roses, goudron, cerises et épices. En bouche: une explosion d'acidité et de tanins – exigeant, défiant, gratifiant.",
            "synonyms": ["Spanna", "Chiavennasca", "Picotener"],
            "body": "vollmundig",
            "acidity": "hoch",
            "tannin": "hoch",
            "aging": "Holz (große Fässer), trocken",
            "primary_aromas": ["Rose", "Kirsche", "Himbeere", "Veilchen"],
            "tertiary_aromas": ["Teer", "Trüffel", "Leder", "Tabak", "Herbstlaub"],
            "perfect_pairings": ["Brasato al Barolo", "Tajarin mit weißen Trüffeln", "Wild-Ragout", "Gereifter Parmigiano"],
            "perfect_pairings_en": ["Brasato al Barolo", "Tajarin with white truffles", "Game ragout", "Aged Parmigiano"],
            "perfect_pairings_fr": ["Brasato al Barolo", "Tajarin aux truffes blanches", "Ragoût de gibier", "Parmigiano affiné"],
            "main_regions": ["Barolo", "Barbaresco", "Langhe", "Valtellina"],
            "image_url": "https://images.unsplash.com/photo-1474722883778-792e7990302f?w=800"
        },
        {
            "slug": "malbec",
            "name": "Malbec",
            "type": "rot",
            "description": "Der argentinische Traum – einst in Frankreich verschmäht, in den Anden zur Weltklasse gereift. Tiefviolett wie der Nachthimmel über Mendoza, mit üppigen Aromen von Brombeere, schwarzer Pflaume, Veilchen und süßen Gewürzen. Am Gaumen samtig und vollmundig, mit weichen Tanninen und einer saftigen Frucht, die nach mehr verlangt. Malbec ist der Wein für Steakliebhaber – geboren fürs Grillen unter freiem Himmel.",
            "description_en": "The Argentine dream – once scorned in France, matured to world class in the Andes. Deep violet like the night sky over Mendoza, with opulent aromas of blackberry, dark plum, violet, and sweet spices. On the palate: velvety and full-bodied, with soft tannins and a juicy fruit that demands more. Malbec is the wine for steak lovers – born for grilling under the open sky.",
            "description_fr": "Le rêve argentin – autrefois dédaigné en France, mûri vers l'excellence mondiale dans les Andes. Violet profond comme le ciel nocturne au-dessus de Mendoza, avec des arômes opulents de mûre, prune noire, violette et épices douces. En bouche: velouté et corsé, avec des tanins souples.",
            "synonyms": ["Côt", "Auxerrois", "Pressac"],
            "body": "vollmundig",
            "acidity": "mittel",
            "tannin": "mittel",
            "aging": "Holz, trocken",
            "primary_aromas": ["Brombeere", "Schwarze Pflaume", "Kirsche", "Veilchen"],
            "tertiary_aromas": ["Vanille", "Kakao", "Tabak", "Mokka", "Leder"],
            "perfect_pairings": ["Argentinisches Asado", "Ribeye Steak", "Empanadas", "Blauschimmelkäse"],
            "perfect_pairings_en": ["Argentine Asado", "Ribeye steak", "Empanadas", "Blue cheese"],
            "perfect_pairings_fr": ["Asado argentin", "Steak Ribeye", "Empanadas", "Fromage bleu"],
            "main_regions": ["Mendoza", "Cahors", "Chile", "Kalifornien"],
            "image_url": "https://images.unsplash.com/photo-1516594915697-87eb3b1c14ea?w=800"
        }
    ]
    
    # Clear existing and insert new
    await db.grape_varieties.delete_many({})
    
    for grape_data in grapes:
        grape = GrapeVariety(**grape_data)
        doc = grape.model_dump()
        doc['created_at'] = doc['created_at'].isoformat()
        await db.grape_varieties.insert_one(doc)
    
    return {"message": f"{len(grapes)} Rebsorten wurden erstellt"}


# ===================== WINE DATABASE ENDPOINTS (DUPLICATE - DISABLED) =====================

# DISABLED DUPLICATE ENDPOINT
# @api_router.get("/wine-database", response_model=List[WineDatabaseEntry])
async def get_wine_database_DISABLED(
    search: Optional[str] = None,
    country: Optional[str] = None,
    region: Optional[str] = None,
    appellation: Optional[str] = None,
    grape_variety: Optional[str] = None,
    wine_color: Optional[str] = None,
    price_category: Optional[str] = None,
    skip: int = 0,
    limit: int = 50
):
    """Get wines from the database with filters and weighted full-text search"""
    
    # If search is provided, use weighted search with prioritization
    if search:
        search_term = search.strip()
        
        # Normalize search term to handle accents (é -> e, à -> a, etc.)
        # This allows "Chateau" to match "Château"
        import unicodedata
        normalized_search = ''.join(
            c for c in unicodedata.normalize('NFD', search_term)
            if unicodedata.category(c) != 'Mn'
        )
        
        # Create regex pattern that matches both with and without accents
        # e.g., "chateau" matches both "chateau" and "château"
        def create_accent_insensitive_pattern(term):
            # Map common accent variations
            replacements = {
                'e': '[eéèêë]',
                'a': '[aàâä]',
                'i': '[iîï]',
                'o': '[oôö]',
                'u': '[uùûü]',
                'c': '[cç]',
                'n': '[nñ]'
            }
            pattern = ''
            for char in term.lower():
                pattern += replacements.get(char, char)
            return pattern
        
        accent_pattern = create_accent_insensitive_pattern(normalized_search)
        
        # Prioritized search: name > appellation > region > country
        # First, find exact or partial matches in name (highest priority)
        name_query = {
            "$or": [
                {"name": {"$regex": f"^{accent_pattern}", "$options": "i"}},  # Starts with
                {"name": {"$regex": accent_pattern, "$options": "i"}}  # Contains
            ]
        }
        
        # Then appellation
        appellation_query = {"appellation": {"$regex": accent_pattern, "$options": "i"}}
        
        # Then region
        region_query = {"region": {"$regex": accent_pattern, "$options": "i"}}
        
        # Then country
        country_query = {"country": {"$regex": accent_pattern, "$options": "i"}}
        
        # Also search in winery and grape variety
        winery_query = {"winery": {"$regex": accent_pattern, "$options": "i"}}
        grape_query = {"grape_variety": {"$regex": accent_pattern, "$options": "i"}}
        
        # Combine with additional filters
        filter_conditions = []
        if country:
            filter_conditions.append({"country": country})
        if region:
            filter_conditions.append({"region": region})
        if appellation:
            filter_conditions.append({"appellation": appellation})
        if grape_variety:
            filter_conditions.append({"grape_variety": grape_variety})
        if wine_color:
            filter_conditions.append({"wine_color": wine_color})
        if price_category:
            filter_conditions.append({"price_category": price_category})
        
        # Fetch results with priority order
        wines = []
        seen_ids = set()
        
        # 1. Name matches (highest priority)
        if len(wines) < limit:
            query = name_query.copy()
            if filter_conditions:
                query = {"$and": [name_query, {"$and": filter_conditions}]}
            name_wines = await db.wine_database.find(query, {"_id": 0}).limit(limit).to_list(limit)
            for w in name_wines:
                if w['id'] not in seen_ids:
                    wines.append(w)
                    seen_ids.add(w['id'])
        
        # 2. Appellation matches
        if len(wines) < limit:
            query = appellation_query.copy()
            if filter_conditions:
                query = {"$and": [appellation_query, {"$and": filter_conditions}]}
            app_wines = await db.wine_database.find(query, {"_id": 0}).limit(limit - len(wines)).to_list(limit - len(wines))
            for w in app_wines:
                if w['id'] not in seen_ids:
                    wines.append(w)
                    seen_ids.add(w['id'])
        
        # 3. Region matches
        if len(wines) < limit:
            query = region_query.copy()
            if filter_conditions:
                query = {"$and": [region_query, {"$and": filter_conditions}]}
            region_wines = await db.wine_database.find(query, {"_id": 0}).limit(limit - len(wines)).to_list(limit - len(wines))
            for w in region_wines:
                if w['id'] not in seen_ids:
                    wines.append(w)
                    seen_ids.add(w['id'])
        
        # 4. Country matches
        if len(wines) < limit:
            query = country_query.copy()
            if filter_conditions:
                query = {"$and": [country_query, {"$and": filter_conditions}]}
            country_wines = await db.wine_database.find(query, {"_id": 0}).limit(limit - len(wines)).to_list(limit - len(wines))
            for w in country_wines:
                if w['id'] not in seen_ids:
                    wines.append(w)
                    seen_ids.add(w['id'])
        
        # 5. Winery and grape matches (lowest priority)
        if len(wines) < limit:
            query = {"$or": [winery_query, grape_query]}
            if filter_conditions:
                query = {"$and": [query, {"$and": filter_conditions}]}
            other_wines = await db.wine_database.find(query, {"_id": 0}).limit(limit - len(wines)).to_list(limit - len(wines))
            for w in other_wines:
                if w['id'] not in seen_ids:
                    wines.append(w)
                    seen_ids.add(w['id'])
        
        # Apply skip/limit
        wines = wines[skip:skip + limit]
        
    else:
        # No search term - use simple filter query
        query = {}
        if country:
            query["country"] = country
        if region:
            query["region"] = region
        if appellation:
            query["appellation"] = appellation
        if grape_variety:
            query["grape_variety"] = grape_variety
        if wine_color:
            query["wine_color"] = wine_color
        if price_category:
            query["price_category"] = price_category
        
        wines = await db.wine_database.find(query, {"_id": 0}).skip(skip).limit(limit).to_list(limit)
    
    # Convert datetime strings
    for wine in wines:
        if isinstance(wine.get('created_at'), str):
            wine['created_at'] = datetime.fromisoformat(wine['created_at'])
    
    return wines

@api_router.get("/wine-database/{wine_id}", response_model=WineDatabaseEntry)
async def get_wine_detail(wine_id: str):
    """Get detailed information about a specific wine"""
    wine = await db.wine_database.find_one({"id": wine_id}, {"_id": 0})
    if not wine:
        raise HTTPException(status_code=404, detail="Wein nicht gefunden")
    
    if isinstance(wine.get('created_at'), str):
        wine['created_at'] = datetime.fromisoformat(wine['created_at'])
    
    return wine

@api_router.get("/wine-database/autocomplete/{query}")
async def autocomplete_wines(query: str, limit: int = 10):
    """Autocomplete for wine search - mit Akzent-Unterstützung"""
    # WICHTIG: Akzent-insensitive Suche verwenden!
    accent_pattern = create_accent_insensitive_pattern(query)
    search_query = {"$or": [
        {"name": {"$regex": accent_pattern, "$options": "i"}},
        {"winery": {"$regex": accent_pattern, "$options": "i"}},
        {"grape_variety": {"$regex": accent_pattern, "$options": "i"}}
    ]}
    
    wines = await db.wine_database.find(search_query, {"_id": 0, "id": 1, "name": 1, "winery": 1, "wine_color": 1}).limit(limit).to_list(limit)
    
    return wines

@api_router.get("/wine-database-filters")
async def get_wine_filters():
    """Get all available filter options"""
    countries = await db.wine_database.distinct("country")
    regions = await db.wine_database.distinct("region")
    appellations = await db.wine_database.distinct("appellation")
    grape_varieties = await db.wine_database.distinct("grape_variety")
    wine_colors = await db.wine_database.distinct("wine_color")
    price_categories = await db.wine_database.distinct("price_category")
    
    return {
        "countries": sorted([c for c in countries if c]),
        "regions": sorted([r for r in regions if r]),
        "appellations": sorted([a for a in appellations if a]),
        "grape_varieties": sorted([g for g in grape_varieties if g]),
        "wine_colors": sorted([w for w in wine_colors if w]),
        "price_categories": sorted([p for p in price_categories if p])
    }

@api_router.post("/seed-wine-database")
async def seed_wine_database(count: int = 2000):
    """Seed the wine database with a mix of real and AI-generated wines"""
    existing_count = await db.wine_database.count_documents({})
    if existing_count > 0:
        return {"message": f"Datenbank enthält bereits {existing_count} Weine"}
    
    logger.info(f"Starting to seed wine database with {count} wines...")
    
    # Base set of real famous wines
    base_wines = [
        # France - Bordeaux
        {"name": "Château Margaux", "winery": "Château Margaux", "country": "Frankreich", "region": "Bordeaux", "appellation": "Margaux", "grape_variety": "Cabernet Sauvignon", "wine_color": "rot", "year": 2015, "price_category": "luxury"},
        {"name": "Château Lafite Rothschild", "winery": "Château Lafite Rothschild", "country": "Frankreich", "region": "Bordeaux", "appellation": "Pauillac", "grape_variety": "Cabernet Sauvignon", "wine_color": "rot", "year": 2016, "price_category": "luxury"},
        {"name": "Château Latour", "winery": "Château Latour", "country": "Frankreich", "region": "Bordeaux", "appellation": "Pauillac", "grape_variety": "Cabernet Sauvignon", "wine_color": "rot", "year": 2014, "price_category": "luxury"},
        {"name": "Château Haut-Brion", "winery": "Château Haut-Brion", "country": "Frankreich", "region": "Bordeaux", "appellation": "Pessac-Léognan", "grape_variety": "Merlot", "wine_color": "rot", "year": 2015, "price_category": "luxury"},
        {"name": "Château Mouton Rothschild", "winery": "Château Mouton Rothschild", "country": "Frankreich", "region": "Bordeaux", "appellation": "Pauillac", "grape_variety": "Cabernet Sauvignon", "wine_color": "rot", "year": 2016, "price_category": "luxury"},
        
        # France - Burgundy
        {"name": "Romanée-Conti", "winery": "Domaine de la Romanée-Conti", "country": "Frankreich", "region": "Burgund", "appellation": "Vosne-Romanée", "grape_variety": "Pinot Noir", "wine_color": "rot", "year": 2018, "price_category": "luxury"},
        {"name": "La Tâche", "winery": "Domaine de la Romanée-Conti", "country": "Frankreich", "region": "Burgund", "appellation": "Vosne-Romanée", "grape_variety": "Pinot Noir", "wine_color": "rot", "year": 2017, "price_category": "luxury"},
        {"name": "Montrachet Grand Cru", "winery": "Domaine de la Romanée-Conti", "country": "Frankreich", "region": "Burgund", "appellation": "Montrachet", "grape_variety": "Chardonnay", "wine_color": "weiss", "year": 2019, "price_category": "luxury"},
        {"name": "Chablis Grand Cru", "winery": "William Fèvre", "country": "Frankreich", "region": "Burgund", "appellation": "Chablis", "grape_variety": "Chardonnay", "wine_color": "weiss", "year": 2020, "price_category": "premium"},
        
        # France - Champagne
        {"name": "Dom Pérignon", "winery": "Moët & Chandon", "country": "Frankreich", "region": "Champagne", "appellation": "Champagne", "grape_variety": "Chardonnay", "wine_color": "schaumwein", "year": 2012, "price_category": "luxury"},
        {"name": "Krug Grande Cuvée", "winery": "Krug", "country": "Frankreich", "region": "Champagne", "appellation": "Champagne", "grape_variety": "Pinot Noir", "wine_color": "schaumwein", "price_category": "luxury"},
        {"name": "Cristal", "winery": "Louis Roederer", "country": "Frankreich", "region": "Champagne", "appellation": "Champagne", "grape_variety": "Chardonnay", "wine_color": "schaumwein", "year": 2013, "price_category": "luxury"},
        
        # Italy - Tuscany
        {"name": "Sassicaia", "winery": "Tenuta San Guido", "country": "Italien", "region": "Toskana", "appellation": "Bolgheri", "grape_variety": "Cabernet Sauvignon", "wine_color": "rot", "year": 2017, "price_category": "premium"},
        {"name": "Tignanello", "winery": "Antinori", "country": "Italien", "region": "Toskana", "appellation": "Toscana IGT", "grape_variety": "Sangiovese", "wine_color": "rot", "year": 2018, "price_category": "premium"},
        {"name": "Brunello di Montalcino", "winery": "Biondi-Santi", "country": "Italien", "region": "Toskana", "appellation": "Montalcino", "grape_variety": "Sangiovese", "wine_color": "rot", "year": 2016, "price_category": "premium"},
        {"name": "Chianti Classico Riserva", "winery": "Castello di Ama", "country": "Italien", "region": "Toskana", "appellation": "Chianti Classico", "grape_variety": "Sangiovese", "wine_color": "rot", "year": 2019, "price_category": "mid-range"},
        
        # Italy - Piedmont
        {"name": "Barolo", "winery": "Giacomo Conterno", "country": "Italien", "region": "Piemont", "appellation": "Barolo", "grape_variety": "Nebbiolo", "wine_color": "rot", "year": 2016, "price_category": "premium"},
        {"name": "Barbaresco", "winery": "Gaja", "country": "Italien", "region": "Piemont", "appellation": "Barbaresco", "grape_variety": "Nebbiolo", "wine_color": "rot", "year": 2018, "price_category": "premium"},
        {"name": "Gavi di Gavi", "winery": "La Scolca", "country": "Italien", "region": "Piemont", "appellation": "Gavi", "grape_variety": "Cortese", "wine_color": "weiss", "year": 2021, "price_category": "mid-range"},
        
        # Spain
        {"name": "Vega Sicilia Único", "winery": "Vega Sicilia", "country": "Spanien", "region": "Ribera del Duero", "appellation": "Ribera del Duero", "grape_variety": "Tempranillo", "wine_color": "rot", "year": 2010, "price_category": "luxury"},
        {"name": "Rioja Gran Reserva", "winery": "Marqués de Riscal", "country": "Spanien", "region": "Rioja", "appellation": "Rioja", "grape_variety": "Tempranillo", "wine_color": "rot", "year": 2015, "price_category": "premium"},
        {"name": "Priorat", "winery": "Clos Mogador", "country": "Spanien", "region": "Priorat", "appellation": "Priorat", "grape_variety": "Garnacha", "wine_color": "rot", "year": 2017, "price_category": "premium"},
        {"name": "Albariño", "winery": "Pazo de Señorans", "country": "Spanien", "region": "Rías Baixas", "appellation": "Rías Baixas", "grape_variety": "Albariño", "wine_color": "weiss", "year": 2021, "price_category": "mid-range"},
        
        # Germany
        {"name": "Riesling Kabinett", "winery": "Weingut Dr. Loosen", "country": "Deutschland", "region": "Mosel", "appellation": "Mosel", "grape_variety": "Riesling", "wine_color": "weiss", "year": 2020, "price_category": "mid-range"},
        {"name": "Riesling Spätlese", "winery": "Egon Müller", "country": "Deutschland", "region": "Mosel", "appellation": "Saar", "grape_variety": "Riesling", "wine_color": "weiss", "year": 2019, "price_category": "premium"},
        {"name": "Riesling Auslese", "winery": "J.J. Prüm", "country": "Deutschland", "region": "Mosel", "appellation": "Mosel", "grape_variety": "Riesling", "wine_color": "suesswein", "year": 2018, "price_category": "premium"},
        {"name": "Spätburgunder", "winery": "Weingut Friedrich Becker", "country": "Deutschland", "region": "Pfalz", "appellation": "Pfalz", "grape_variety": "Pinot Noir", "wine_color": "rot", "year": 2019, "price_category": "premium"},
        
        # Austria
        {"name": "Grüner Veltliner", "winery": "Weingut FX Pichler", "country": "Österreich", "region": "Wachau", "appellation": "Wachau", "grape_variety": "Grüner Veltliner", "wine_color": "weiss", "year": 2020, "price_category": "premium"},
        {"name": "Riesling Smaragd", "winery": "Domäne Wachau", "country": "Österreich", "region": "Wachau", "appellation": "Wachau", "grape_variety": "Riesling", "wine_color": "weiss", "year": 2019, "price_category": "premium"},
        
        # USA - California
        {"name": "Opus One", "winery": "Opus One Winery", "country": "USA", "region": "Kalifornien", "appellation": "Napa Valley", "grape_variety": "Cabernet Sauvignon", "wine_color": "rot", "year": 2016, "price_category": "luxury"},
        {"name": "Screaming Eagle", "winery": "Screaming Eagle Winery", "country": "USA", "region": "Kalifornien", "appellation": "Napa Valley", "grape_variety": "Cabernet Sauvignon", "wine_color": "rot", "year": 2015, "price_category": "luxury"},
        {"name": "Caymus Special Selection", "winery": "Caymus Vineyards", "country": "USA", "region": "Kalifornien", "appellation": "Napa Valley", "grape_variety": "Cabernet Sauvignon", "wine_color": "rot", "year": 2018, "price_category": "premium"},
        {"name": "Stag's Leap Wine Cellars", "winery": "Stag's Leap Wine Cellars", "country": "USA", "region": "Kalifornien", "appellation": "Napa Valley", "grape_variety": "Cabernet Sauvignon", "wine_color": "rot", "year": 2017, "price_category": "premium"},
        {"name": "Ridge Monte Bello", "winery": "Ridge Vineyards", "country": "USA", "region": "Kalifornien", "appellation": "Santa Cruz Mountains", "grape_variety": "Cabernet Sauvignon", "wine_color": "rot", "year": 2016, "price_category": "luxury"},
        {"name": "Kistler Chardonnay", "winery": "Kistler Vineyards", "country": "USA", "region": "Kalifornien", "appellation": "Sonoma Coast", "grape_variety": "Chardonnay", "wine_color": "weiss", "year": 2019, "price_category": "premium"},
        
        # USA - Oregon
        {"name": "Domaine Drouhin Pinot Noir", "winery": "Domaine Drouhin", "country": "USA", "region": "Oregon", "appellation": "Willamette Valley", "grape_variety": "Pinot Noir", "wine_color": "rot", "year": 2018, "price_category": "premium"},
        
        # Australia
        {"name": "Penfolds Grange", "winery": "Penfolds", "country": "Australien", "region": "South Australia", "appellation": "Barossa Valley", "grape_variety": "Shiraz", "wine_color": "rot", "year": 2016, "price_category": "luxury"},
        {"name": "Henschke Hill of Grace", "winery": "Henschke", "country": "Australien", "region": "South Australia", "appellation": "Eden Valley", "grape_variety": "Shiraz", "wine_color": "rot", "year": 2015, "price_category": "luxury"},
        
        # New Zealand
        {"name": "Cloudy Bay Sauvignon Blanc", "winery": "Cloudy Bay", "country": "Neuseeland", "region": "Marlborough", "appellation": "Marlborough", "grape_variety": "Sauvignon Blanc", "wine_color": "weiss", "year": 2021, "price_category": "mid-range"},
        
        # Argentina
        {"name": "Catena Zapata Malbec", "winery": "Catena Zapata", "country": "Argentinien", "region": "Mendoza", "appellation": "Mendoza", "grape_variety": "Malbec", "wine_color": "rot", "year": 2018, "price_category": "premium"},
        
        # Chile
        {"name": "Concha y Toro Don Melchor", "winery": "Concha y Toro", "country": "Chile", "region": "Maipo Valley", "appellation": "Puente Alto", "grape_variety": "Cabernet Sauvignon", "wine_color": "rot", "year": 2017, "price_category": "premium"},
        
        # South Africa
        {"name": "Kanonkop Paul Sauer", "winery": "Kanonkop", "country": "Südafrika", "region": "Stellenbosch", "appellation": "Stellenbosch", "grape_variety": "Cabernet Sauvignon", "wine_color": "rot", "year": 2017, "price_category": "premium"},
        
        # Portugal
        {"name": "Quinta do Noval Vintage Port", "winery": "Quinta do Noval", "country": "Portugal", "region": "Douro", "appellation": "Porto", "grape_variety": "Touriga Nacional", "wine_color": "suesswein", "year": 2016, "price_category": "luxury"},
        {"name": "Vinho Verde", "winery": "Quinta da Aveleda", "country": "Portugal", "region": "Minho", "appellation": "Vinho Verde", "grape_variety": "Alvarinho", "wine_color": "weiss", "year": 2021, "price_category": "budget"},
    ]
    
    # Add emotional descriptions and pairings to base wines
    wines_to_insert = []
    
    for base_wine in base_wines:
        # Generate emotional description with GPT-5.1 (I'll create a simpler version for speed)
        wine_entry = WineDatabaseEntry(
            **base_wine,
            description=f"Ein außergewöhnlicher Wein aus {base_wine['region']}, der die Essenz von {base_wine['grape_variety']} perfekt einfängt.",
            tasting_notes="Aromen von dunklen Früchten, elegant und komplex",
            food_pairings=["Gegrilltes Fleisch", "Käse", "Wildgerichte"],
            alcohol_content=13.5,
            image_url="/placeholder-wine.png",
            rating=4.5
        )
        wines_to_insert.append(wine_entry.model_dump())
    
    # Insert base wines
    for wine_data in wines_to_insert:
        wine_data['created_at'] = wine_data['created_at'].isoformat()
        await db.wine_database.insert_one(wine_data)
    
    inserted_count = len(wines_to_insert)
    logger.info(f"Inserted {inserted_count} base wines")
    
    # Generate additional wines to reach target count
    # For now, return the base wines count
    return {"message": f"{inserted_count} Weine wurden erstellt (Basis-Set). Weitere {count - inserted_count} werden nach und nach generiert."}


@api_router.post("/generate-wines")
async def generate_additional_wines(batch_size: int = 50):
    """Generate additional wines using AI"""
    try:
        current_count = await db.wine_database.count_documents({})
        logger.info(f"Current wine count: {current_count}, generating {batch_size} more...")
        
        # Wine generation templates for variety
        regions_templates = [
            {"country": "Frankreich", "region": "Bordeaux", "grapes": ["Cabernet Sauvignon", "Merlot", "Cabernet Franc"], "color": "rot"},
            {"country": "Frankreich", "region": "Burgund", "grapes": ["Pinot Noir", "Chardonnay"], "color": "rot"},
            {"country": "Frankreich", "region": "Rhône", "grapes": ["Syrah", "Grenache", "Mourvèdre"], "color": "rot"},
            {"country": "Frankreich", "region": "Loire", "grapes": ["Sauvignon Blanc", "Chenin Blanc"], "color": "weiss"},
            {"country": "Frankreich", "region": "Elsass", "grapes": ["Riesling", "Gewürztraminer", "Pinot Gris"], "color": "weiss"},
            {"country": "Italien", "region": "Toskana", "grapes": ["Sangiovese", "Cabernet Sauvignon"], "color": "rot"},
            {"country": "Italien", "region": "Piemont", "grapes": ["Nebbiolo", "Barbera", "Dolcetto"], "color": "rot"},
            {"country": "Italien", "region": "Venetien", "grapes": ["Corvina", "Garganega"], "color": "rot"},
            {"country": "Spanien", "region": "Rioja", "grapes": ["Tempranillo", "Garnacha"], "color": "rot"},
            {"country": "Spanien", "region": "Ribera del Duero", "grapes": ["Tempranillo"], "color": "rot"},
            {"country": "Spanien", "region": "Priorat", "grapes": ["Garnacha", "Cariñena"], "color": "rot"},
            {"country": "Deutschland", "region": "Mosel", "grapes": ["Riesling"], "color": "weiss"},
            {"country": "Deutschland", "region": "Rheingau", "grapes": ["Riesling"], "color": "weiss"},
            {"country": "Deutschland", "region": "Pfalz", "grapes": ["Riesling", "Pinot Noir"], "color": "weiss"},
            {"country": "Österreich", "region": "Wachau", "grapes": ["Grüner Veltliner", "Riesling"], "color": "weiss"},
            {"country": "USA", "region": "Kalifornien", "grapes": ["Cabernet Sauvignon", "Chardonnay", "Pinot Noir"], "color": "rot"},
            {"country": "USA", "region": "Oregon", "grapes": ["Pinot Noir"], "color": "rot"},
            {"country": "Australien", "region": "Barossa Valley", "grapes": ["Shiraz"], "color": "rot"},
            {"country": "Neuseeland", "region": "Marlborough", "grapes": ["Sauvignon Blanc"], "color": "weiss"},
            {"country": "Argentinien", "region": "Mendoza", "grapes": ["Malbec"], "color": "rot"},
            {"country": "Chile", "region": "Maipo Valley", "grapes": ["Cabernet Sauvignon", "Carmenère"], "color": "rot"},
            {"country": "Südafrika", "region": "Stellenbosch", "grapes": ["Pinotage", "Cabernet Sauvignon"], "color": "rot"},
            {"country": "Portugal", "region": "Douro", "grapes": ["Touriga Nacional", "Tinta Roriz"], "color": "rot"},
        ]
        
        price_categories = ["budget", "mid-range", "premium", "luxury"]
        
        wines_generated = []
        
        # Generate wines in batches
        for i in range(0, batch_size, 5):
            # Select random region template
            import random
            template = random.choice(regions_templates)
            grape = random.choice(template["grapes"])
            price_cat = random.choice(price_categories)
            year = random.randint(2015, 2022)
            
            # Use GPT-5.1 to generate wine details
            chat = LlmChat(
                api_key=EMERGENT_LLM_KEY,
                session_id=str(uuid.uuid4()),
                system_message="Du bist ein Weinexperte. Generiere realistische Wein-Informationen im JSON-Format."
            ).with_model("openai", "gpt-5.1")
            
            prompt = f"""Generiere 5 realistische Weine aus {template['region']}, {template['country']} mit folgenden Eigenschaften:
- Rebsorte: {grape}
- Weinfarbe: {template['color']}
- Preiskategorie: {price_cat}
- Jahrgang: {year}

Für jeden Wein generiere:
1. Einen authentischen Weinnamen (z.B. "Château...", "Domaine...", "Estate...")
2. Einen realistischen Weingut-Namen
3. Eine emotionale, poetische Beschreibung (2-3 Sätze auf Deutsch)
4. 4-6 Food Pairings
5. Appellations-Name (wenn zutreffend)

Antwort NUR als JSON-Array:
[
  {{
    "name": "Weinname",
    "winery": "Weingut Name",
    "appellation": "Appellation Name",
    "description": "Emotionale Beschreibung...",
    "food_pairings": ["Pairing 1", "Pairing 2", ...]
  }},
  ...
]"""
            
            user_message = UserMessage(text=prompt)
            response = await chat.send_message(user_message)
            
            # Parse JSON response
            try:
                json_match = re.search(r'\[.*\]', response, re.DOTALL)
                if json_match:
                    wines_data = json.loads(json_match.group())
                    
                    for wine_data in wines_data[:5]:  # Ensure max 5
                        wine_entry = WineDatabaseEntry(
                            name=wine_data.get('name', f'Wein {current_count + len(wines_generated) + 1}'),
                            winery=wine_data.get('winery', 'Unbekannt'),
                            country=template['country'],
                            region=template['region'],
                            appellation=wine_data.get('appellation'),
                            grape_variety=grape,
                            wine_color=template['color'],
                            year=year,
                            description=wine_data.get('description', 'Ein bemerkenswerter Wein.'),
                            food_pairings=wine_data.get('food_pairings', []),
                            price_category=price_cat,
                            alcohol_content=round(random.uniform(11.5, 15.0), 1),
                            image_url="/placeholder-wine.png",
                            rating=round(random.uniform(3.5, 5.0), 1)
                        )
                        wines_generated.append(wine_entry.model_dump())
            except Exception as e:
                logger.warning(f"Failed to parse AI wine generation: {e}")
                continue
        
        # Insert generated wines
        if wines_generated:
            for wine_data in wines_generated:
                wine_data['created_at'] = wine_data['created_at'].isoformat()
                await db.wine_database.insert_one(wine_data)
        
        new_count = await db.wine_database.count_documents({})
        logger.info(f"Generated {len(wines_generated)} wines. Total count: {new_count}")
        
        return {"message": f"{len(wines_generated)} Weine generiert. Gesamt: {new_count}"}
        
    except Exception as e:
        logger.error(f"Wine generation error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Fehler bei der Wein-Generierung: {str(e)}")



# ===================== WINE FAVORITES ENDPOINTS =====================

class FavoriteWine(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    wine_id: str
    wine_name: str
    winery: str
    wine_color: str
    country: str
    region: str
    image_url: Optional[str] = None
    is_wishlist: bool = False  # False = favorite, True = wishlist
    notes: Optional[str] = None
    added_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

@api_router.get("/favorites")
async def get_favorites(wishlist_only: bool = False):
    """Get user's favorite wines or wishlist"""
    query = {}
    if wishlist_only:
        query["is_wishlist"] = True
    else:
        query["is_wishlist"] = False
    
    favorites = await db.wine_favorites.find(query, {"_id": 0}).sort("added_at", -1).to_list(1000)
    
    for fav in favorites:
        if isinstance(fav.get('added_at'), str):
            fav['added_at'] = datetime.fromisoformat(fav['added_at'])
    
    return favorites

@api_router.post("/favorites/{wine_id}")
async def add_to_favorites(wine_id: str, is_wishlist: bool = False):
    """Add a wine to favorites or wishlist"""
    # Get wine details from database
    wine = await db.wine_database.find_one({"id": wine_id}, {"_id": 0})
    if not wine:
        raise HTTPException(status_code=404, detail="Wein nicht gefunden")
    
    # Check if already in favorites/wishlist
    existing = await db.wine_favorites.find_one({"wine_id": wine_id})
    if existing:
        # Update is_wishlist status if different
        if existing.get('is_wishlist') != is_wishlist:
            await db.wine_favorites.update_one(
                {"wine_id": wine_id},
                {"$set": {"is_wishlist": is_wishlist}}
            )
            return {"message": f"Wein zu {'Merkliste' if is_wishlist else 'Favoriten'} verschoben"}
        raise HTTPException(status_code=400, detail="Wein bereits in der Liste")
    
    # Create favorite entry
    favorite = FavoriteWine(
        wine_id=wine_id,
        wine_name=wine['name'],
        winery=wine['winery'],
        wine_color=wine['wine_color'],
        country=wine['country'],
        region=wine['region'],
        image_url=wine.get('image_url'),
        is_wishlist=is_wishlist
    )
    
    fav_dict = favorite.model_dump()
    fav_dict['added_at'] = fav_dict['added_at'].isoformat()
    await db.wine_favorites.insert_one(fav_dict)
    
    return {"message": f"Wein zu {'Merkliste' if is_wishlist else 'Favoriten'} hinzugefügt"}

@api_router.delete("/favorites/{wine_id}")
async def remove_from_favorites(wine_id: str):
    """Remove a wine from favorites/wishlist"""
    result = await db.wine_favorites.delete_one({"wine_id": wine_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Wein nicht in Favoriten")
    
    return {"message": "Wein aus Favoriten entfernt"}

@api_router.get("/favorites/check/{wine_id}")
async def check_favorite_status(wine_id: str):
    """Check if a wine is in favorites or wishlist"""
    favorite = await db.wine_favorites.find_one({"wine_id": wine_id}, {"_id": 0})
    if favorite:
        return {
            "is_favorite": not favorite.get('is_wishlist', False),
            "is_wishlist": favorite.get('is_wishlist', False)
        }
    return {"is_favorite": False, "is_wishlist": False}

# ===================== WEEKLY TIP ENDPOINTS =====================

@api_router.get("/weekly-tips")
async def get_weekly_tips(limit: int = 4, include_archived: bool = False):
    """
    Hole die neuesten Wochen-Tipps.
    - limit: Anzahl der Tipps (Standard: 4)
    - include_archived: Auch inaktive Tipps laden (für Archiv)
    """
    query = {} if include_archived else {"is_active": True}
    tips = await db.weekly_tips.find(query, {"_id": 0}).sort("created_at", -1).to_list(limit)
    return tips


@api_router.get("/weekly-tips/archive")
async def get_weekly_tips_archive(page: int = 1, per_page: int = 12):
    """
    Hole alle Tipps für das Archiv (paginiert).
    """
    skip = (page - 1) * per_page
    total = await db.weekly_tips.count_documents({})
    tips = await db.weekly_tips.find({}, {"_id": 0}).sort("created_at", -1).skip(skip).limit(per_page).to_list(per_page)
    
    return {
        "tips": tips,
        "total": total,
        "page": page,
        "per_page": per_page,
        "total_pages": (total + per_page - 1) // per_page
    }


@api_router.get("/weekly-tips/{tip_id}")
async def get_weekly_tip(tip_id: str):
    """Hole einen einzelnen Tipp nach ID."""
    tip = await db.weekly_tips.find_one({"id": tip_id}, {"_id": 0})
    if not tip:
        raise HTTPException(status_code=404, detail="Tipp nicht gefunden")
    return tip


@api_router.post("/admin/generate-weekly-tip")
async def generate_weekly_tip(language: str = "de"):
    """
    Generiert einen neuen Wochen-Tipp mit KI.
    Wird vom Cron-Job oder manuell aufgerufen.
    """
    from datetime import date
    import calendar
    
    # Aktuelle Kalenderwoche ermitteln
    today = date.today()
    week_number = today.isocalendar()[1]
    year = today.year
    
    # Prüfen ob diese Woche schon ein Tipp existiert
    existing = await db.weekly_tips.find_one({
        "week_number": week_number,
        "year": year,
        "language": language
    })
    
    if existing:
        return {
            "status": "exists",
            "message": f"Tipp für KW {week_number}/{year} existiert bereits",
            "tip": {k: v for k, v in existing.items() if k != "_id"}
        }
    
    # KI-Prompt für kreative Tipp-Generierung
    prompt = f"""Du bist ein kreativer Sommelier. Generiere einen überraschenden, aber gut funktionierenden Wein-Pairing-Tipp der Woche.

WICHTIG: Wähle eine UNGEWÖHNLICHE aber FUNKTIONIERENDE Kombination, die ein "Aha-Erlebnis" auslöst.

Antworte NUR im folgenden JSON-Format (keine Markdown-Codeblocks):
{{
    "dish": "Name des Gerichts (kurz)",
    "dish_emoji": "passendes Emoji",
    "wine": "Konkreter Weinname mit Region",
    "wine_type": "rot|weiss|rose|schaumwein",
    "region": "Weinregion, Land",
    "why": "Kurze, emotionale Begründung (max 2 Sätze) warum das funktioniert",
    "fun_fact": "Ein interessanter Fakt über diese Kombination (optional)"
}}

Beispiele für gute Tipps:
- Scharfes Thai-Curry + Gewürztraminer Spätlese → "Die exotische Süße zähmt die Schärfe perfekt"
- Pasta Carbonara + Pinot Grigio → "Frische trifft auf Cremigkeit"
- Dunkle Schokolade + Syrah → "Fruchtbomben-Harmonie"
- BBQ Ribs + Zinfandel → "Rauch trifft Rauch"

Sei kreativ! Vermeide langweilige Standard-Pairings."""

    try:
        # KI aufrufen
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        
        chat = LlmChat(
            api_key=EMERGENT_API_KEY,
            model="claude-sonnet-4-20250514",
            system_prompt="Du bist ein kreativer Sommelier, der überraschende Wein-Pairings empfiehlt."
        )
        
        response = await chat.send_message(UserMessage(text=prompt))
        response_text = response.message.text.strip()
        
        # JSON parsen
        import json
        # Entferne mögliche Markdown-Codeblocks
        if response_text.startswith("```"):
            response_text = response_text.split("```")[1]
            if response_text.startswith("json"):
                response_text = response_text[4:]
        response_text = response_text.strip()
        
        tip_data = json.loads(response_text)
        
        # Neuen Tipp erstellen
        new_tip = WeeklyTip(
            dish=tip_data.get("dish", "Unbekannt"),
            dish_emoji=tip_data.get("dish_emoji", "🍽️"),
            wine=tip_data.get("wine", "Unbekannt"),
            wine_type=tip_data.get("wine_type", "weiss"),
            region=tip_data.get("region"),
            why=tip_data.get("why", ""),
            fun_fact=tip_data.get("fun_fact"),
            week_number=week_number,
            year=year,
            language=language
        )
        
        # In Datenbank speichern
        await db.weekly_tips.insert_one(new_tip.model_dump())
        
        logger.info(f"Neuer Wochen-Tipp generiert: {new_tip.dish} + {new_tip.wine}")
        
        return {
            "status": "created",
            "message": f"Neuer Tipp für KW {week_number}/{year} erstellt",
            "tip": new_tip.model_dump()
        }
        
    except json.JSONDecodeError as e:
        logger.error(f"JSON Parse Error bei Tipp-Generierung: {e}")
        raise HTTPException(status_code=500, detail=f"KI-Antwort konnte nicht geparst werden: {str(e)}")
    except Exception as e:
        logger.error(f"Fehler bei Tipp-Generierung: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@api_router.post("/admin/seed-initial-tips")
async def seed_initial_tips():
    """
    Erstellt initiale Tipps für den Start (einmalig).
    """
    initial_tips = [
        {
            "dish": "Scharfes Thai-Curry",
            "dish_emoji": "🍛",
            "wine": "Gewürztraminer Spätlese",
            "wine_type": "weiss",
            "region": "Elsass, Frankreich",
            "why": "Die exotische Süße zähmt die Schärfe perfekt – ein echtes Aha-Erlebnis!",
            "fun_fact": "Der Gewürztraminer hat seinen Namen von den intensiven Gewürzaromen wie Litschi und Rose.",
            "week_number": 51,
            "year": 2025
        },
        {
            "dish": "Pasta Carbonara",
            "dish_emoji": "🍝",
            "wine": "Pinot Grigio",
            "wine_type": "weiss",
            "region": "Friaul, Italien",
            "why": "Frische trifft auf Cremigkeit – die knackige Säure schneidet durch die reichhaltige Sauce.",
            "fun_fact": "In Italien wird Carbonara traditionell NUR mit Guanciale (Schweinebacke) und Pecorino gemacht.",
            "week_number": 50,
            "year": 2025
        },
        {
            "dish": "Dunkle Schokolade",
            "dish_emoji": "🍫",
            "wine": "Kräftiger Syrah",
            "wine_type": "rot",
            "region": "Rhône-Tal, Frankreich",
            "why": "Ein Dessert-Traum – die dunklen Beerenaromen des Syrah umarmen die Bitterkeit der Schokolade.",
            "fun_fact": "Syrah und Shiraz sind derselbe Wein! In Frankreich heißt er Syrah, in Australien Shiraz.",
            "week_number": 49,
            "year": 2025
        },
        {
            "dish": "BBQ Spare Ribs",
            "dish_emoji": "🍖",
            "wine": "Zinfandel",
            "wine_type": "rot",
            "region": "Kalifornien, USA",
            "why": "Rauch trifft Rauch – die würzigen Brombeer-Noten ergänzen die süß-rauchige BBQ-Sauce perfekt.",
            "fun_fact": "Zinfandel ist genetisch identisch mit der italienischen Primitivo-Traube!",
            "week_number": 48,
            "year": 2025
        }
    ]
    
    inserted = 0
    for tip_data in initial_tips:
        # Prüfen ob Tipp schon existiert
        existing = await db.weekly_tips.find_one({
            "week_number": tip_data["week_number"],
            "year": tip_data["year"]
        })
        if not existing:
            tip = WeeklyTip(**tip_data, language="de")
            await db.weekly_tips.insert_one(tip.model_dump())
            inserted += 1
    
    return {
        "status": "success",
        "message": f"{inserted} initiale Tipps erstellt",
        "total_tips": await db.weekly_tips.count_documents({})
    }


# ===================== BLOG ENDPOINTS =====================

@api_router.get("/blog", response_model=List[BlogPost])
async def get_blog_posts(category: Optional[str] = None, limit: int = 20):
    """Get all published blog posts"""
    query = {"published": True}
    if category:
        query["category"] = category
    
    posts = await db.blog_posts.find(query, {"_id": 0}).sort("created_at", -1).to_list(limit)
    for post in posts:
        if isinstance(post.get('created_at'), str):
            post['created_at'] = datetime.fromisoformat(post['created_at'])
        if isinstance(post.get('updated_at'), str):
            post['updated_at'] = datetime.fromisoformat(post['updated_at'])
    return posts


@api_router.get("/blog-search")
async def search_blog_posts(q: str, limit: int = 50):
    """Search blog posts by query - searches title, excerpt, tags, region, content"""
    if not q or len(q) < 2:
        return []
    
    # Suche in mehreren Feldern mit Regex
    search_regex = {"$regex": q, "$options": "i"}
    query = {
        "published": True,
        "$or": [
            {"title": search_regex},
            {"title_en": search_regex},
            {"title_fr": search_regex},
            {"excerpt": search_regex},
            {"excerpt_en": search_regex},
            {"excerpt_fr": search_regex},
            {"region": search_regex},
            {"country": search_regex},
            {"tags": search_regex},
            {"content": search_regex},
            {"content_en": search_regex},
            {"content_fr": search_regex},
        ]
    }
    
    # Hole Posts mit Relevanz (Titel-Matches zuerst)
    posts = await db.blog_posts.find(query, {"_id": 0}).to_list(limit)
    
    # Sortiere nach Relevanz (Titel/Region Match zuerst)
    def relevance_score(post):
        score = 0
        q_lower = q.lower()
        if q_lower in (post.get('title') or '').lower():
            score += 100
        if q_lower in (post.get('region') or '').lower():
            score += 100
        if q_lower in ' '.join(post.get('tags') or []).lower():
            score += 50
        if q_lower in (post.get('excerpt') or '').lower():
            score += 30
        return score
    
    posts.sort(key=relevance_score, reverse=True)
    
    for post in posts:
        if isinstance(post.get('created_at'), str):
            post['created_at'] = datetime.fromisoformat(post['created_at'])
        if isinstance(post.get('updated_at'), str):
            post['updated_at'] = datetime.fromisoformat(post['updated_at'])
    
    return posts

@api_router.get("/blog/{slug}", response_model=BlogPost)
async def get_blog_post(slug: str):
    """Get a specific blog post by slug"""
    post = await db.blog_posts.find_one({"slug": slug, "published": True}, {"_id": 0})
    if not post:
        raise HTTPException(status_code=404, detail="Artikel nicht gefunden")
    if isinstance(post.get('created_at'), str):
        post['created_at'] = datetime.fromisoformat(post['created_at'])
    if isinstance(post.get('updated_at'), str):
        post['updated_at'] = datetime.fromisoformat(post['updated_at'])
    return post

@api_router.post("/blog", response_model=BlogPost)
async def create_blog_post(post_data: BlogPostCreate):
    """Create a new blog post"""
    # Check if slug exists
    existing = await db.blog_posts.find_one({"slug": post_data.slug})
    if existing:
        raise HTTPException(status_code=400, detail="Slug bereits vorhanden")
    
    post = BlogPost(**post_data.model_dump(), published=True)
    doc = post.model_dump()
    doc['created_at'] = doc['created_at'].isoformat()
    doc['updated_at'] = doc['updated_at'].isoformat()
    await db.blog_posts.insert_one(doc)
    return post

@api_router.get("/blog-categories")
async def get_blog_categories():
    """Get all blog categories with counts"""
    pipeline = [
        {"$match": {"published": True}},
        {"$group": {"_id": "$category", "count": {"$sum": 1}}},
        {"$sort": {"count": -1}}
    ]
    categories = await db.blog_posts.aggregate(pipeline).to_list(100)
    return [{"category": c["_id"], "count": c["count"]} for c in categories]

# ===================== SEO SITEMAP =====================

@api_router.get("/sitemap")
async def get_sitemap():
    """Generate sitemap data for SEO"""
    base_url = os.environ.get("FRONTEND_BASE_URL", FRONTEND_URL).rstrip("/")
    
    # Static pages
    pages = [
        {"url": f"{base_url}/", "priority": 1.0, "changefreq": "weekly"},
        {"url": f"{base_url}/pairing", "priority": 0.9, "changefreq": "weekly"},
        {"url": f"{base_url}/cellar", "priority": 0.8, "changefreq": "daily"},
        {"url": f"{base_url}/chat", "priority": 0.8, "changefreq": "weekly"},
        {"url": f"{base_url}/blog", "priority": 0.9, "changefreq": "daily"},
    ]
    
    # Blog posts
    posts = await db.blog_posts.find({"published": True}, {"slug": 1, "updated_at": 1, "_id": 0}).to_list(1000)
    for post in posts:
        pages.append({
            "url": f"{base_url}/blog/{post['slug']}",
            "priority": 0.7,
            "changefreq": "monthly",
            "lastmod": post.get('updated_at', '')
        })
    
    return {"pages": pages}

# ===================== SEED BLOG DATA =====================

@api_router.post("/seed-blog")
async def seed_blog_posts():
    """Seed initial blog posts for demonstration"""
    posts = [
        {
            "slug": "perfekte-weintemperatur",
            "title": "Die perfekte Weintemperatur – Der unterschätzte Genussfaktor",
            "title_en": "The Perfect Wine Temperature – The Underrated Pleasure Factor",
            "title_fr": "La température parfaite du vin – Le facteur plaisir sous-estimé",
            "excerpt": "Warum die richtige Temperatur über Genuss oder Enttäuschung entscheidet und wie Sie jeden Wein optimal servieren.",
            "excerpt_en": "Why the right temperature determines enjoyment or disappointment and how to serve every wine perfectly.",
            "excerpt_fr": "Pourquoi la bonne température détermine le plaisir ou la déception et comment servir chaque vin parfaitement.",
            "content": """## Die Wissenschaft hinter der Weintemperatur

Die Temperatur beeinflusst maßgeblich, wie wir Aromen wahrnehmen. Ein zu kalter Rotwein verschließt sich, seine Tannine wirken hart und die Frucht bleibt verborgen. Ein zu warmer Weißwein verliert seine Frische und wirkt plump.

### Die goldenen Regeln:

**Rotweine (16-18°C)**
- Leichte Rotweine wie Beaujolais: 14-16°C
- Mittelkräftige wie Pinot Noir: 15-17°C
- Kräftige wie Barolo oder Bordeaux: 17-18°C

**Weißweine (8-12°C)**
- Leichte, frische Weine: 8-10°C
- Gehaltvolle Weißweine mit Holz: 10-12°C
- Champagner & Schaumweine: 6-8°C

### Der Praxis-Tipp

Nehmen Sie Rotwein 30 Minuten vor dem Servieren aus dem Keller. Weißwein sollte etwa 20 Minuten vor dem Genuss aus dem Kühlschrank – nicht eiskalt, sondern mit spürbarer Kühle.""",
            "content_en": """## The Science Behind Wine Temperature

Temperature significantly influences how we perceive aromas. A too-cold red wine closes up, its tannins seem harsh, and the fruit remains hidden. A too-warm white wine loses its freshness and appears clumsy.

### The Golden Rules:

**Red Wines (16-18°C)**
- Light reds like Beaujolais: 14-16°C
- Medium-bodied like Pinot Noir: 15-17°C
- Full-bodied like Barolo or Bordeaux: 17-18°C

**White Wines (8-12°C)**
- Light, fresh wines: 8-10°C
- Full-bodied whites with oak: 10-12°C
- Champagne & sparkling: 6-8°C

### Practical Tip

Take red wine out of the cellar 30 minutes before serving. White wine should come out of the fridge about 20 minutes before – not ice cold, but with noticeable coolness.""",
            "image_url": "https://images.unsplash.com/photo-1510812431401-41d2bd2722f3?w=800",
            "category": "tipps",
            "tags": ["Temperatur", "Servieren", "Grundlagen"]
        },
        {
            "slug": "rotwein-zu-fisch",
            "title": "Rotwein zu Fisch? Warum alte Regeln nicht mehr gelten",
            "title_en": "Red Wine with Fish? Why Old Rules No Longer Apply",
            "title_fr": "Vin rouge avec du poisson? Pourquoi les anciennes règles ne s'appliquent plus",
            "excerpt": "Die Wein-Dogmen der Vergangenheit brechen auf. Entdecken Sie, welche Rotweine erstaunlich gut zu Fisch passen.",
            "excerpt_en": "The wine dogmas of the past are breaking down. Discover which red wines pair surprisingly well with fish.",
            "excerpt_fr": "Les dogmes vinicoles du passé s'effondrent. Découvrez quels vins rouges s'accordent étonnamment bien avec le poisson.",
            "content": """## Das Ende eines Mythos

„Weißwein zu Fisch, Rotwein zu Fleisch" – diese Regel hat Generationen von Weintrinkern geprägt. Doch die moderne Sommelierkunst hat erkannt: Es kommt auf die Zubereitung an, nicht nur auf das Hauptprodukt.

### Wann Rotwein zu Fisch funktioniert:

**1. Gegrillter oder gebratener Fisch**
Die Röstaromen vertragen sich wunderbar mit einem leichten Pinot Noir oder einem kühlen Gamay.

**2. Fisch in Rotwein-Sauce**
Logisch: Wenn Rotwein im Gericht ist, sollte er auch im Glas sein.

**3. Thunfisch und Lachs**
Diese fetteren Fische mit ihrem kräftigen Eigengeschmack harmonieren mit leichten, fruchtigen Rotweinen.

### Die Faustregel

Je mehr Umami und Röstaromen im Gericht, desto eher funktioniert ein leichter Rotwein. Meiden Sie tanninreiche Weine – die Gerbstoffe können mit Fischölen metallisch schmecken.""",
            "image_url": "https://images.unsplash.com/photo-1534604973900-c43ab4c2e0ab?w=800",
            "category": "pairings",
            "tags": ["Fisch", "Rotwein", "Pairing", "Mythen"]
        },
        {
            "slug": "weinregion-burgund",
            "title": "Burgund verstehen: Eine Reise durch Frankreichs Herzstück",
            "title_en": "Understanding Burgundy: A Journey Through France's Heartland",
            "title_fr": "Comprendre la Bourgogne: Un voyage au cœur de la France",
            "excerpt": "Von Chablis bis Beaujolais – wie Sie die komplexe Welt burgundischer Weine entschlüsseln.",
            "excerpt_en": "From Chablis to Beaujolais – how to decode the complex world of Burgundy wines.",
            "excerpt_fr": "De Chablis au Beaujolais – comment décoder le monde complexe des vins de Bourgogne.",
            "content": """## Warum Burgund so besonders ist

Keine andere Weinregion der Welt hat die Idee des Terroirs so perfektioniert wie Burgund. Hier zählt jeder Meter Boden, jede Hangneigung, jedes Mikroklima.

### Die Hierarchie verstehen:

**Grand Cru** (2% der Produktion)
Die Spitze: 33 Lagen für Rotwein, 8 für Weißwein. Namen wie Romanée-Conti oder Montrachet.

**Premier Cru** (10% der Produktion)
Exzellente Einzellagen, oft mit bestem Preis-Leistungs-Verhältnis.

**Village** (35% der Produktion)
Weine aus benannten Gemeinden: Gevrey-Chambertin, Meursault, Pommard.

**Bourgogne** (53% der Produktion)
Regionale Weine – der Einstieg in die burgundische Welt.

### Mein Geheimtipp

Suchen Sie nach Premier Crus aus weniger bekannten Dörfern wie Savigny-lès-Beaune oder Saint-Romain. Hier finden Sie großartige Qualität zu vernünftigen Preisen.""",
            "image_url": "https://images.unsplash.com/photo-1560493676-04071c5f467b?w=800",
            "category": "regionen",
            "tags": ["Burgund", "Frankreich", "Pinot Noir", "Chardonnay"]
        },
        {
            "slug": "dekantieren-wann-warum",
            "title": "Dekantieren: Wann es Sinn macht und wann nicht",
            "title_en": "Decanting: When It Makes Sense and When It Doesn't",
            "title_fr": "Décanter: Quand c'est utile et quand ça ne l'est pas",
            "excerpt": "Nicht jeder Wein braucht eine Karaffe. Lernen Sie, welche Weine vom Dekantieren profitieren.",
            "excerpt_en": "Not every wine needs a decanter. Learn which wines benefit from decanting.",
            "excerpt_fr": "Tous les vins n'ont pas besoin d'une carafe. Apprenez quels vins bénéficient de la décantation.",
            "content": """## Die Kunst des Dekantierens

Dekantieren hat zwei Funktionen: Belüftung und Trennung vom Depot. Doch nicht jeder Wein braucht beides – oder überhaupt eines davon.

### Wann Sie dekantieren sollten:

**Junge, tanninreiche Rotweine**
- Bordeaux unter 10 Jahren: 1-2 Stunden
- Barolo, Barbaresco: 2-3 Stunden
- Cabernet Sauvignon aus Übersee: 1-2 Stunden

**Alte Weine mit Depot**
Vorsichtig umfüllen, Depot im Flaschenhals stoppen. Aber: nicht zu lange atmen lassen – alte Weine sind empfindlich!

### Wann Sie NICHT dekantieren sollten:

- **Leichte Rotweine** wie Beaujolais oder Valpolicella
- **Alte, fragile Weine** über 20 Jahre
- **Die meisten Weißweine** (Ausnahme: sehr junge, hochwertige Burgunder)
- **Schaumweine** – niemals!

### Die Alternative

Kein Dekanter zur Hand? Schwenken Sie den Wein kräftig im Glas. Das beschleunigt die Belüftung erstaunlich effektiv.""",
            "image_url": "https://images.unsplash.com/photo-1569919659476-f0852f9f8ede?w=800",
            "category": "wissen",
            "tags": ["Dekantieren", "Karaffe", "Servieren", "Tipps"]
        },
        {
            "slug": "wein-lagerung-zuhause",
            "title": "Wein richtig lagern: So bauen Sie Ihren Heimkeller auf",
            "title_en": "Storing Wine Properly: How to Build Your Home Cellar",
            "title_fr": "Bien conserver le vin: Comment aménager votre cave à domicile",
            "excerpt": "Die wichtigsten Regeln für die Weinlagerung zu Hause – auch ohne echten Weinkeller.",
            "excerpt_en": "The most important rules for storing wine at home – even without a real wine cellar.",
            "excerpt_fr": "Les règles les plus importantes pour conserver le vin à la maison – même sans vraie cave.",
            "content": """## Die vier Feinde des Weins

**1. Licht**
UV-Strahlen zerstören Aromen. Dunkle Flaschen schützen besser, aber Dunkelheit ist immer am besten.

**2. Temperaturschwankungen**
Konstante 12-14°C sind ideal. Schwankungen sind schlimmer als eine etwas zu hohe Durchschnittstemperatur.

**3. Erschütterungen**
Vibrationen stören die Reifung. Nicht neben der Waschmaschine lagern!

**4. Trockene Luft**
Korken können austrocknen. Idealfeuchte: 70%.

### Praktische Lösungen:

**Für Einsteiger:**
Ein temperierter Kleiderschrank in einem kühlen Raum reicht für Weine, die Sie innerhalb eines Jahres trinken.

**Für Ambitionierte:**
Ein Weintemperierschrank (ab 300€) hält konstante Temperatur und Luftfeuchtigkeit.

**Für Sammler:**
Ein echter Keller mit Klimatisierung ist die Königsklasse.

### Mein Tipp

Lagern Sie Flaschen liegend, damit der Korken feucht bleibt. Schraubverschluss? Stehend ist auch okay.""",
            "image_url": "https://images.unsplash.com/photo-1560493676-04071c5f467b?w=800",
            "category": "tipps",
            "tags": ["Lagerung", "Weinkeller", "Aufbewahrung", "Grundlagen"]
        }
    ]
    
    # Clear existing and insert new
    await db.blog_posts.delete_many({})
    
    for post_data in posts:
        post = BlogPost(**post_data, published=True, author="Sommelier Team")
        doc = post.model_dump()
        doc['created_at'] = doc['created_at'].isoformat()
        doc['updated_at'] = doc['updated_at'].isoformat()
        await db.blog_posts.insert_one(doc)
    
    return {"message": f"{len(posts)} Blog-Artikel wurden erstellt"}


# ===================== NEW PUBLIC WINES ENDPOINT (CLEAN) =====================

@api_router.get("/public-wines", response_model=List[WineDatabaseEntry])
async def get_public_wines(
    search: Optional[str] = None,
    country: Optional[str] = None,
    region: Optional[str] = None,
    wine_color: Optional[str] = None,
    price_category: Optional[str] = None,
    skip: int = 0,
    limit: int = 50
):
    """
    Get wines from the new public_wines collection.
    WICHTIG: Verwendet create_accent_insensitive_pattern() für Akzent-Suche!
    """
    logger.info(f"PUBLIC WINES ENDPOINT called: search={search}, limit={limit}")
    
    query = {}
    
    if search:
        # WICHTIG: Globale Hilfsfunktion für Akzent-insensitive Suche verwenden!
        # "Chateau" findet "Château", "Cotes" findet "Côtes"
        accent_pattern = create_accent_insensitive_pattern(search)
        regex = {"$regex": accent_pattern, "$options": "i"}
        
        query["$or"] = [
            {"name": regex},
            {"winery": regex},
            {"region": regex},
            {"grape_variety": regex},
        ]
    
    if country:
        query["country"] = country
    if region:
        # =================================================================
        # REGION FILTER FIX: Match in multiple fields
        # =================================================================
        # Wines can have region info in different fields:
        # - region: "Barbaresco", "Südtirol"
        # - appellation: "Barbaresco DOCG", "Alto Adige"
        # - anbaugebiet: some imports use this field
        # 
        # Solution: Search in all region-related fields
        # =================================================================
        region_regex = {"$regex": re.escape(region), "$options": "i"}
        region_conditions = [
            {"region": region_regex},
            {"appellation": region_regex},
            {"anbaugebiet": region_regex}
        ]
        
        # Combine with existing $or if present (from search)
        if "$or" in query:
            # Wrap existing query in $and with region conditions
            existing_or = query.pop("$or")
            query["$and"] = [
                {"$or": existing_or},
                {"$or": region_conditions}
            ]
        else:
            query["$or"] = region_conditions
    if wine_color:
        query["wine_color"] = wine_color
    if price_category:
        query["price_category"] = price_category
    
    wines = await db.public_wines.find(query, {"_id": 0}).skip(skip).limit(limit).to_list(limit)
    
    logger.info(f"PUBLIC WINES: Found {len(wines)} wines")
    
    # Convert created_at strings to datetime
    for wine in wines:
        if isinstance(wine.get('created_at'), str):
            wine['created_at'] = datetime.fromisoformat(wine['created_at'])
    
    return wines


@api_router.get("/public-wines/{wine_id}", response_model=WineDatabaseEntry)
async def get_public_wine_detail(wine_id: str):
    """Get details of a specific wine from public_wines collection"""
    wine = await db.public_wines.find_one({"id": wine_id}, {"_id": 0})
    
    if not wine:
        raise HTTPException(status_code=404, detail="Wein nicht gefunden")
    
    if isinstance(wine.get('created_at'), str):
        wine['created_at'] = datetime.fromisoformat(wine['created_at'])
    
    return wine


class AutoAddWineRequest(BaseModel):
    """Request model for auto-adding wines from Claude recommendations"""
    name: str
    grape: Optional[str] = None
    region: Optional[str] = None
    country: Optional[str] = None
    color: Optional[str] = None
    description_de: Optional[str] = None
    description_en: Optional[str] = None
    description_fr: Optional[str] = None
    source: Optional[str] = "claude_recommendation"


@api_router.post("/public-wines/auto-add")
async def auto_add_wine_from_recommendation(request: AutoAddWineRequest):
    """
    Automatisch einen Wein zur öffentlichen Datenbank hinzufügen.
    Wird aufgerufen wenn Claude einen Wein empfiehlt, der noch nicht in der DB ist.
    """
    # Prüfe ob Wein bereits existiert
    existing = await db.public_wines.find_one({
        "name": {"$regex": f"^{request.name}$", "$options": "i"}
    })
    
    if existing:
        # Wein existiert bereits - gib ihn zurück
        existing.pop('_id', None)
        return existing
    
    # Neuen Wein erstellen
    new_wine = {
        "id": str(uuid.uuid4()),
        "name": request.name,
        "grape": request.grape or "",
        "region": request.region or "",
        "country": request.country or "",
        "color": request.color or "",
        "description_de": request.description_de or f"{request.name} - Ein von Claude empfohlener Wein.",
        "description_en": request.description_en or f"{request.name} - A wine recommended by Claude.",
        "description_fr": request.description_fr or f"{request.name} - Un vin recommandé par Claude.",
        "source": request.source,
        "auto_added": True,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.public_wines.insert_one(new_wine)
    new_wine.pop('_id', None)
    
    logger.info(f"🍷 Auto-added wine to database: {request.name}")
    
    return new_wine


# ==============================================================================
# REGION SIMPLIFICATION FIX (Dezember 2024)
# ==============================================================================
# PROBLEM: Swiss wines are stored with detailed sub-regions like:
#   - "Genf - Anières", "Genf - Satigny", "Genf - Dardagny" (9 sub-regions)
#   - "Wallis - Sion", "Wallis - Salgesch", "Wallis - Fully" (20+ sub-regions)
#   - "Waadt - Aigle", "Waadt - Lavaux", "Waadt - Yvorne" (20+ sub-regions)
#   - "Tessin - Lugano", "Tessin - Mendrisio", etc. (20+ sub-regions)
#
# This made the region dropdown unreadable with 133 entries.
#
# SOLUTION: 
#   1. simplify_region() extracts just the canton name (before " - ")
#   2. Region filter uses regex "^Genf" to match all sub-regions
#   3. Result: Clean dropdown with only 64 entries, Swiss cantons as single items
#
# IMPORTANT: When adding new Swiss wines, use format "Kanton - Unterregion"
#            e.g., "Genf - Satigny", "Wallis - Sion"
# ==============================================================================

def simplify_region(region: str) -> str:
    """
    Simplify region names by extracting just the canton/main region.
    
    Swiss wines are stored with detailed sub-regions (e.g., 'Genf - Satigny').
    For the filter dropdown, we only show the canton (e.g., 'Genf').
    
    Examples:
        'Genf - Satigny' -> 'Genf'
        'Wallis - Sion' -> 'Wallis'
        'Tessin - Lugano' -> 'Tessin'
        'Burgund' -> 'Burgund' (no change for non-Swiss regions)
    
    Args:
        region: The full region name from database
        
    Returns:
        Simplified region name (canton only for Swiss wines)
    """
    if not region:
        return region
    # Check if region contains " - " separator (Swiss canton format)
    if " - " in region:
        return region.split(" - ")[0].strip()
    return region


@api_router.post("/admin/estimate-wine-prices")
async def estimate_wine_prices():
    """
    Admin endpoint to estimate price categories for public wines based on heuristics.
    Uses region, appellation, and wine type to estimate prices.
    
    Categories:
    - '1' = 🍷 bis €20 (everyday wines)
    - '2' = 🍷🍷 €20-50 (mid-range)
    - '3' = 🍷🍷🍷 ab €50 (premium/luxury)
    """
    
    # Premium regions/appellations that typically cost €50+
    luxury_indicators = [
        'romanée', 'montrachet', 'chambertin', 'musigny', 'richebourg',
        'la tâche', 'petrus', 'margaux', 'pauillac', 'saint-julien',
        'saint-estèphe', 'pessac-léognan', 'pomerol', 'saint-émilion grand cru',
        'barolo', 'barbaresco', 'brunello', 'amarone', 'sassicaia',
        'tignanello', 'ornellaia', 'masseto', 'dom pérignon', 'krug',
        'cristal', 'château', 'grand cru', 'premier cru'
    ]
    
    # Mid-range regions/appellations typically €20-50
    midrange_indicators = [
        'chablis', 'meursault', 'puligny', 'chassagne', 'gevrey',
        'nuits-saint-georges', 'beaune', 'côte de beaune', 'côte de nuits',
        'saint-joseph', 'crozes-hermitage', 'gigondas', 'vacqueyras',
        'châteauneuf-du-pape', 'hermitage', 'côte-rôtie',
        'chianti classico', 'valpolicella ripasso', 'langhe', 'montalcino',
        'rioja reserva', 'ribera del duero', 'priorat',
        'mosel', 'rheingau', 'pfalz', 'nahe', 'spätlese', 'auslese',
        'grüner veltliner smaragd', 'wachau'
    ]
    
    updated_count = 0
    wines_without_category = await db.public_wines.find(
        {"$or": [{"price_category": None}, {"price_category": {"$exists": False}}, {"price_category": ""}]}
    ).to_list(10000)
    
    logger.info(f"Found {len(wines_without_category)} wines without price category")
    
    for wine in wines_without_category:
        # Build search string from wine attributes
        search_text = ' '.join([
            str(wine.get('name', '')),
            str(wine.get('region', '')),
            str(wine.get('appellation', '')),
            str(wine.get('winery', ''))
        ]).lower()
        
        # Determine category based on indicators
        category = '1'  # Default: everyday wine
        
        # Check for luxury indicators
        for indicator in luxury_indicators:
            if indicator in search_text:
                category = '3'
                break
        
        # If not luxury, check for mid-range
        if category == '1':
            for indicator in midrange_indicators:
                if indicator in search_text:
                    category = '2'
                    break
        
        # Update wine
        await db.public_wines.update_one(
            {"id": wine['id']},
            {"$set": {"price_category": category}}
        )
        updated_count += 1
    
    return {
        "status": "success",
        "message": f"Updated {updated_count} wines with estimated price categories",
        "details": {
            "total_processed": len(wines_without_category),
            "updated": updated_count
        }
    }


@api_router.get("/public-wines-filters")
async def get_public_wines_filters(country: Optional[str] = None, region: Optional[str] = None):
    """
    Get available filter options for public wines with cascading support.
    
    NOTE: Regions are simplified for display (see simplify_region function).
    When filtering by a simplified region like 'Genf', the API will match
    all sub-regions (Genf - Satigny, Genf - Dardagny, etc.) using regex.
    """
    
    # Base query
    query = {}
    if country and country != 'all':
        query["country"] = country
    if region and region != 'all':
        # When filtering by simplified region (e.g., "Genf"), match all sub-regions
        # Use proper $regex operator for MongoDB distinct() query
        query["region"] = {"$regex": f"^{re.escape(region)}", "$options": "i"}
    
    # Get all distinct values
    countries = await db.public_wines.distinct("country", {})
    
    # Get regions from multiple fields
    country_filter = {"country": country} if country and country != 'all' else {}
    raw_regions = await db.public_wines.distinct("region", country_filter)
    raw_appellations = await db.public_wines.distinct("appellation", country_filter)
    
    appellations = await db.public_wines.distinct("appellation", query)
    colors = await db.public_wines.distinct("wine_color")
    price_categories = await db.public_wines.distinct("price_category")
    grape_varieties = await db.public_wines.distinct("grape_variety")
    
    # Combine regions and appellations for the filter dropdown
    # This ensures "Barbaresco" appears even if it's only in appellation field
    simplified_regions = set()
    for r in raw_regions:
        if r and r != 'Unbekannt' and r.strip():
            simplified_regions.add(simplify_region(r))
    
    # For countries with well-defined regions (France, Germany, etc.), 
    # DON'T add appellations to region dropdown - keep them separate
    # Only add appellations as region options for countries like Italy where 
    # the region field might be empty but appellation contains location info
    countries_with_clean_regions = {'Frankreich', 'Deutschland', 'Österreich', 'Schweiz', 'Spanien', 'Italien'}
    
    if country not in countries_with_clean_regions:
        # Add major appellations as region options (like Barbaresco, Barolo, Chianti)
        for a in raw_appellations:
            if a and a != 'Unbekannt' and a.strip():
                # Only add short appellation names (not full DOCG names)
                if len(a) < 30 and 'DOCG' not in a and 'DOC' not in a:
                    simplified_regions.add(a)
    
    # Build hierarchy map
    hierarchy = {}
    if not country or country == 'all':
        # Get all countries with their regions
        all_wines = await db.public_wines.find({}, {"_id": 0, "country": 1, "region": 1, "appellation": 1}).to_list(10000)
        for wine in all_wines:
            c = wine.get('country')
            r = wine.get('region')
            a = wine.get('appellation')
            if c and c != 'Unbekannt':
                if c not in hierarchy:
                    hierarchy[c] = {}
                if r and r != 'Unbekannt':
                    # Use simplified region for hierarchy
                    simplified_r = simplify_region(r)
                    if simplified_r not in hierarchy[c]:
                        hierarchy[c][simplified_r] = set()
                    if a and a != 'Unbekannt':
                        hierarchy[c][simplified_r].add(a)
        
        # Convert sets to sorted lists
        for c in hierarchy:
            for r in hierarchy[c]:
                hierarchy[c][r] = sorted(list(hierarchy[c][r]))
    
    return {
        "countries": sorted([c for c in countries if c and c != 'Unbekannt']),
        "regions": sorted(list(simplified_regions)),
        "appellations": sorted([a for a in appellations if a and a != 'Unbekannt']),
        "wine_colors": sorted([c for c in colors if c]),
        "price_categories": sorted([p for p in price_categories if p]),
        "grape_varieties": sorted([g for g in grape_varieties if g and g != 'Unbekannt']),
        "hierarchy": hierarchy
    }


# Add CORS middleware BEFORE including router (critical for proper request handling)
# Get allowed origins - if wildcard, use specific origins for credentials support
cors_origins_env = os.environ.get('CORS_ORIGINS', '*')
if cors_origins_env == '*':
    # Default allowed origins for credentials
    allowed_origins = [
        "http://localhost:3000",
        "https://localhost:3000",
        "https://wine-promo-suite.preview.emergentagent.com",
        "https://wine-pairing.online",
        "https://www.wine-pairing.online"
    ]
else:
    allowed_origins = cors_origins_env.split(',')

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=allowed_origins,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ===================== BACKUP DOWNLOAD ENDPOINTS =====================

@api_router.get("/backup/list")
async def list_backup_files():
    """List all available backup files"""
    data_dir = ROOT_DIR / "data"
    backup_files = []
    
    for f in sorted(data_dir.glob("*.json")):
        size_kb = f.stat().st_size / 1024
        backup_files.append({
            "filename": f.name,
            "size_kb": round(size_kb, 1),
            "download_url": f"/api/backup/download/{f.name}"
        })
    
    return {"files": backup_files, "total_files": len(backup_files)}


@api_router.get("/backup/download/{filename}")
async def download_backup_file(filename: str):
    """Download a specific backup file"""
    # Security: Only allow .json files from data directory
    if not filename.endswith('.json'):
        raise HTTPException(status_code=400, detail="Only JSON files are allowed")
    
    # Prevent path traversal
    if '/' in filename or '\\' in filename or '..' in filename:
        raise HTTPException(status_code=400, detail="Invalid filename")
    
    data_dir = ROOT_DIR / "data"
    file_path = data_dir / filename
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail=f"File not found: {filename}")
    
    # Read file content
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    return Response(
        content=content,
        media_type="application/json",
        headers={
            "Content-Disposition": f"attachment; filename={filename}"
        }
    )


@api_router.get("/backup/download-all")
async def download_all_backups():
    """Get all backup data as a single JSON object"""
    data_dir = ROOT_DIR / "data"
    all_data = {}
    
    for f in sorted(data_dir.glob("*.json")):
        try:
            with open(f, 'r', encoding='utf-8') as file:
                all_data[f.stem] = json.load(file)
        except Exception as e:
            all_data[f.stem] = {"error": str(e)}
    
    return all_data


@api_router.get("/export/excel/{collection_name}")
async def export_collection_excel(collection_name: str):
    """Export a collection as Excel file for download"""
    import pandas as pd
    from io import BytesIO
    
    valid_collections = [
        'public_wines', 'wine_database', 'grape_varieties', 'regional_pairings',
        'blog_posts', 'feed_posts', 'dishes', 'seo_pairings', 'coupons'
    ]
    
    if collection_name not in valid_collections:
        raise HTTPException(status_code=400, detail=f"Collection nicht verfügbar. Gültig: {valid_collections}")
    
    try:
        docs = await db[collection_name].find({}, {'_id': 0}).to_list(None)
        if not docs:
            raise HTTPException(status_code=404, detail="Collection ist leer")
        
        df = pd.DataFrame(docs)
        
        # Excel in Memory erstellen
        output = BytesIO()
        df.to_excel(output, index=False, engine='openpyxl')
        output.seek(0)
        
        from fastapi.responses import StreamingResponse
        
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={
                "Content-Disposition": f"attachment; filename={collection_name}.xlsx"
            }
        )
    except Exception as e:
        logger.error(f"Excel export error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@api_router.get("/export/excel-links")
async def get_excel_download_links():
    """Get all Excel download links with current counts"""
    collections = [
        'public_wines', 'wine_database', 'grape_varieties', 'regional_pairings',
        'blog_posts', 'feed_posts', 'dishes', 'seo_pairings', 'coupons'
    ]
    
    links = []
    total = 0
    
    for coll in collections:
        count = await db[coll].count_documents({})
        total += count
        links.append({
            "collection": coll,
            "count": count,
            "excel_url": f"/api/export/excel/{coll}",
            "json_url": f"/api/backup/download/{coll}.json"
        })
    
    return {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "total_documents": total,
        "downloads": links
    }


@api_router.get("/docs/download")
async def download_documentation():
    """Download the complete app documentation as Excel"""
    import pandas as pd
    from io import BytesIO
    
    # Lade Dokumentation
    doc_path = ROOT_DIR / "docs" / "APP_DOKUMENTATION_KOMPLETT.md"
    if not doc_path.exists():
        doc_path = Path("/app/docs/APP_DOKUMENTATION_KOMPLETT.md")
    
    if not doc_path.exists():
        raise HTTPException(status_code=404, detail="Dokumentation nicht gefunden")
    
    # Erstelle Excel mit mehreren Sheets
    output = BytesIO()
    
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        # Sheet 1: Übersicht
        overview_data = {
            'Metrik': ['Weine', 'Rebsorten', 'Sommelier Kompass', 'Blog-Artikel', 'Community Feed', 'Sprachen', 'Stand'],
            'Wert': ['7,066 (wächst dynamisch)', '313', '1,652 Gerichte', '233', '268 Beiträge', 'DE, EN, FR', '17.12.2025']
        }
        pd.DataFrame(overview_data).to_excel(writer, sheet_name='Übersicht', index=False)
        
        # Sheet 2: Features
        features_data = {
            'Feature': ['Pairing', 'Weinkeller', 'Chat', 'Rebsorten-Lexikon', 'Wein-Datenbank', 'Sommelier Kompass', 'Community Feed', 'Blog'],
            'Route': ['/pairing', '/cellar', '/chat', '/grapes', '/wine-database', '/sommelier-kompass', '/feed', '/blog'],
            'Zugriff Basic': ['5/Tag', '10 Weine', '5/Tag', 'Unbegrenzt', 'Unbegrenzt', 'Unbegrenzt', 'Unbegrenzt', 'Unbegrenzt'],
            'Zugriff Pro': ['Unbegrenzt', 'Unbegrenzt', 'Unbegrenzt', 'Unbegrenzt', 'Unbegrenzt', 'Unbegrenzt', 'Unbegrenzt', 'Unbegrenzt']
        }
        pd.DataFrame(features_data).to_excel(writer, sheet_name='Features', index=False)
        
        # Sheet 3: API Endpoints
        api_data = {
            'Methode': ['POST', 'POST', 'POST', 'GET', 'POST', 'GET', 'POST', 'PUT', 'DELETE', 'GET', 'GET', 'GET', 'GET'],
            'Endpoint': ['/api/auth/register', '/api/auth/login', '/api/pairing', '/api/chat', '/api/wines', '/api/wines', '/api/wines', '/api/wines/{id}', '/api/wines/{id}', '/api/public-wines', '/api/grape-varieties', '/api/regional-pairings', '/api/export/excel/{collection}'],
            'Beschreibung': ['Registrieren', 'Einloggen', 'Weinempfehlung', 'Chat', 'Wein hinzufügen', 'Eigene Weine', 'Wein hinzufügen', 'Bearbeiten', 'Löschen', 'Wein-Datenbank', 'Rebsorten', 'Sommelier Kompass', 'Excel-Export']
        }
        pd.DataFrame(api_data).to_excel(writer, sheet_name='API Endpoints', index=False)
        
        # Sheet 4: Datenbank
        db_data = {
            'Collection': ['public_wines', 'wine_database', 'grape_varieties', 'regional_pairings', 'blog_posts', 'feed_posts', 'dishes', 'seo_pairings', 'users', 'wines', 'coupons'],
            'Anzahl': ['7,066', '494', '313', '1,652', '233', '268', '40', '500', '~20', '~40', '100'],
            'Beschreibung': ['Öffentliche Weine', 'Erweiterte Wein-Infos', 'Rebsorten', 'Sommelier Kompass', 'Blog-Artikel', 'Community Feed', 'Gerichte', 'SEO Pairings', 'Benutzer', 'Persönliche Weinkeller', 'Gutscheine'],
            'Wachstum': ['Dynamisch (KI)', 'Statisch', 'Statisch', 'Statisch', 'Manuell', 'User-generiert', 'Statisch', 'Statisch', 'User-generiert', 'User-generiert', 'Manuell']
        }
        pd.DataFrame(db_data).to_excel(writer, sheet_name='Datenbank', index=False)
        
        # Sheet 5: Preise
        pricing_data = {
            'Plan': ['Basic', 'Pro Monatlich', 'Pro Jährlich'],
            'Preis': ['Kostenlos', '4,99€/Monat', '39,99€/Jahr'],
            'Pairing/Tag': ['5', 'Unbegrenzt', 'Unbegrenzt'],
            'Chat/Tag': ['5', 'Unbegrenzt', 'Unbegrenzt'],
            'Weinkeller': ['Max. 10', 'Unbegrenzt', 'Unbegrenzt'],
            'Favoriten': ['Max. 10', 'Unbegrenzt', 'Unbegrenzt']
        }
        pd.DataFrame(pricing_data).to_excel(writer, sheet_name='Preise', index=False)
    
    output.seek(0)
    
    from fastapi.responses import StreamingResponse
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={
            "Content-Disposition": "attachment; filename=Wine_Pairing_App_Dokumentation.xlsx"
        }
    )


@api_router.get("/docs/download-md")
async def download_documentation_markdown():
    """Download the complete app documentation as Markdown file"""
    from fastapi.responses import FileResponse
    
    doc_path = Path("/app/docs/APP_DOKUMENTATION_KOMPLETT.md")
    
    if not doc_path.exists():
        raise HTTPException(status_code=404, detail="Dokumentation nicht gefunden")
    
    return FileResponse(
        path=str(doc_path),
        media_type="text/markdown",
        filename="Wine_Pairing_App_Dokumentation.md"
    )


@api_router.get("/docs/download-word")
async def download_documentation_word():
    """Download the complete app documentation as Word document"""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO
    import re
    
    doc_path = Path("/app/docs/APP_DOKUMENTATION_KOMPLETT.md")
    
    if not doc_path.exists():
        raise HTTPException(status_code=404, detail="Dokumentation nicht gefunden")
    
    # Lese Markdown
    with open(doc_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Erstelle Word-Dokument
    doc = Document()
    
    # Titel
    title = doc.add_heading('WINE PAIRING APP', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Vollständige Dokumentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Stand: 17. Dezember 2025')
    doc.add_paragraph(f'Domain: https://wine-pairing.online')
    doc.add_paragraph('---')
    
    # Parse Markdown und füge zum Word-Dokument hinzu
    lines = md_content.split('\n')
    current_table = []
    in_table = False
    
    for line in lines:
        line = line.strip()
        
        # Skip leere Zeilen
        if not line:
            if in_table and current_table:
                # Tabelle beenden
                in_table = False
                current_table = []
            continue
        
        # Überschriften
        if line.startswith('# '):
            doc.add_heading(line[2:].replace('📖 ', '').replace('📊 ', ''), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].replace('💳 ', '').replace('🏠 ', '').replace('📚 ', '').replace('🔐 ', '').replace('🤖 ', '').replace('💾 ', '').replace('🌐 ', '').replace('🔌 ', '').replace('🚀 ', '').replace('📋 ', '').replace('📥 ', ''), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].replace('🍷 ', '').replace('🍇 ', '').replace('🧭 ', '').replace('👥 ', '').replace('📝 ', '').replace('❤️ ', '').replace('💬 ', '').replace('🎟️ ', ''), level=3)
        elif line.startswith('---'):
            doc.add_paragraph('─' * 50)
        elif line.startswith('|') and '|' in line[1:]:
            # Tabellen-Zeile - als Text einfügen
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells and not all(c.replace('-', '').replace(':', '') == '' for c in cells):
                doc.add_paragraph('  │  '.join(cells))
        elif line.startswith('- '):
            # Liste
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('**') and line.endswith('**'):
            # Fett
            p = doc.add_paragraph()
            p.add_run(line.replace('**', '')).bold = True
        else:
            # Normaler Text
            clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)  # Remove bold markers
            clean_line = re.sub(r'`(.*?)`', r'\1', clean_line)  # Remove code markers
            if clean_line and not clean_line.startswith('```'):
                doc.add_paragraph(clean_line)
    
    # Speichere in BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    from fastapi.responses import StreamingResponse
    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": "attachment; filename=Wine_Pairing_App_Dokumentation.docx"
        }
    )


# ===================== AUTHENTICATION & SUBSCRIPTION ENDPOINTS =====================

async def get_current_user(request: Request) -> Optional[User]:
    """Get current user from JWT token (cookie or header)"""
    # Try cookie first
    token = request.cookies.get("session_token")
    
    # Fallback to Authorization header
    if not token:
        auth_header = request.headers.get("Authorization")
        if auth_header and auth_header.startswith("Bearer "):
            token = auth_header[7:]
    
    if not token:
        return None
    
    # Decode JWT token
    payload = decode_jwt_token(token)
    if not payload:
        return None
    
    user_id = payload.get("user_id")
    if not user_id:
        return None
    
    # Get user from database
    user_doc = await db.users.find_one(
        {"user_id": user_id},
        {"_id": 0, "password_hash": 0}  # Exclude password
    )
    
    if not user_doc:
        return None
    
    return User(**user_doc)

async def reset_daily_usage_if_needed(user: User) -> User:
    """Reset daily usage counters if it's a new day"""
    today = datetime.now(timezone.utc).date().isoformat()
    last_date = user.usage.get("last_usage_date")
    
    if last_date != today:
        # New day - reset counters
        user.usage = {
            "pairing_requests_today": 0,
            "chat_messages_today": 0,
            "last_usage_date": today
        }
        await db.users.update_one(
            {"user_id": user.user_id},
            {"$set": {"usage": user.usage}}
        )
    
    return user

async def check_limit(user: Optional[User], limit_type: str) -> tuple[bool, str]:
    """Check if user has reached their limit. Returns (allowed, message)"""
    if user is None:
        # Anonymous user - use basic limits
        plan = "basic"
        usage_count = 0
    else:
        user = await reset_daily_usage_if_needed(user)
        plan = user.plan
        if limit_type == "pairing":
            usage_count = user.usage.get("pairing_requests_today", 0)
        elif limit_type == "chat":
            usage_count = user.usage.get("chat_messages_today", 0)
        else:
            usage_count = 0
    
    limits = FREEMIUM_LIMITS[plan]
    
    if limit_type == "pairing":
        limit = limits["pairing_requests_per_day"]
        if usage_count >= limit:
            return False, f"Tageslimit erreicht ({int(limit)} Anfragen). Upgrade auf Pro für unbegrenzte Nutzung!"
    elif limit_type == "chat":
        limit = limits["chat_messages_per_day"]
        if usage_count >= limit:
            return False, f"Tageslimit erreicht ({int(limit)} Nachrichten). Upgrade auf Pro für unbegrenzte Nutzung!"
    elif limit_type == "cellar":
        limit = limits["max_cellar_wines"]
        if user:
            cellar_count = await db.wines.count_documents({"user_id": user.user_id})
        else:
            cellar_count = await db.wines.count_documents({})
        if cellar_count >= limit:
            return False, f"Maximale Anzahl Weine erreicht ({int(limit)}). Upgrade auf Pro für unbegrenzten Keller!"
    
    return True, ""

async def increment_usage(user: User, usage_type: str):
    """Increment usage counter for user"""
    if usage_type == "pairing":
        field = "usage.pairing_requests_today"
    elif usage_type == "chat":
        field = "usage.chat_messages_today"
    else:
        return
    
    await db.users.update_one(
        {"user_id": user.user_id},
        {
            "$inc": {field: 1},
            "$set": {"usage.last_usage_date": datetime.now(timezone.utc).date().isoformat()}
        }
    )

# Auth endpoints - JWT Email/Password

def hash_password(password: str) -> str:
    """Hash password using bcrypt"""
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def verify_password(password: str, hashed: str) -> bool:
    """Verify password against hash"""
    return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))

def create_jwt_token(user_id: str, email: str) -> str:
    """Create JWT token"""
    payload = {
        "user_id": user_id,
        "email": email,
        "exp": datetime.now(timezone.utc) + timedelta(days=JWT_EXPIRY_DAYS),
        "iat": datetime.now(timezone.utc)
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

def decode_jwt_token(token: str) -> Optional[dict]:
    """Decode and verify JWT token"""
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        return payload
    except jwt.ExpiredSignatureError:
        return None
    except jwt.InvalidTokenError:
        return None

async def verify_jwt_token(request: Request) -> dict:
    """Verify JWT token from Authorization header"""
    auth_header = request.headers.get("Authorization")
    if not auth_header or not auth_header.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Token fehlt")
    
    token = auth_header.split(" ")[1]
    payload = decode_jwt_token(token)
    if not payload:
        raise HTTPException(status_code=401, detail="Ungültiger Token")
    
    # Get user data
    user = await db.users.find_one({"user_id": payload["user_id"]}, {"_id": 0})
    if not user:
        raise HTTPException(status_code=401, detail="User nicht gefunden")
    
    return user

def validate_user_document(user: dict) -> tuple[bool, str]:
    """
    Validiert ein User-Dokument auf Vollständigkeit.
    Returns: (is_valid, error_message)
    """
    required_fields = ['user_id', 'email', 'password_hash', 'plan']
    
    for field in required_fields:
        if not user.get(field):
            return False, f"Feld '{field}' fehlt oder ist leer"
    
    # Validate email format
    if not re.match(r"[^@]+@[^@]+\.[^@]+", user.get('email', '')):
        return False, "Ungültiges E-Mail-Format"
    
    # Validate plan
    if user.get('plan') not in ['basic', 'pro']:
        return False, f"Ungültiger Plan: {user.get('plan')}"
    
    return True, ""

async def repair_user_if_needed(user: dict) -> dict:
    """
    Repariert einen User, falls Felder fehlen.
    Wird bei Login aufgerufen um alte User zu migrieren.
    """
    updates = {}
    
    # Fix missing user_id
    if not user.get('user_id'):
        updates['user_id'] = f"user_{uuid.uuid4().hex[:12]}"
    
    # Fix missing plan
    if not user.get('plan'):
        updates['plan'] = 'basic'
    
    # Fix missing usage
    if not user.get('usage'):
        updates['usage'] = {
            "pairing_requests_today": 0,
            "chat_messages_today": 0,
            "last_usage_date": None
        }
    
    # Apply updates if needed
    if updates:
        await db.users.update_one(
            {"email": user['email']},
            {"$set": updates}
        )
        logger.info(f"🔧 User '{user['email']}' repariert: {list(updates.keys())}")
        user.update(updates)
    
    return user

@api_router.post("/auth/register")
async def register_user(req: RegisterRequest, response: Response):
    """
    Register a new user with email and password.
    
    ROBUSTER ABLAUF FÜR TAUSENDE USER:
    1. Email-Validierung (Format + Lowercase)
    2. Passwort-Validierung (min 6 Zeichen)
    3. Duplikat-Check (Email bereits registriert?)
    4. User-ID Generierung (eindeutig, kurz, URL-safe)
    5. Password-Hashing (bcrypt, sicher)
    6. Vollständiges User-Dokument erstellen
    7. Validierung des Dokuments vor Insert
    8. Datenbank-Insert
    9. JWT Token erstellen
    10. Session Cookie setzen
    """
    # 1. Email-Validierung
    email = req.email.lower().strip()
    if not re.match(r"[^@]+@[^@]+\.[^@]+", email):
        raise HTTPException(status_code=400, detail="Ungültige E-Mail-Adresse")
    
    # 2. Passwort-Validierung
    if len(req.password) < 6:
        raise HTTPException(status_code=400, detail="Passwort muss mindestens 6 Zeichen haben")
    
    # 3. Duplikat-Check
    existing = await db.users.find_one({"email": email})
    if existing:
        raise HTTPException(status_code=400, detail="Diese E-Mail ist bereits registriert")
    
    # 4. User-ID generieren (eindeutig)
    user_id = f"user_{uuid.uuid4().hex[:12]}"
    
    # 5. Password hashen
    hashed_pw = hash_password(req.password)
    
    # 6. Vollständiges User-Dokument
    new_user = {
        "user_id": user_id,
        "email": email,
        "name": req.name.strip() if req.name else email.split('@')[0],
        "password_hash": hashed_pw,
        "picture": None,
        "plan": "basic",
        "subscription_id": None,
        "subscription_status": None,
        "subscription_end_date": None,
        "stripe_customer_id": None,
        "usage": {
            "pairing_requests_today": 0,
            "chat_messages_today": 0,
            "last_usage_date": None
        },
        "created_at": datetime.now(timezone.utc),
        "last_login": datetime.now(timezone.utc),
        "login_count": 1
    }
    
    # 7. Validierung vor Insert
    is_valid, error = validate_user_document(new_user)
    if not is_valid:
        logger.error(f"❌ User-Validierung fehlgeschlagen: {error}")
        raise HTTPException(status_code=500, detail="Registrierung fehlgeschlagen - bitte erneut versuchen")
    
    # 8. Datenbank-Insert
    try:
        await db.users.insert_one(new_user)
        logger.info(f"✅ Neuer User registriert: {email} (ID: {user_id})")
    except Exception as e:
        logger.error(f"❌ Datenbank-Fehler bei Registrierung: {e}")
        raise HTTPException(status_code=500, detail="Registrierung fehlgeschlagen - bitte erneut versuchen")
    
    # 9. JWT Token erstellen
    token = create_jwt_token(user_id, email)
    
    # 10. Session Cookie setzen
    response.set_cookie(
        key="session_token",
        value=token,
        httponly=True,
        secure=True,
        samesite="none",
        path="/",
        max_age=JWT_EXPIRY_DAYS * 24 * 60 * 60
    )
    
    return {
        "user_id": user_id,
        "email": email,
        "name": new_user["name"],
        "plan": "basic",
        "token": token,  # Token für localStorage (Safari/iOS)
        "message": "Registrierung erfolgreich!"
    }

@api_router.post("/auth/login")
async def login_user(req: LoginRequest, response: Response):
    """
    Login with email and password.
    
    ROBUSTER ABLAUF:
    1. Email normalisieren (lowercase, trim)
    2. User in DB suchen
    3. Passwort verifizieren
    4. User reparieren falls Felder fehlen (Migration alter User)
    5. Login-Statistik aktualisieren
    6. JWT Token erstellen
    7. Session Cookie setzen
    """
    # 1. Email normalisieren
    email = req.email.lower().strip()
    
    # 2. User suchen
    user = await db.users.find_one({"email": email})
    
    if not user:
        raise HTTPException(status_code=401, detail="E-Mail oder Passwort falsch")
    
    # 3. Passwort verifizieren
    password_hash = user.get("password_hash", "")
    if not password_hash or not verify_password(req.password, password_hash):
        raise HTTPException(status_code=401, detail="E-Mail oder Passwort falsch")
    
    # 4. User reparieren falls nötig (Migration alter User)
    user = await repair_user_if_needed(user)
    
    # 5. Login-Statistik aktualisieren
    await db.users.update_one(
        {"email": email},
        {
            "$set": {"last_login": datetime.now(timezone.utc)},
            "$inc": {"login_count": 1}
        }
    )
    
    # 6. JWT Token erstellen
    token = create_jwt_token(user["user_id"], user["email"])
    
    # 7. Session Cookie setzen (für Browser die Cookies unterstützen)
    response.set_cookie(
        key="session_token",
        value=token,
        httponly=True,
        secure=True,
        samesite="none",
        path="/",
        max_age=JWT_EXPIRY_DAYS * 24 * 60 * 60
    )
    
    # Return user (ohne Passwort) + Token für localStorage (Safari/iOS-kompatibel)
    return {
        "user_id": user["user_id"],
        "email": user["email"],
        "name": user.get("name", email.split('@')[0]),
        "picture": user.get("picture"),
        "plan": user.get("plan", "basic"),
        "usage": user.get("usage", {}),
        "token": token,  # Token für localStorage (Safari/iOS)
        "message": "Anmeldung erfolgreich!"
    }

@api_router.get("/auth/me")
async def get_current_user_endpoint(request: Request):
    """Get current authenticated user"""
    user = await get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    user = await reset_daily_usage_if_needed(user)
    return user.model_dump()

@api_router.post("/auth/logout")
async def logout(request: Request, response: Response):
    """Logout user and clear session"""
    user = await get_current_user(request)
    if user:
        await db.user_sessions.delete_many({"user_id": user.user_id})
    
    response.delete_cookie("session_token", path="/")
    # Auch localStorage-Token löschen (Frontend kümmert sich darum)
    return {"message": "Logged out"}


# ===================== WINE PROFILE API =====================

class WineProfileUpdate(BaseModel):
    """Request model for updating wine profile"""
    red_wine_style: Optional[str] = None
    white_wine_style: Optional[str] = None
    acidity_tolerance: Optional[str] = None
    tannin_preference: Optional[str] = None
    sweetness_preference: Optional[str] = None
    favorite_regions: Optional[List[str]] = None
    budget_everyday: Optional[str] = None
    budget_restaurant: Optional[str] = None
    no_gos: Optional[List[str]] = None
    dietary_preferences: Optional[List[str]] = None
    adventure_level: Optional[str] = None

@api_router.get("/profile/wine")
async def get_wine_profile(request: Request):
    """Get user's wine taste profile (Pro feature)"""
    user = await get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Nicht angemeldet")
    
    # Check if user is Pro
    if user.plan != "pro":
        raise HTTPException(status_code=403, detail="Diese Funktion ist nur für Pro-Mitglieder verfügbar")
    
    # Get profile from database
    profile = await db.wine_profiles.find_one({"user_id": user.user_id}, {"_id": 0})
    
    if not profile:
        # Return default empty profile
        return {
            "user_id": user.user_id,
            "red_wine_style": None,
            "white_wine_style": None,
            "acidity_tolerance": None,
            "tannin_preference": None,
            "sweetness_preference": None,
            "favorite_regions": [],
            "budget_everyday": None,
            "budget_restaurant": None,
            "no_gos": [],
            "dietary_preferences": [],
            "adventure_level": None,
            "updated_at": None
        }
    
    return profile

@api_router.put("/profile/wine")
async def update_wine_profile(request: Request, profile_update: WineProfileUpdate):
    """Update user's wine taste profile (Pro feature)"""
    user = await get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Nicht angemeldet")
    
    # Check if user is Pro
    if user.plan != "pro":
        raise HTTPException(status_code=403, detail="Diese Funktion ist nur für Pro-Mitglieder verfügbar")
    
    # Prepare update data
    update_data = {
        "user_id": user.user_id,
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    
    # Only update fields that are provided
    for field, value in profile_update.model_dump().items():
        if value is not None:
            update_data[field] = value
    
    # Upsert profile
    await db.wine_profiles.update_one(
        {"user_id": user.user_id},
        {"$set": update_data},
        upsert=True
    )
    
    # Return updated profile
    profile = await db.wine_profiles.find_one({"user_id": user.user_id}, {"_id": 0})
    logger.info(f"🍷 Wine profile updated for user {user.user_id}")
    
    return {
        "success": True,
        "message": "Weinprofil erfolgreich aktualisiert",
        "profile": profile
    }

@api_router.delete("/profile/wine")
async def reset_wine_profile(request: Request):
    """Reset user's wine taste profile to defaults (Pro feature)"""
    user = await get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Nicht angemeldet")
    
    if user.plan != "pro":
        raise HTTPException(status_code=403, detail="Diese Funktion ist nur für Pro-Mitglieder verfügbar")
    
    await db.wine_profiles.delete_one({"user_id": user.user_id})
    logger.info(f"🍷 Wine profile reset for user {user.user_id}")
    
    return {"success": True, "message": "Weinprofil zurückgesetzt"}


def get_profile_context_for_ai(profile: dict, language: str = "de") -> str:
    """Generate AI context from user's wine profile"""
    if not profile:
        return ""
    
    context_parts = []
    
    # Red wine style
    if profile.get("red_wine_style"):
        styles = {
            "kraftig_wurzig": {"de": "kräftige, würzige Rotweine (Bordeaux, Rhône-Stil)", "en": "bold, spicy red wines (Bordeaux, Rhône style)"},
            "fruchtig_elegant": {"de": "fruchtige, elegante Rotweine (Burgunder-Stil)", "en": "fruity, elegant red wines (Burgundy style)"},
            "beides": {"de": "sowohl kräftige als auch elegante Rotweine", "en": "both bold and elegant red wines"}
        }
        if profile["red_wine_style"] in styles:
            context_parts.append(f"Bevorzugt bei Rotwein: {styles[profile['red_wine_style']].get(language, styles[profile['red_wine_style']]['de'])}")
    
    # White wine style
    if profile.get("white_wine_style"):
        styles = {
            "mineralisch_frisch": {"de": "mineralische, frische Weißweine (Chablis-Stil)", "en": "mineral, fresh white wines (Chablis style)"},
            "cremig_textur": {"de": "cremige Weißweine mit Textur (Meursault-Stil)", "en": "creamy white wines with texture (Meursault style)"},
            "aromatisch_verspielt": {"de": "aromatische, verspielte Weißweine (Riesling-Stil)", "en": "aromatic, playful white wines (Riesling style)"},
            "beides": {"de": "verschiedene Weißwein-Stile", "en": "various white wine styles"}
        }
        if profile["white_wine_style"] in styles:
            context_parts.append(f"Bevorzugt bei Weißwein: {styles[profile['white_wine_style']].get(language, styles[profile['white_wine_style']]['de'])}")
    
    # Acidity tolerance
    if profile.get("acidity_tolerance"):
        tolerances = {
            "niedrig": {"de": "Mag keine säurebetonten Weine", "en": "Doesn't like high-acid wines"},
            "mittel": {"de": "Moderate Säure ist akzeptabel", "en": "Moderate acidity is acceptable"},
            "hoch": {"de": "Liebt säurebetonte, frische Weine", "en": "Loves high-acid, fresh wines"}
        }
        if profile["acidity_tolerance"] in tolerances:
            context_parts.append(tolerances[profile["acidity_tolerance"]].get(language, tolerances[profile["acidity_tolerance"]]["de"]))
    
    # Tannin preference
    if profile.get("tannin_preference"):
        prefs = {
            "weich_seidig": {"de": "Bevorzugt weiche, seidige Tannine – keine harten Gerbstoffe", "en": "Prefers soft, silky tannins – no harsh tannins"},
            "mittel": {"de": "Mittlere Tannine sind akzeptabel", "en": "Medium tannins are acceptable"},
            "markant_griffig": {"de": "Mag markante, griffige Tannine", "en": "Likes bold, grippy tannins"}
        }
        if profile["tannin_preference"] in prefs:
            context_parts.append(prefs[profile["tannin_preference"]].get(language, prefs[profile["tannin_preference"]]["de"]))
    
    # Sweetness
    if profile.get("sweetness_preference"):
        prefs = {
            "knochentrocken": {"de": "Nur knochentrockene Weine", "en": "Only bone-dry wines"},
            "trocken": {"de": "Bevorzugt trockene Weine", "en": "Prefers dry wines"},
            "halbtrocken": {"de": "Halbtrockene Weine sind willkommen", "en": "Off-dry wines are welcome"},
            "lieblich": {"de": "Mag auch liebliche Weine", "en": "Also likes sweet wines"},
            "edelsuss": {"de": "Liebt edelsüße Weine", "en": "Loves noble sweet wines"}
        }
        if profile["sweetness_preference"] in prefs:
            context_parts.append(prefs[profile["sweetness_preference"]].get(language, prefs[profile["sweetness_preference"]]["de"]))
    
    # Favorite regions
    if profile.get("favorite_regions") and len(profile["favorite_regions"]) > 0:
        regions = ", ".join(profile["favorite_regions"])
        context_parts.append(f"Lieblingsregionen: {regions}")
    
    # Budget
    if profile.get("budget_everyday"):
        budgets = {
            "unter_10": "unter 10€",
            "10_20": "10-20€",
            "20_35": "20-35€",
            "35_50": "35-50€",
            "ueber_50": "über 50€"
        }
        if profile["budget_everyday"] in budgets:
            context_parts.append(f"Budget für Alltag: {budgets[profile['budget_everyday']]}")
    
    if profile.get("budget_restaurant"):
        budgets = {
            "unter_30": "unter 30€",
            "30_50": "30-50€",
            "50_80": "50-80€",
            "80_120": "80-120€",
            "ueber_120": "über 120€"
        }
        if profile["budget_restaurant"] in budgets:
            context_parts.append(f"Budget im Restaurant: {budgets[profile['budget_restaurant']]}")
    
    # No-gos
    if profile.get("no_gos") and len(profile["no_gos"]) > 0:
        no_gos = ", ".join(profile["no_gos"])
        context_parts.append(f"⚠️ WICHTIG - KEINE Empfehlungen für: {no_gos}")
    
    # Dietary preferences
    if profile.get("dietary_preferences") and len(profile["dietary_preferences"]) > 0:
        prefs = ", ".join(profile["dietary_preferences"])
        context_parts.append(f"Kulinarische Vorlieben: {prefs}")
    
    # Adventure level
    if profile.get("adventure_level"):
        levels = {
            "klassiker": {"de": "Bevorzugt klassische, bekannte Weine – keine exotischen Empfehlungen", "en": "Prefers classic, well-known wines – no exotic recommendations"},
            "ausgewogen": {"de": "Mix aus Klassikern und neuen Entdeckungen", "en": "Mix of classics and new discoveries"},
            "abenteuerlich": {"de": "Liebt mutige Wildcard-Empfehlungen und unbekannte Regionen", "en": "Loves bold wildcard recommendations and unknown regions"}
        }
        if profile["adventure_level"] in levels:
            context_parts.append(levels[profile["adventure_level"]].get(language, levels[profile["adventure_level"]]["de"]))
    
    if context_parts:
        return "\n\n🍷 BENUTZERPROFIL (Bitte bei der Empfehlung berücksichtigen):\n" + "\n".join(f"• {part}" for part in context_parts)
    
    return ""


# ===================== GOOGLE OAUTH =====================

class GoogleSessionRequest(BaseModel):
    session_id: str

@api_router.post("/auth/google/session")
async def process_google_session(req: GoogleSessionRequest, response: Response):
    """
    Verarbeitet Google OAuth Session von Emergent Auth.
    
    Flow:
    1. Frontend erhält session_id von auth.emergentagent.com
    2. Frontend sendet session_id an diesen Endpoint
    3. Backend tauscht session_id gegen User-Daten
    4. Backend erstellt/aktualisiert User in DB
    5. Backend setzt Session-Cookie + gibt Token zurück
    """
    import httpx
    
    try:
        # Session-Daten von Emergent Auth abrufen
        async with httpx.AsyncClient() as client:
            auth_response = await client.get(
                "https://demobackend.emergentagent.com/auth/v1/env/oauth/session-data",
                headers={"X-Session-ID": req.session_id},
                timeout=10.0
            )
            
            if auth_response.status_code != 200:
                logger.error(f"Google Auth failed: {auth_response.status_code}")
                raise HTTPException(status_code=401, detail="Google-Anmeldung fehlgeschlagen")
            
            auth_data = auth_response.json()
    except httpx.RequestError as e:
        logger.error(f"Google Auth request error: {e}")
        raise HTTPException(status_code=500, detail="Verbindungsfehler bei Google-Anmeldung")
    
    # User-Daten extrahieren
    google_id = auth_data.get("id")
    email = auth_data.get("email", "").lower().strip()
    name = auth_data.get("name", "")
    picture = auth_data.get("picture")
    emergent_session_token = auth_data.get("session_token")
    
    if not email:
        raise HTTPException(status_code=400, detail="Keine E-Mail von Google erhalten")
    
    # Prüfen ob User existiert (per Email)
    existing_user = await db.users.find_one({"email": email}, {"_id": 0})
    
    if existing_user:
        # User existiert - aktualisieren
        user_id = existing_user.get("user_id")
        await db.users.update_one(
            {"email": email},
            {
                "$set": {
                    "google_id": google_id,
                    "picture": picture or existing_user.get("picture"),
                    "name": name or existing_user.get("name"),
                    "last_login": datetime.now(timezone.utc),
                    "auth_provider": "google"
                },
                "$inc": {"login_count": 1}
            }
        )
        plan = existing_user.get("plan", "basic")
        usage = existing_user.get("usage", {})
        logger.info(f"Google login: existing user {email}")
    else:
        # Neuer User erstellen
        user_id = f"user_{uuid.uuid4().hex[:12]}"
        new_user = {
            "user_id": user_id,
            "email": email,
            "name": name or email.split('@')[0],
            "picture": picture,
            "google_id": google_id,
            "auth_provider": "google",
            "plan": "basic",
            "subscription_status": "active",
            "usage": {
                "pairing_requests_today": 0,
                "chat_messages_today": 0,
                "last_usage_date": datetime.now(timezone.utc).date().isoformat()
            },
            "created_at": datetime.now(timezone.utc),
            "last_login": datetime.now(timezone.utc),
            "login_count": 1
        }
        await db.users.insert_one(new_user)
        plan = "basic"
        usage = new_user["usage"]
        logger.info(f"Google login: new user created {email}")
    
    # JWT Token erstellen
    token = create_jwt_token(user_id, email)
    
    # Session-Cookie setzen
    response.set_cookie(
        key="session_token",
        value=token,
        httponly=True,
        secure=True,
        samesite="none",
        path="/",
        max_age=JWT_EXPIRY_DAYS * 24 * 60 * 60
    )
    
    # Emergent Session speichern (optional, für spätere Verwendung)
    if emergent_session_token:
        await db.user_sessions.update_one(
            {"user_id": user_id},
            {
                "$set": {
                    "session_token": emergent_session_token,
                    "expires_at": datetime.now(timezone.utc) + timedelta(days=7),
                    "updated_at": datetime.now(timezone.utc)
                }
            },
            upsert=True
        )
    
    return {
        "user_id": user_id,
        "email": email,
        "name": name,
        "picture": picture,
        "plan": plan,
        "usage": usage,
        "token": token,  # Für localStorage (Safari/iOS)
        "auth_provider": "google",
        "message": "Google-Anmeldung erfolgreich!"
    }


# ===================== PASSWORD RESET =====================

class PasswordResetRequest(BaseModel):
    email: str

class PasswordResetConfirm(BaseModel):
    token: str
    new_password: str

@api_router.post("/auth/forgot-password")
async def forgot_password(req: PasswordResetRequest):
    """
    Passwort vergessen - sendet Reset-Email.
    
    Flow:
    1. User gibt Email ein
    2. System generiert Reset-Token (gültig 1 Stunde)
    3. Email mit Reset-Link wird gesendet
    4. User klickt Link und setzt neues Passwort
    """
    email = req.email.lower().strip()
    
    # Find user
    user = await db.users.find_one({"email": email})
    
    # WICHTIG: Immer gleiche Antwort geben (Security - verhindert Email-Enumeration)
    success_message = {
        "message": "Falls ein Account mit dieser E-Mail existiert, wurde ein Reset-Link gesendet.",
        "message_en": "If an account with this email exists, a reset link has been sent."
    }
    
    if not user:
        return success_message
    
    # Generate secure reset token
    reset_token = secrets.token_urlsafe(32)
    reset_expiry = datetime.now(timezone.utc) + timedelta(hours=1)
    
    # Store token in database
    await db.users.update_one(
        {"email": email},
        {"$set": {
            "password_reset_token": reset_token,
            "password_reset_expiry": reset_expiry
        }}
    )
    
    # Build reset URL
    reset_url = f"{FRONTEND_URL}/reset-password?token={reset_token}"
    
    # Send email via Resend
    if RESEND_API_KEY:
        try:
            resend.Emails.send({
                "from": f"Wine Pairing <{SENDER_EMAIL}>",
                "to": [email],
                "subject": "🍷 Passwort zurücksetzen - Wine Pairing",
                "html": f"""
                <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto; padding: 20px;">
                    <h1 style="color: #722F37;">🍷 Passwort zurücksetzen</h1>
                    <p>Hallo,</p>
                    <p>Sie haben eine Anfrage zum Zurücksetzen Ihres Passworts gestellt.</p>
                    <p>Klicken Sie auf den folgenden Button, um ein neues Passwort zu setzen:</p>
                    <p style="margin: 30px 0;">
                        <a href="{reset_url}" 
                           style="background-color: #722F37; color: white; padding: 15px 30px; 
                                  text-decoration: none; border-radius: 5px; font-weight: bold;">
                            Neues Passwort setzen
                        </a>
                    </p>
                    <p style="color: #666; font-size: 14px;">
                        Dieser Link ist 1 Stunde gültig.<br>
                        Falls Sie diese Anfrage nicht gestellt haben, ignorieren Sie diese E-Mail.
                    </p>
                    <hr style="border: none; border-top: 1px solid #eee; margin: 30px 0;">
                    <p style="color: #999; font-size: 12px;">
                        Wine Pairing - Ihr persönlicher Wein-Sommelier<br>
                        <a href="https://wine-pairing.online" style="color: #722F37;">wine-pairing.online</a>
                    </p>
                </div>
                """
            })
            logger.info(f"✅ Password reset email sent to {email}")
        except Exception as e:
            logger.error(f"❌ Failed to send reset email: {e}")
            # Still return success (don't reveal email exists)
    
    return success_message

@api_router.post("/auth/reset-password")
async def reset_password(req: PasswordResetConfirm):
    """
    Setzt das Passwort mit dem Reset-Token zurück.
    """
    # Validate password
    if len(req.new_password) < 6:
        raise HTTPException(status_code=400, detail="Passwort muss mindestens 6 Zeichen haben")
    
    # Find user with valid token
    user = await db.users.find_one({
        "password_reset_token": req.token,
        "password_reset_expiry": {"$gt": datetime.now(timezone.utc)}
    })
    
    if not user:
        raise HTTPException(
            status_code=400, 
            detail="Ungültiger oder abgelaufener Reset-Link. Bitte fordern Sie einen neuen an."
        )
    
    # Hash new password
    new_hash = hash_password(req.new_password)
    
    # Update password and remove reset token
    await db.users.update_one(
        {"_id": user["_id"]},
        {
            "$set": {"password_hash": new_hash},
            "$unset": {"password_reset_token": "", "password_reset_expiry": ""}
        }
    )
    
    logger.info(f"✅ Password reset successful for {user.get('email')}")
    
    return {
        "message": "Passwort erfolgreich geändert! Sie können sich jetzt einloggen.",
        "message_en": "Password changed successfully! You can now log in."
    }

@api_router.get("/auth/verify-reset-token/{token}")
async def verify_reset_token(token: str):
    """
    Prüft ob ein Reset-Token gültig ist.
    Wird vom Frontend aufgerufen bevor das Reset-Formular angezeigt wird.
    """
    user = await db.users.find_one({
        "password_reset_token": token,
        "password_reset_expiry": {"$gt": datetime.now(timezone.utc)}
    })
    
    if not user:
        raise HTTPException(status_code=400, detail="Ungültiger oder abgelaufener Link")
    
    return {"valid": True, "email": user.get("email", "")[:3] + "***"}

# Subscription endpoints
@api_router.post("/subscription/checkout")
async def create_checkout_session(checkout_req: CheckoutRequest, request: Request):
    """Create Stripe checkout session for subscription"""
    user = await get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Login required")
    
    if checkout_req.plan not in SUBSCRIPTION_PLANS:
        raise HTTPException(status_code=400, detail="Invalid plan")
    
    plan = SUBSCRIPTION_PLANS[checkout_req.plan]
    
    # Create success/cancel URLs
    success_url = f"{checkout_req.origin_url}/subscription/success?session_id={{CHECKOUT_SESSION_ID}}"
    cancel_url = f"{checkout_req.origin_url}/subscription/cancel"
    
    # Initialize Stripe
    host_url = str(request.base_url)
    webhook_url = f"{host_url}api/webhook/stripe"
    stripe_checkout = StripeCheckout(api_key=STRIPE_API_KEY, webhook_url=webhook_url)
    
    # Create checkout session
    checkout_request = CheckoutSessionRequest(
        amount=plan["price"],
        currency=plan["currency"],
        success_url=success_url,
        cancel_url=cancel_url,
        metadata={
            "user_id": user.user_id,
            "email": user.email,
            "plan": checkout_req.plan
        }
    )
    
    session = await stripe_checkout.create_checkout_session(checkout_request)
    
    # Create payment transaction record
    await db.payment_transactions.insert_one({
        "transaction_id": f"txn_{uuid.uuid4().hex[:12]}",
        "user_id": user.user_id,
        "email": user.email,
        "session_id": session.session_id,
        "plan": checkout_req.plan,
        "amount": plan["price"],
        "currency": plan["currency"],
        "payment_status": "pending",
        "created_at": datetime.now(timezone.utc)
    })
    
    return {"url": session.url, "session_id": session.session_id}

@api_router.get("/subscription/status/{session_id}")
async def get_subscription_status(session_id: str, request: Request):
    """Check subscription payment status"""
    user = await get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Login required")
    
    # Find transaction
    transaction = await db.payment_transactions.find_one(
        {"session_id": session_id, "user_id": user.user_id},
        {"_id": 0}
    )
    
    if not transaction:
        raise HTTPException(status_code=404, detail="Transaction not found")
    
    # If already processed, return cached status
    if transaction["payment_status"] in ["paid", "failed", "expired"]:
        return transaction
    
    # Check with Stripe
    host_url = str(request.base_url)
    webhook_url = f"{host_url}api/webhook/stripe"
    stripe_checkout = StripeCheckout(api_key=STRIPE_API_KEY, webhook_url=webhook_url)
    
    status = await stripe_checkout.get_checkout_status(session_id)
    
    # Update transaction
    new_status = status.payment_status
    await db.payment_transactions.update_one(
        {"session_id": session_id},
        {"$set": {"payment_status": new_status}}
    )
    
    # If paid, upgrade user
    if new_status == "paid":
        plan = transaction["plan"]
        if "yearly" in plan:
            end_date = datetime.now(timezone.utc) + timedelta(days=365)
        else:
            end_date = datetime.now(timezone.utc) + timedelta(days=30)
        
        await db.users.update_one(
            {"user_id": user.user_id},
            {"$set": {
                "plan": "pro",
                "subscription_id": session_id,
                "subscription_status": "active",
                "subscription_end_date": end_date
            }}
        )
    
    transaction["payment_status"] = new_status
    return transaction

@api_router.get("/subscription/plans")
async def get_subscription_plans():
    """Get available subscription plans"""
    return {
        "plans": SUBSCRIPTION_PLANS,
        "limits": {
            "basic": FREEMIUM_LIMITS["basic"],
            "pro": {k: "unlimited" if v == float('inf') else v for k, v in FREEMIUM_LIMITS["pro"].items()}
        }
    }

@api_router.post("/webhook/stripe")
async def stripe_webhook(request: Request):
    """Handle Stripe webhooks"""
    body = await request.body()
    signature = request.headers.get("Stripe-Signature")
    
    try:
        host_url = str(request.base_url)
        webhook_url = f"{host_url}api/webhook/stripe"
        stripe_checkout = StripeCheckout(api_key=STRIPE_API_KEY, webhook_url=webhook_url)
        
        webhook_response = await stripe_checkout.handle_webhook(body, signature)
        
        if webhook_response.payment_status == "paid":
            session_id = webhook_response.session_id
            metadata = webhook_response.metadata
            
            # Update transaction
            await db.payment_transactions.update_one(
                {"session_id": session_id},
                {"$set": {"payment_status": "paid"}}
            )
            
            # Upgrade user
            user_id = metadata.get("user_id")
            plan = metadata.get("plan", "pro_monthly")
            
            if "yearly" in plan:
                end_date = datetime.now(timezone.utc) + timedelta(days=365)
            else:
                end_date = datetime.now(timezone.utc) + timedelta(days=30)
            
            await db.users.update_one(
                {"user_id": user_id},
                {"$set": {
                    "plan": "pro",
                    "subscription_id": session_id,
                    "subscription_status": "active",
                    "subscription_end_date": end_date
                }}
            )
        
        return {"status": "ok"}
    except Exception as e:
        logger.error(f"Webhook error: {e}")
        return {"status": "error", "message": str(e)}


# ===================== BACKUP SYSTEM ENDPOINTS =====================

@api_router.get("/backup/status")
async def get_backup_status():
    """
    Gibt den aktuellen Backup-Status zurück.
    Zeigt alle verfügbaren Backups, aktuelle Daten-Counts und Auto-Backup-Status.
    """
    global backup_manager
    if not backup_manager:
        raise HTTPException(status_code=503, detail="Backup-Manager nicht initialisiert")
    
    status = await backup_manager.get_backup_status()
    
    # Füge Auto-Backup-Informationen hinzu
    status['auto_backup'] = {
        'enabled': getattr(backup_manager, '_auto_backup_running', False),
        'interval_hours': 6,
        'next_backup': await backup_manager.get_next_backup_time()
    }
    
    return status


@api_router.post("/backup/create")
async def create_backup(background_tasks: BackgroundTasks, user_data_only: bool = False):
    """
    Erstellt ein neues Backup.
    
    Args:
        user_data_only: Wenn True, werden nur User-Daten gesichert (schneller)
    """
    global backup_manager
    if not backup_manager:
        raise HTTPException(status_code=503, detail="Backup-Manager nicht initialisiert")
    
    try:
        if user_data_only:
            result = await backup_manager.backup_user_data_only()
        else:
            result = await backup_manager.create_full_backup()
        
        return {
            "success": True,
            "message": "Backup erfolgreich erstellt",
            "backup_dir": result.get('backup_dir'),
            "timestamp": result.get('timestamp'),
            "collections": result.get('collections', {})
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Backup fehlgeschlagen: {str(e)}")


@api_router.get("/backup/user-data-counts")
async def get_user_data_counts():
    """
    Gibt die aktuellen Counts aller User-Daten Collections zurück.
    Nützlich für Quick-Health-Checks.
    """
    counts = {}
    user_collections = ['users', 'wines', 'pairings', 'chats', 'wine_favorites', 'payment_transactions']
    
    for col in user_collections:
        counts[col] = await db[col].count_documents({})
    
    return {
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "user_data_counts": counts,
        "total_user_documents": sum(counts.values())
    }


# ===================== COUPON SYSTEM MODELS =====================

class CouponRequest(BaseModel):
    code: str

class CouponResponse(BaseModel):
    success: bool
    message: str
    plan_upgraded_to: Optional[str] = None
    expires_at: Optional[datetime] = None

# ===================== STRIPE ENDPOINTS =====================

@api_router.get("/subscription-plans")
async def get_subscription_plans():
    """Get available subscription plans"""
    return SUBSCRIPTION_PLANS

@api_router.post("/create-checkout-session")
async def create_checkout_session(
    request: dict,
    current_user: dict = Depends(verify_jwt_token)
):
    """Create Stripe checkout session"""
    try:
        plan = request.get("plan")
        if not plan or plan not in SUBSCRIPTION_PLANS:
            raise HTTPException(status_code=400, detail="Invalid plan")
            
        plan_info = SUBSCRIPTION_PLANS[plan]
        
        # Create Stripe checkout session
        stripe_checkout = StripeCheckout(api_key=STRIPE_API_KEY)
        
        frontend_url = os.environ.get('FRONTEND_URL', 'https://wine-pairing.online')
        checkout_request = CheckoutSessionRequest(
            mode="subscription",
            success_url=f"{frontend_url}/subscription/success",
            cancel_url=f"{frontend_url}/subscription/cancel",
            line_items=[{
                "price_data": {
                    "currency": plan_info["currency"],
                    "product_data": {
                        "name": f"Wine Pairing Pro - {plan_info['interval'].title()}"
                    },
                    "unit_amount": int(plan_info["price"] * 100),
                    "recurring": {
                        "interval": plan_info["interval"]
                    }
                },
                "quantity": 1
            }],
            customer_email=current_user["email"]
        )
        
        response = await stripe_checkout.create_checkout_session(checkout_request)
        return {"checkout_url": response.url}
        
    except Exception as e:
        print(f"Stripe error: {e}")
        raise HTTPException(status_code=500, detail=f"Stripe error: {str(e)}")

@api_router.post("/coupon/redeem", response_model=CouponResponse)
async def redeem_coupon(
    coupon_request: CouponRequest,
    request: Request
):
    """Redeem a coupon code"""
    try:
        # Get current user from session cookie
        current_user = await get_current_user(request)
        if not current_user:
            raise HTTPException(status_code=401, detail="Anmeldung erforderlich")
            
        coupon_code = coupon_request.code.upper().strip()
        
        # Find coupon in database
        coupon = await db.coupons.find_one({"code": coupon_code}, {"_id": 0})
        
        if not coupon:
            return CouponResponse(
                success=False,
                message="Gutschein-Code nicht gefunden"
            )
        
        if coupon.get("used", False):
            return CouponResponse(
                success=False,
                message="Gutschein-Code bereits verwendet"
            )
        
        # Calculate expiry date
        expires_at = datetime.now(timezone.utc) + timedelta(days=365)  # 1 Jahr
        
        # Update user plan
        await db.users.update_one(
            {"user_id": current_user.user_id},
            {
                "$set": {
                    "plan": "pro",
                    "plan_expires_at": expires_at.isoformat(),
                    "upgraded_via": "coupon",
                    "coupon_code": coupon_code
                }
            }
        )
        
        # Mark coupon as used
        await db.coupons.update_one(
            {"code": coupon_code},
            {
                "$set": {
                    "used": True,
                    "used_by": current_user.email,
                    "used_at": datetime.now(timezone.utc).isoformat()
                }
            }
        )
        
        return CouponResponse(
            success=True,
            message="Gutschein erfolgreich eingelöst! Sie haben jetzt 1 Jahr kostenlosen Pro-Zugang.",
            plan_upgraded_to="pro",
            expires_at=expires_at
        )
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"Coupon redemption error: {e}")
        raise HTTPException(status_code=500, detail=f"Fehler beim Einlösen: {str(e)}")

@api_router.get("/coupon/stats")
async def get_coupon_stats():
    """Get coupon statistics (admin only)"""
    try:
        total_coupons = await db.coupons.count_documents({})
        used_coupons = await db.coupons.count_documents({"used": True})
        unused_coupons = total_coupons - used_coupons
        
        return {
            "total": total_coupons,
            "used": used_coupons,
            "unused": unused_coupons,
            "usage_rate": round((used_coupons / total_coupons * 100), 2) if total_coupons > 0 else 0
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Include the router AFTER middleware
app.include_router(api_router)

# Mount exports directory for database backups
exports_path = Path("/app/exports")
if exports_path.exists():
    app.mount("/api/exports", StaticFiles(directory=str(exports_path)), name="exports")

@app.on_event("startup")
async def startup_seed_data():
    """
    ROBUSTE DATENBANK-INITIALISIERUNG
    Prüft gegen das Backup-Manifest und stellt sicher, dass ALLE Daten korrekt sind.
    Version 3.2 - Mit Datenbank-Indizes für Skalierbarkeit
    """
    global backup_manager
    
    print("\n" + "=" * 60)
    print("🚀 WINE-PAIRING.ONLINE - SERVER STARTUP")
    print("=" * 60)
    
    # ===================================================================
    # DATENBANK-INDIZES FÜR SCHNELLE ABFRAGEN BEI VIELEN USERN
    # Wird bei jedem Start geprüft (idempotent - erstellt nur wenn nicht existiert)
    # ===================================================================
    print("\n📊 ERSTELLE DATENBANK-INDIZES...")
    try:
        # Users Collection - kritisch für Login-Performance
        await db.users.create_index("email", unique=True, background=True)
        await db.users.create_index("user_id", unique=True, background=True)
        await db.users.create_index("stripe_customer_id", sparse=True, background=True)
        
        # Wines Collection - für Weinkeller-Abfragen
        await db.wines.create_index("user_id", background=True)
        await db.wines.create_index([("user_id", 1), ("name", 1)], background=True)
        
        # Public Wines - für Wein-Datenbank Filter
        await db.public_wines.create_index("country", background=True)
        await db.public_wines.create_index([("country", 1), ("region", 1)], background=True)
        
        # Regional Pairings - für Sommelier Kompass
        await db.regional_pairings.create_index("country", background=True)
        await db.regional_pairings.create_index([("country", 1), ("region", 1)], background=True)
        
        # Grape Varieties - für Rebsorten-Lexikon
        await db.grape_varieties.create_index("slug", unique=True, sparse=True, background=True)
        await db.grape_varieties.create_index("name", background=True)
        
        # Feed Posts - für Community Feed
        await db.feed_posts.create_index("user_id", background=True)
        await db.feed_posts.create_index("created_at", background=True)
        
        # Chats - für Chat-History
        await db.chats.create_index("user_id", background=True)
        
        print("   ✅ Alle Indizes erstellt/verifiziert")
    except Exception as e:
        print(f"   ⚠️ Index-Erstellung: {e}")
    
    # ===================================================================
    # KRITISCH: User-Daten aus Backup wiederherstellen wenn DB leer ist!
    # Dies stellt sicher, dass bei einem Deployment die Daten erhalten bleiben.
    # ===================================================================
    print("\n🔐 PRÜFE USER-DATEN...")
    
    user_data_collections = [
        ('users', 'users.json'),
        ('wines', 'wines.json'),
        ('pairings', 'pairings.json'),
        ('chats', 'chats.json'),
        ('wine_favorites', 'wine_favorites.json'),
        ('payment_transactions', 'payment_transactions.json'),
        ('regional_pairings', 'regional_pairings.json'),  # Sommelier-Kompass Gerichte
    ]
    
    for collection_name, backup_file in user_data_collections:
        try:
            count = await db[collection_name].count_documents({})
            backup_path = ROOT_DIR / "data" / backup_file
            
            if count == 0 and backup_path.exists():
                # Collection ist leer aber Backup existiert -> WIEDERHERSTELLEN!
                with open(backup_path, 'r', encoding='utf-8') as f:
                    backup_data = json.load(f)
                
                if backup_data and len(backup_data) > 0:
                    await db[collection_name].insert_many(backup_data)
                    print(f"   🔄 {collection_name}: {len(backup_data)} Einträge aus Backup wiederhergestellt!")
                else:
                    print(f"   ⚠️ {collection_name}: Backup leer")
            elif count > 0:
                print(f"   ✅ {collection_name}: {count} Einträge vorhanden")
            else:
                print(f"   ⚠️ {collection_name}: leer (kein Backup vorhanden)")
        except Exception as e:
            print(f"   ❌ {collection_name}: Fehler - {e}")
    
    print()
    
    # Initialisiere Backup-Manager mit automatischem Backup-Task
    backup_manager = await create_startup_backup(db, ROOT_DIR / "data")
    print("📦 Backup-Manager initialisiert (Auto-Backup alle 6 Stunden)")
    
    # ===================================================================
    # WEEKLY TIP: Prüfe und generiere wöchentlichen Tipp
    # ===================================================================
    print("\n💡 PRÜFE WOCHEN-TIPPS...")
    try:
        import datetime as dt_module
        today = dt_module.date.today()
        current_week = today.isocalendar()[1]
        current_year = today.year
        
        # Prüfe ob diese Woche schon ein Tipp existiert
        tip_exists = await db.weekly_tips.find_one({
            "week_number": current_week,
            "year": current_year
        })
        
        tip_count = await db.weekly_tips.count_documents({})
        
        if tip_count == 0:
            # Keine Tipps vorhanden - Seed initial tips
            print("   📝 Keine Tipps vorhanden - erstelle initiale Tipps...")
            now = dt_module.datetime.now(timezone.utc)
            initial_tips = [
                {
                    "id": str(uuid.uuid4()),
                    "dish": "Scharfes Thai-Curry",
                    "dish_emoji": "🍛",
                    "wine": "Gewürztraminer Spätlese",
                    "wine_type": "weiss",
                    "region": "Elsass, Frankreich",
                    "why": "Die exotische Süße zähmt die Schärfe perfekt – ein echtes Aha-Erlebnis!",
                    "fun_fact": "Der Gewürztraminer hat seinen Namen von den intensiven Gewürzaromen wie Litschi und Rose.",
                    "week_number": current_week,
                    "year": current_year,
                    "language": "de",
                    "is_active": True,
                    "created_at": now
                },
                {
                    "id": str(uuid.uuid4()),
                    "dish": "Pasta Carbonara",
                    "dish_emoji": "🍝",
                    "wine": "Pinot Grigio",
                    "wine_type": "weiss",
                    "region": "Friaul, Italien",
                    "why": "Frische trifft auf Cremigkeit – die knackige Säure schneidet durch die reichhaltige Sauce.",
                    "fun_fact": "In Italien wird Carbonara traditionell NUR mit Guanciale (Schweinebacke) und Pecorino gemacht.",
                    "week_number": current_week - 1 if current_week > 1 else 52,
                    "year": current_year if current_week > 1 else current_year - 1,
                    "language": "de",
                    "is_active": True,
                    "created_at": now - timedelta(days=7)
                },
                {
                    "id": str(uuid.uuid4()),
                    "dish": "Dunkle Schokolade",
                    "dish_emoji": "🍫",
                    "wine": "Kräftiger Syrah",
                    "wine_type": "rot",
                    "region": "Rhône-Tal, Frankreich",
                    "why": "Ein Dessert-Traum – die dunklen Beerenaromen des Syrah umarmen die Bitterkeit der Schokolade.",
                    "fun_fact": "Syrah und Shiraz sind derselbe Wein! In Frankreich heißt er Syrah, in Australien Shiraz.",
                    "week_number": current_week - 2 if current_week > 2 else 52 - (2 - current_week),
                    "year": current_year if current_week > 2 else current_year - 1,
                    "language": "de",
                    "is_active": True,
                    "created_at": now - timedelta(days=14)
                },
                {
                    "id": str(uuid.uuid4()),
                    "dish": "BBQ Spare Ribs",
                    "dish_emoji": "🍖",
                    "wine": "Zinfandel",
                    "wine_type": "rot",
                    "region": "Kalifornien, USA",
                    "why": "Rauch trifft Rauch – die würzigen Brombeer-Noten ergänzen die süß-rauchige BBQ-Sauce perfekt.",
                    "fun_fact": "Zinfandel ist genetisch identisch mit der italienischen Primitivo-Traube!",
                    "week_number": current_week - 3 if current_week > 3 else 52 - (3 - current_week),
                    "year": current_year if current_week > 3 else current_year - 1,
                    "language": "de",
                    "is_active": True,
                    "created_at": now - timedelta(days=21)
                }
            ]
            await db.weekly_tips.insert_many(initial_tips)
            print(f"   ✅ {len(initial_tips)} initiale Wochen-Tipps erstellt")
        elif tip_exists:
            print(f"   ✅ Wochen-Tipp für KW {current_week}/{current_year} existiert bereits")
            print(f"   📊 Gesamt: {tip_count} Tipps im Archiv")
        else:
            print(f"   ⚠️ Kein Tipp für KW {current_week}/{current_year} - wird bei nächstem API-Call generiert")
            print(f"   📊 Gesamt: {tip_count} Tipps im Archiv")
            
        # Index für Weekly Tips
        await db.weekly_tips.create_index([("week_number", -1), ("year", -1)])
        await db.weekly_tips.create_index("created_at")
        
    except Exception as e:
        print(f"   ⚠️ Weekly Tip Check: {e}")
    
    # ===================================================================
    # WICHTIG: Datenbank-Indizes für Performance erstellen
    # ===================================================================
    print("\n🔧 ERSTELLE DATENBANK-INDIZES...")
    try:
        # Index für Weinkeller - KRITISCH für Multi-User-Skalierung
        await db.wines.create_index("user_id")
        print("   ✅ Index 'user_id' auf 'wines' Collection erstellt")
        
        # Index für User-Suche
        await db.users.create_index("user_id", unique=True)
        await db.users.create_index("email", unique=True)
        print("   ✅ Index 'user_id' und 'email' auf 'users' Collection erstellt")
        
    except Exception as e:
        print(f"   ⚠️ Index-Erstellung: {e} (Index existiert möglicherweise bereits)")
    
    # Lade das Backup-Manifest für erwartete Werte
    manifest_path = ROOT_DIR / "data" / "backup_manifest.json"
    expected = {
        'blog_posts': 233,
        'public_wines': 1751,
        'grape_varieties': 140,
        'regional_pairings': 44,
        'dishes': 40,
        'feed_posts': 268,
        'wine_database': 494,
        'seo_pairings': 500,
        'wines': 11
    }
    
    if manifest_path.exists():
        try:
            with open(manifest_path, 'r') as f:
                manifest = json.load(f)
                expected = manifest.get('expected', expected)
                print(f"📋 Manifest Version: {manifest.get('version', 'unknown')}")
                print(f"   Erstellt: {manifest.get('timestamp', 'unknown')}")
        except:
            pass
    
    # Prüfe aktuelle Datenbank
    print("\n🔍 DATENBANK-PRÜFUNG:")
    
    checks = {}
    needs_reseed = False
    
    # Blog-Posts mit Kategorie-Check
    blog_count = await db.blog_posts.count_documents({})
    regionen_count = await db.blog_posts.count_documents({"category": "regionen"})
    checks['blog_posts'] = blog_count
    if blog_count < expected['blog_posts'] or regionen_count < 80:
        needs_reseed = True
        print(f"   ❌ blog_posts: {blog_count}/{expected['blog_posts']} (regionen: {regionen_count}/84)")
    else:
        print(f"   ✅ blog_posts: {blog_count}/{expected['blog_posts']}")
    
    # Alle anderen Collections (System-Daten)
    for col_name in ['public_wines', 'grape_varieties', 'regional_pairings', 'dishes', 'feed_posts', 'wine_database', 'seo_pairings']:
        count = await db[col_name].count_documents({})
        checks[col_name] = count
        exp = expected.get(col_name, 0)
        
        # Toleranz von 5% nach unten erlaubt
        min_expected = int(exp * 0.95)
        
        if count < min_expected:
            needs_reseed = True
            print(f"   ❌ {col_name}: {count}/{exp}")
        else:
            print(f"   ✅ {col_name}: {count}/{exp}")
    
    # User-Collections separat prüfen (triggern KEIN Reseed)
    # KRITISCH: Diese Collections werden NIEMALS überschrieben!
    wines_count = await db.wines.count_documents({})
    users_count = await db.users.count_documents({})
    print(f"   🔒 wines (User-Keller): {wines_count} Flaschen")
    print(f"   🔒 users (Benutzerkonten): {users_count} Konten")
    
    # Wenn IRGENDETWAS fehlt -> Prüfe und lade nur LEERE Collections
    if needs_reseed:
        print("\n" + "=" * 60)
        print("📦 DATENBANK-CHECK - NUR LEERE COLLECTIONS WERDEN GEFÜLLT")
        print("=" * 60)
        
        # ALLE Collections die geschützt werden sollen
        # NIEMALS werden existierende Daten überschrieben!
        all_protected_collections = [
            # Content-Collections (Weindatenbank, Blogs, etc.)
            ("regional_pairings", "regional_pairings.json"),   # Sommelier-Kompass
            ("grape_varieties", "grape_varieties.json"),       # Rebsorten
            ("blog_posts", "blog_posts.json"),                 # Blogs
            ("dishes", "dishes.json"),                         # Gerichte
            ("wine_database", "wine_database.json"),           # Wein-Datenbank
            ("public_wines", "public_wines.json"),             # Öffentliche Weine
            ("feed_posts", "feed_posts.json"),                 # Community Feed
            ("seo_pairings", "seo_pairings.json"),             # SEO-Pairings
            # User-Collections
            ("wines", "wines.json"),                           # Persönlicher Weinkeller
            ("users", "users.json"),                           # Benutzerkonten
            ("pairings", "pairings.json"),                     # Pairing-History
            ("chats", "chats.json"),                           # Chat-Verläufe
            ("wine_favorites", "wine_favorites.json"),         # Favoriten
            ("payment_transactions", "payment_transactions.json"),  # Zahlungen
            # System-Collections
            ("coupons", "coupons.json"),                       # Gutschein-Codes
        ]
        
        for collection_name, json_filename in all_protected_collections:
            try:
                # KRITISCH: Prüfe ob Collection bereits Daten hat
                existing_count = await db[collection_name].count_documents({})
                
                if existing_count > 0:
                    # Collection hat Daten -> NIEMALS überschreiben!
                    print(f"   🔒 {collection_name}: {existing_count} Dokumente GESCHÜTZT")
                    continue
                
                # Collection ist leer -> Aus Backup laden
                data_file = ROOT_DIR / "data" / json_filename
                
                if data_file.exists():
                    with open(data_file, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                    
                    if data:
                        await db[collection_name].insert_many(data)
                        print(f"   ✅ {collection_name}: {len(data)} Dokumente aus Backup geladen")
                    else:
                        print(f"   ⚠️ {json_filename} ist leer - Collection bleibt leer")
                else:
                    print(f"   ⚠️ Backup-Datei fehlt: {json_filename} - Collection bleibt leer")
                    
            except Exception as e:
                print(f"   ❌ {collection_name}: Fehler - {e}")
    else:
        print("\n✅ Alle Daten sind korrekt - kein Seeding nötig")
    
    # Finale Verifizierung
    print("\n" + "=" * 60)
    print("📊 FINALE VERIFIZIERUNG:")
    final_blog = await db.blog_posts.count_documents({})
    final_regionen = await db.blog_posts.count_documents({"category": "regionen"})
    final_pairings = await db.regional_pairings.count_documents({})
    final_grapes = await db.grape_varieties.count_documents({})
    final_wines = await db.public_wines.count_documents({})
    final_cellar = await db.wines.count_documents({})
    final_seo = await db.seo_pairings.count_documents({})
    
    print(f"   Blogs: {final_blog} (Regionen: {final_regionen})")
    print(f"   Regional Pairings: {final_pairings}")
    print(f"   Rebsorten: {final_grapes}")
    print(f"   Weine (DB): {final_wines}")
    print(f"   Weinkeller: {final_cellar}")
    print(f"   SEO Pairings: {final_seo}")
    
    # ===================================================================
    # AGENT HANDOFF: Aktualisiere Statistiken für nächsten Agenten
    # ===================================================================
    try:
        handoff_path = ROOT_DIR.parent / "AGENT_HANDOFF.md"
        if handoff_path.exists():
            import re
            from datetime import datetime, timezone as tz
            
            with open(handoff_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Aktualisiere Datum
            now = datetime.now(tz.utc).strftime('%d.%m.%Y %H:%M UTC')
            content = re.sub(
                r'\*Letzte Aktualisierung:.*\*',
                f'*Letzte Aktualisierung: {now}*',
                content
            )
            
            with open(handoff_path, 'w', encoding='utf-8') as f:
                f.write(content)
            print("\n   📋 AGENT_HANDOFF.md aktualisiert")
    except Exception as e:
        print(f"\n   ⚠️ AGENT_HANDOFF.md Update fehlgeschlagen: {e}")
    
    print("=" * 60)
    print("🍷 SERVER BEREIT!")
    print("=" * 60 + "\n")

@app.on_event("shutdown")
async def shutdown_db_client():
    # Stoppe automatischen Backup-Task
    await stop_backup_task()
    # Schließe MongoDB-Verbindung
    client.close()
